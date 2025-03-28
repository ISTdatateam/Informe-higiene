#!/usr/bin/env python3
import logging
import re
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from pythermalcomfort.models import pmv_ppd_iso
import qrcode
from PIL import ImageOps  # Asegúrate de tener Pillow instalado
from io import BytesIO
from docx.oxml.ns import qn
from datetime import datetime
from docx.enum.table import WD_ALIGN_VERTICAL

# Configuración básica del logging
# logging.basicConfig(level=logging.INFO)
logging.basicConfig(level=logging.DEBUG)


# -----------------------------------------------
# Nuevas funciones
# -----------------------------------------------


def calcular_analisis_area(group):
    """Determina si un área cumple o no basado en mediciones."""
    if len(group) > 1:
        # Cálculo de promedios
        avg_t_bul = group["t_bul_seco"].astype(float).mean()
        avg_t_globo = group["t_globo"].astype(float).mean()
        avg_hum = group["hum_rel"].astype(float).mean()
        avg_vel = group["vel_air"].astype(float).mean()
        avg_met = group["met"].astype(float).mean() or 1.1
        avg_clo = group["clo"].astype(float).mean() or 0.5

        # Cálculo PMV/PPD
        try:
            results = pmv_ppd_iso(
                tdb=avg_t_bul,
                tr=avg_t_globo,
                vr=avg_vel,
                rh=avg_hum,
                met=avg_met,
                clo=avg_clo,
                model="7730-2005",
                limit_inputs=False
            )
            pmv = results.pmv if hasattr(results, 'pmv') else results.get('pmv', 0)
            analisis = interpret_pmv(pmv)
        except Exception:
            analisis = "NO CUMPLE"
    else:
        analisis = group.iloc[0].get("resultado_medicion", "NO CUMPLE").upper()

    return analisis


def procesar_areas(df_mediciones):
    """Procesa todas las áreas para determinar cumplimiento."""
    areas_cumplen = []
    areas_no_cumplen = []

    if not df_mediciones.empty:
        grouped = df_mediciones.groupby("nombre_area")
        for area, group in grouped:
            analisis = calcular_analisis_area(group)
            if analisis == "CUMPLE":
                areas_cumplen.append(area)
            else:
                areas_no_cumplen.append(area)
    return areas_cumplen, areas_no_cumplen


# -----------------------------------------------
# FUNCIONES AUXILIARES PARA EL DOCUMENTO WORD
# -----------------------------------------------


def format_columns(df, columns, mode="title"):
    """
    Formatea las columnas especificadas en un DataFrame eliminando espacios al inicio y al final,
    y transformando el texto según el modo seleccionado.

    Parámetros:
        df (DataFrame): El DataFrame a modificar.
        columns (list): Lista de nombres de columnas a formatear.
        mode (str): Modo de transformación, puede ser:
            - "title": Convierte el texto a Título (cada palabra con su primera letra en mayúscula).
            - "capitalize": Solo la primera letra de la primera palabra en mayúscula.
            - "upper": Todo el texto en mayúsculas.
            (El valor por defecto es "title".)

    Retorna:
        DataFrame: El DataFrame con las columnas formateadas.
    """
    for col in columns:
        if col in df.columns:
            # Convertir a cadena y eliminar espacios al inicio y final
            df[col] = df[col].astype(str).str.strip()
            # Aplicar la transformación según el modo
            if mode == "title":
                df[col] = df[col].str.lower().str.title()
            elif mode == "capitalize":
                df[col] = df[col].str.lower().str.capitalize()
            elif mode == "upper":
                df[col] = df[col].str.upper()
            else:
                raise ValueError(f"Modo '{mode}' no reconocido. Use 'title', 'capitalize' o 'upper'.")
    return df


def set_column_width(table, col_index, width):
    """
    Establece el ancho de la columna 'col_index' de la tabla 'table' al valor 'width' (un objeto de docx.shared, por ejemplo Cm(4)).
    """
    # width._value retorna el ancho en unidades dxa (1 dxa = 1/20 de punto)
    for cell in table.columns[col_index].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width.inches * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)


def merge_column_cells(tabla, col_index, start_row=1):
    """
    Fusiona celdas consecutivas en la columna 'col_index' de la tabla 'tabla'
    a partir de la fila 'start_row' (por defecto se salta el encabezado).
    Se fusionan las celdas cuando se encuentran celdas vacías, es decir, se asume
    que solo la primera celda del grupo contiene el valor.
    """
    num_filas = len(tabla.rows)
    current_value = None
    start_index = None

    for i in range(start_row, num_filas):
        cell = tabla.cell(i, col_index)
        text = cell.text.strip()
        # Si la celda tiene contenido, se considera el inicio de un grupo
        if text:
            # Si ya había un grupo iniciado y abarca más de una fila, se fusionan las celdas anteriores
            if current_value is not None and start_index is not None and i - start_index > 1:
                primera_celda = tabla.cell(start_index, col_index)
                for j in range(start_index + 1, i):
                    primera_celda = primera_celda.merge(tabla.cell(j, col_index))
            current_value = text
            start_index = i
    # Fusionar el último grupo pendiente (si abarca más de una fila)
    if current_value is not None and start_index is not None and num_filas - start_index > 1:
        primera_celda = tabla.cell(start_index, col_index)
        for j in range(start_index + 1, num_filas):
            primera_celda = primera_celda.merge(tabla.cell(j, col_index))


def generate_qr_code(url, border=10, box_size=2):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=box_size,
        border=2  # Borde interno del QR
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
    # Agregar un borde (padding) blanco alrededor de la imagen
    padded_img = ImageOps.expand(img, border=border, fill='white')
    bio = BytesIO()
    padded_img.save(bio, format="PNG")
    bio.seek(0)
    return bio


def join_with_and(items):
    """
    Une los elementos de la lista usando la conjunción "y" para el último elemento.
    Ejemplos:
      - ['Oficina']             => "Oficina"
      - ['Oficina', 'Sala']     => "Oficina y Sala"
      - ['Oficina', 'Sala', 'Bodega'] => "Oficina, Sala y Bodega"
    """
    items = list(items)  # Convertir a lista para evitar ambigüedades
    if not items:
        return ""
    elif len(items) == 1:
        return items[0]
    elif len(items) == 2:
        return f"{items[0]} y {items[1]}"
    else:
        return ", ".join(items[:-1]) + " y " + items[-1]


def get_or_add_lang(rPr):
    """
    Obtiene o agrega el elemento de idioma en la configuración de la fuente.
    """
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    return lang


def interpret_pmv(pmv_value):
    if pmv_value >= 1:
        return "NO CUMPLE"
    elif pmv_value > -1:
        return "CUMPLE"
    else:
        return "NO CUMPLE"


def descargar_imagen_gdrive(url_foto: str) -> BytesIO or None:
    """
    Dada una URL de Google Drive, intenta descargar el archivo y devolverlo como BytesIO.
    Retorna None si no se puede descargar o no tiene permisos adecuados.
    """
    match = re.search(r"id=(.*)", url_foto)
    if not match:
        return None
    file_id = match.group(1)
    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
    try:
        response = requests.get(download_url)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            return None
    except Exception as e:
        return None


# -----------------------------------------------------
# NUEVAS FUNCIONES PARA CONFIGURAR ESTILOS DE MANERA MODULAR
# -----------------------------------------------------

def set_style_language(style, lang_code="es-CL"):
    """
    Asigna el código de idioma a un estilo dado.
    Se configura para el texto normal, Asia Oriental y bidireccional.
    """
    rPr = style.element.get_or_add_rPr()
    lang = get_or_add_lang(rPr)
    lang.set(qn('w:val'), lang_code)
    lang.set(qn('w:eastAsia'), lang_code)
    lang.set(qn('w:bidi'), lang_code)


def apply_style_properties(style, properties):
    """
    Aplica las propiedades definidas en el diccionario 'properties' al estilo 'style'.
    Las propiedades pueden incluir: font_name, font_size, font_color, bold, italic,
    space_before, space_after, alignment y lang_code.
    """
    font = style.font
    font.name = properties.get("font_name", font.name)
    font.size = properties.get("font_size", font.size)
    font.color.rgb = properties.get("font_color", font.color.rgb)
    font.bold = properties.get("bold", font.bold)
    font.italic = properties.get("italic", font.italic)

    para_format = style.paragraph_format
    para_format.space_before = properties.get("space_before", para_format.space_before)
    para_format.space_after = properties.get("space_after", para_format.space_after)
    para_format.alignment = properties.get("alignment", para_format.alignment)

    # Aplicar idioma
    set_style_language(style, properties.get("lang_code", "es-CL"))


# Diccionario de configuración de estilos
from docx.enum.text import WD_ALIGN_PARAGRAPH

style_configurations = {
    "Normal": {
        "font_name": "Calibri",
        "font_size": Pt(10),
        "font_color": RGBColor(0x00, 0x00, 0x00),
        "bold": False,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "lang_code": "es-CL"
    },
    "Heading 1": {
        "font_name": "Calibri",
        "font_size": Pt(14),
        "font_color": RGBColor(0, 0, 0),
        "bold": True,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(12),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "lang_code": "es-CL"
    },
    "Heading 2": {
        "font_name": "Calibri",
        "font_size": Pt(12),
        "font_color": RGBColor(79, 11, 123),
        "bold": True,
        "italic": False,
        "space_before": Pt(12),
        "space_after": Pt(12),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "lang_code": "es-CL"
    },
    "Heading 3": {
        "font_name": "Calibri",
        "font_size": Pt(11),
        "font_color": RGBColor(0x30, 0x30, 0x30),
        "bold": True,
        "italic": False,
        "space_before": Pt(8),
        "space_after": Pt(8),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "lang_code": "es-CL"
    },
    "TablaTexto": {  # Estilo personalizado para el contenido de las tablas
        "font_name": "Calibri",
        "font_size": Pt(10),
        "font_color": RGBColor(0, 0, 0),
        "bold": False,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "lang_code": "es-CL"
    },
    "Centrado": {
        "font_name": "Calibri",
        "font_size": Pt(10),
        "font_color": RGBColor(0, 0, 0),
        "bold": False,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "lang_code": "es-CL"
    },
    "Centrado Bold": {
        "font_name": "Calibri",
        "font_size": Pt(10),
        "font_color": RGBColor(0, 0, 0),
        "bold": True,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "lang_code": "es-CL"
    }
}


def look_informe(doc: Document):
    """
    Configura el documento: establece orientación, márgenes, estilos y el idioma.
    """
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Aplicar configuración a los estilos definidos en nuestro diccionario
    for style_name, properties in style_configurations.items():
        try:
            style = doc.styles[style_name]
            apply_style_properties(style, properties)
        except KeyError:
            # Si el estilo no existe, se continúa
            continue



def set_vertical_alignment(doc: Document, section_index=0, alignment='top'):
    """
    Ajusta la alineación vertical de la sección especificada.
    """
    section = doc.sections[section_index]
    sectPr = section._sectPr
    vAlign = sectPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        sectPr.append(vAlign)
    alignment = alignment.lower()
    if alignment not in ['top', 'center', 'both']:
        alignment = 'top'
    vAlign.set(qn('w:val'), alignment)


# Función auxiliar para poner en negrita todas las celdas de una fila
def set_row_bold(row):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def apply_style(cell, shading_color="4f0b7b"):
    """
    Aplica un formato de fondo (con el color indicado), negrita y tamaño de fuente a una celda.
    """
    # Crear el elemento de sombreado con el color especificado
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shading_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    # Para cada párrafo y cada "run" dentro de la celda, aplicar color de fuente blanco, negrita y tamaño 12
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)


def format_row(row, shading_color="4f0b7b"):
    """
    Aplica el formato de estilo (fondo, negrita, color de fuente y tamaño) a todas las celdas de una fila.
    """
    for cell in row.cells:
        apply_style(cell, shading_color)


def add_row(table, label, value="", first=False):
    """
    Agrega una fila a la tabla.

    Si 'value' es una cadena vacía, se interpreta como una fila cabecera;
    se realiza un merge de la primera y segunda celda, se asigna el texto 'label'
    y se aplica el estilo (usando apply_style).

    Si 'value' no es vacío, se asigna 'label' a la primera celda y 'value' a la segunda.
    Además, si se especifica first=True, se aplica el mismo estilo a todas las celdas de la fila.
    """
    row = table.add_row()
    cells = row.cells
    if value == "":
        # Fila cabecera: se fusionan las dos primeras celdas y se formatea el resultado.
        merged_cell = cells[0].merge(cells[1])
        merged_cell.text = label
        apply_style(merged_cell)
    else:
        # Fila normal: se asigna el texto a la primera y segunda celda.
        cells[0].text = label
        cells[1].text = "" if (value is None or (isinstance(value, str) and value.strip() == "")) else str(value)
        if first:
            format_row(row)


def agregar_contenido(cell, items):
    """
    Agrega cada elemento de 'items' a la celda como párrafo y, posteriormente,
    elimina el primer párrafo si está vacío y limpia la celda para eliminar
    espacios y saltos de línea extra, de modo que solo se conserven los párrafos
    con contenido.
    """
    # 1) Agregar cada elemento como párrafo
    for item in items:
        cell.add_paragraph(item)

    if cell.paragraphs:
        parrafo_a_eliminar = cell.paragraphs[0]
        # Accede al elemento XML y elimínalo del padre
        p_element = parrafo_a_eliminar._element
        p_element.getparent().remove(p_element)


####
# Añade al final de doc_utils.py

def generar_recomendaciones(pmv, tdb_initial, tr_initial, vr, rh, met, clo):
    recomendaciones = []

    # Calcular diferencias térmicas importantes
    dif_temp = tr_initial - tdb_initial

    # 1. Estrategias principales según condición térmica
    estrategias_base = []

    if pmv > 1.0:  # Ambiente caluroso
        estrategias_base = [
            {
                'condicion': True,  # Siempre aplica para calor
                'tipo': 'enfriamiento',
                'mensaje': "Refrigeración activa requerida",
                'acciones': [
                    "- Implementar [__] equipos enfriadores de aire (según las dimensiones de las áreas), con el fin de enfriar el aire.",
                    "- Consultar con el proveedor el óptimo uso del equipo por ejemplo: periodicidad de suministrar agua helada, hielo o implemento refrigerante autorizado para el equipamiento adquirido, con el fin de estar constantemente enfriando durante toda la jornada laboral el área, especialmente en periodo de mayor temperaturas o época estival.",
                    "- Llevar una Bitácora o Registro de la actividad en lo referido al uso de enfriador(es)."
                ],
                'plazo': '6 meses desde la recepción del presente informe técnico'
            },
            {
                'condicion': dif_temp > 2.0,
                'tipo': 'aislamiento',
                'mensaje': "Reducción de carga térmica",
                'acciones': [
                    "- (REVISAR SI CORRESPONDE) Instalar materiales aislantes en techos/paredes",
                    "- (REVISAR SI CORRESPONDE) Implementar protecciones solares reflectivas",
                    "- (REVISAR SI CORRESPONDE) Aislar fuentes de calor radiante"
                ],
                'plazo': '3 meses desde la recepción del presente informe técnico'
            },
            {
                'condicion': vr < 0.2,
                'tipo': 'ventilacion',
                'mensaje': f"Aumentar ventilación",
                'acciones': [
                    "- Implementar [__] ventiladores industriales, con el fin de generar corrientes de aire, las cuales ayudarán a mejorar condiciones de confort térmico en dicha área.",
                    "- (REVISAR SI CORRESPONDE) Implementar sistemas de extracción forzada",
                    "- (REVISAR SI CORRESPONDE) Optimizar ventilación cruzada"
                ],
                'plazo': '3 meses desde la recepción del presente informe técnico'
            }
        ]
    '''
    #NOTA:
    #Se desactivan las recomendaciones por frio, dado que el foco de las evaluaciones son por calor en esta temporada.
    #NOTA
    
    
    elif pmv < -1.0:  # Ambiente frío
        estrategias_base = [
            {
                'condicion': True,
                'tipo': 'calefaccion',
                'mensaje': "Protección contra el frío",
                'acciones': [
                    "- (REVISAR SI CORRESPONDE) Implementar sistemas de calefacción radiante",
                    "- (REVISAR SI CORRESPONDE) Mejorar aislamiento térmico en envolvente",
                    "- (REVISAR SI CORRESPONDE) Optimizar sellado de infiltraciones"
                ],
                'plazo': '3 meses desde la recepción del presente informe técnico'
            }
        ]
    '''

    # Filtrar y agregar estrategias válidas
    for estrategia in estrategias_base:
        if estrategia['condicion']:
            recomendaciones.append({
                'tipo': estrategia['tipo'],
                'categoria': estrategia['mensaje'],
                'nivel': 'Prioridad 1',
                'mensaje': estrategia['mensaje'],
                'acciones': estrategia.get('acciones', []),
                'plazo': estrategia.get('plazo', 'Inmediato')
            })

    '''
    # 2. Mantenimiento (siempre aplica)
    mantenimiento = {
        'preventivo': [
            "Programar mantención de equipos con proveedores",
            "Calendario de limpieza de filtros/paneles",
            "Verificación mensual de sistemas"
        ],
        'correctivo': [
            "Reparación de sistemas de ventilación",
            "Ajuste de equipos de climatización",
            "Registro de intervenciones técnicas"
        ],
        'control': [
            f"Regulación térmica (23-26°C) con registro",
            "Monitoreo continuo de parámetros ambientales"
        ]
    }
    recomendaciones.append({
        'tipo': 'mantenimiento',
        'categoria': 'Gestión Técnica',
        'nivel': 'Prioridad 2',
        'mensaje': "Programa de mantenimiento integral",
        'acciones': mantenimiento,
        'plazo': 'Continuo'
    })

    # 3. Medidas administrativas (siempre aplican)
    recomendaciones.append({
        'tipo': 'administrativa',
        'categoria': 'Comunicación',
        'nivel': 'Prioridad 3',
        'acciones': [
            "Informar formalmente a todo el personal sobre riesgos térmicos (DS N°44)",
            "Establecer registros firmados de capacitación"
        ],
        'plazo': '30 días desde la recepción del presente informe técnico'
    })
    '''

    return recomendaciones

def crear_tabla_recomendaciones(doc, tipo_medida, medidas):
    """Crea tabla de recomendaciones por tipo de medida"""
    # Configurar encabezados según el tipo
    headers = ["Área", "Acción correctiva", "Prescripción de medidas", "Plazo"]
    col_widths = [Cm(2), Cm(2), Cm(12), Cm(2)]

    # Crear tabla
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Configurar anchos de columna
    for idx, width in enumerate(col_widths):
        set_column_width(table, idx, width)

    # Formatear cabecera
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    format_row(table.rows[0], shading_color="4F0B7B")  # Morado corporativo

    # Llenar con medidas
    for medida in medidas:
        row_cells = table.add_row().cells
        agregar_contenido(row_cells[0],medida['areas'])
        row_cells[1].text = medida['tipo_medida']
        agregar_contenido(row_cells[2], medida['acciones'])
        row_cells[3].text = medida['plazo']

        # Formato vertical centrado
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table

def agregar_medidas_correctivas(doc, df_mediciones, areas_no_cumplen):
    # 1. Medidas Ingenieriles (solo para áreas no conformes)
    medidas_ingenieriles = []
    if areas_no_cumplen:
        for area in areas_no_cumplen:
            grupo = df_mediciones[df_mediciones['nombre_area'] == area]
            avg_params = {
                'tdb': grupo['t_bul_seco'].mean(),
                'tr': grupo['t_globo'].mean(),
                'vr': grupo['vel_air'].mean(),
                'rh': grupo['hum_rel'].mean(),
                'met': grupo['met'].mean(),
                'clo': grupo['clo'].mean(),
            }

            try:
                results = pmv_ppd_iso(
                    tdb=avg_params['tdb'],
                    tr=avg_params['tr'],
                    vr=avg_params['vr'],
                    rh=avg_params['rh'],
                    met=avg_params['met'],
                    clo=avg_params['clo'],
                    model="7730-2005",
                    limit_inputs=False,
                    round_output=True
                )
            except Exception as e:
                logging.error("Error al calcular pmv_ppd_iso para el área %s: %s", area, e)
                results = None

            if results is not None:
                if isinstance(results, dict):
                    avg_ppd = float(results.get("ppd", 0))
                    avg_pmv = float(results.get("pmv", 0))
                elif hasattr(results, "ppd") and hasattr(results, "pmv"):
                    avg_ppd = float(results.ppd)
                    avg_pmv = float(results.pmv)
                else:
                    avg_ppd = 0
                    avg_pmv = 0
            else:
                avg_ppd = 0
                avg_pmv = 0

            avg_params.update({
                'pmv': avg_pmv,
                'ppd':avg_ppd
            })

            recs = generar_recomendaciones(
                pmv=avg_params['pmv'],
                tdb_initial=avg_params['tdb'],
                tr_initial=avg_params['tr'],
                vr=avg_params['vr'],
                rh=avg_params['rh'],
                met=avg_params['met'],
                clo=avg_params['clo']
            )

            for rec in recs:
                if rec['tipo'] in ['ventilacion', 'enfriamiento', 'aislamiento', 'calefaccion']:
                    medidas_ingenieriles.append({
                        'tipo_medida': rec['categoria'],
                        'areas': [area],
                        'acciones': rec['acciones'],
                        'plazo': rec['plazo']
                    })

    # 2. Medidas Administrativas (siempre se incluyen)
    medidas_administrativas = [
        {
            'tipo_medida': 'Comunicación',
            'areas': ['Todas'],
            'acciones': [
                "- Informar a cada persona trabajadora acerca de los riesgos que entrañan sus labores, de las medidas preventivas, de los métodos y/o procedimientos de trabajo correctos, acorde a lo identificado por la empresa. Además de lo señalado previamente, la entidad empleadora deberá informar de manera oportuna y adecuada el resultado del presente informe técnico.",
                "- Complementario a lo anterior, se deberán realizar capacitaciones (teóricas/prácticas) periódicas en prevención de riesgos laborales, con la finalidad de garantizar el aprendizaje efectivo y eficaz, dejando registro de dichas capacitaciones y evaluaciones.",
                "- Lo señalado previamente se enmarca en los artículos 15° y 16° del Párrafo IV del D.S 44 “Aprueba nuevo reglamento sobre gestión preventiva de los riesgos laborales para un entorno de trabajo seguro y saludable."
            ],
            'plazo': '30 días desde la recepción del presente informe técnico'
        },
        {
            'tipo_medida': 'Mantenimiento',
            'areas': areas_no_cumplen if areas_no_cumplen else ['Todas'],
            'acciones': [
                "- Realizar mantención preventiva en los equipos de climatización, con el fin de identificar desgastes y prevenir fallas. Se debe seguir un cronograma establecido y registrar cada intervención.",
                "- Ejecutar reparaciones en equipos de climatización al detectar fallas en su funcionamiento, restableciendo su operatividad de manera oportuna y registrando las acciones realizadas",
                "- Implementar un monitoreo continuo de los parámetros de confort térmico entre 23 a 26°C en verano y 18 a 21 en invierno, manteniendo un registro sistemático de las mediciones y ajustes efectuados"
            ],
            'plazo': 'Continuo'
        }
    ]

    # Agregar secciones al documento
    doc.add_heading("Medidas de Carácter Ingenieril", level=2)
    if medidas_ingenieriles:
        crear_tabla_recomendaciones(doc, "Ingenieril", medidas_ingenieriles)
    else:
        doc.add_paragraph("No se requieren medidas ingenieriles para este caso.")

    doc.add_heading("Medidas de Carácter Administrativo", level=2)
    crear_tabla_recomendaciones(doc, "Administrativa", medidas_administrativas)

# Función para generar texto de áreas cumplen/no cumplen
def generar_texto_areas(areas, tipo):
    if not areas:
        return ""
    areas_formateadas = join_with_and(areas)
    if tipo == "cumplen":
        return f"las áreas {areas_formateadas} cumplen con el estándar de confort térmico"
    else:
        return f"las áreas {areas_formateadas} no cumplen con el estándar"

######

# -----------------------------------------------
# FUNCIÓN PARA GENERAR EL DOCUMENTO WORD
# -----------------------------------------------
def generar_informe_en_word(df_centros, df_visitas, df_mediciones, df_equipos) -> BytesIO:
    """
    Genera el informe en Word utilizando:
      - df_centros: información del centro de trabajo (tabla higiene_Centros_Trabajo)
      - df_visitas: información de visitas (tabla higiene_Visitas); se selecciona la visita más reciente.
      - df_mediciones: mediciones asociadas a la visita (tabla higiene_Mediciones)
      - df_equipos: información de equipos de medición (tabla higiene_Equipos_Medicion)
    """

    format_columns(df_visitas, ['nombre_personal_visita', 'consultor_ist'], mode="title")
    format_columns(df_visitas, 'cargo_personal_visita', mode="capitalize")
    format_columns(df_mediciones, ['nombre_area', 'sector_especifico','puesto_trabajo'], mode="capitalize")



    doc = Document()
    look_informe(doc)
    set_vertical_alignment(doc, section_index=0, alignment='top')

    # Cabecera con logo
    section = doc.sections[0]
    section.header_distance = Inches(0.5)
    header = section.header
    if header.paragraphs:
        paragraph = header.paragraphs[0]
    else:
        paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.add_picture('IST.jpg', width=Inches(1))

    # Título del informe
    paragraph = doc.add_heading("INFORME EVALUACIÓN CONFORT TÉRMICO", level=1)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    # -------------------------------
    # 1) IDENTIFICACIÓN ACTIVIDAD
    # -------------------------------
    table_inicial = doc.add_table(rows=0, cols=2)
    table_inicial.style = 'Table Grid'

    if not df_centros.empty:
        row_centro = df_centros.iloc[0]
        # Información de la empresa
        add_row(table_inicial, "Información empresa")
        add_row(table_inicial, "Razón Social", row_centro.get('razon_social', ''))
        add_row(table_inicial, "RUT", row_centro.get('rut', ''))
        # Información centro de trabajo
        add_row(table_inicial, "Información centro de trabajo")
        add_row(table_inicial, "CUV", row_centro.get('cuv', ''))
        add_row(table_inicial, "Nombre de Local", row_centro.get('nombre_ct', ''))
        direccion_ct = row_centro.get('direccion_ct', '')
        comuna_ct = row_centro.get('comuna_ct', '')
        add_row(table_inicial, "Dirección", f"{direccion_ct}, {comuna_ct}")
    else:
        add_row(table_inicial, "Información de empresa/centro", "No se encontró información para este CUV.")


    # Procesar áreas antes del resumen
    areas_cumplen, areas_no_cumplen = procesar_areas(df_mediciones)

    # Encabezado principal del contenido: Resumen
    doc.add_paragraph()
    paragraph = doc.add_heading("Resumen", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def generar_texto_areas(areas, condicion):
        """
        Genera el texto correspondiente a las áreas evaluadas según la condición recibida.
        Si 'condicion' contiene la palabra 'no', se genera un mensaje para áreas que no cumplen.
        """
        areas_str = ", ".join(areas)
        if "no" in condicion.lower():
            if len(areas) == 1:
                return f"el área de {areas_str} NO cumple con el estándar de confort térmico"
            else:
                return f"las áreas {areas_str}, éstas NO cumplen con el estándar confort térmico"
        else:
            if len(areas) == 1:
                return f"el área de {areas_str} cumple con el estándar confort térmico"
            else:
                return f"las áreas {areas_str}, éstas cumplen con el estándar confort térmico"

    if not df_visitas.empty:

        nombre_ct = df_centros.iloc[0].get("nombre_ct", "") if not df_centros.empty else ""
        texto_base = f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, "

        # Escenario 1: Todas cumplen
        if areas_cumplen and not areas_no_cumplen:
            forma_cumple = "cumple" if len(areas_cumplen) == 1 else "cumplen"
            texto_areas = generar_texto_areas(areas_cumplen, forma_cumple)
            connector = (
                ", por lo que se debe mantener la condición actual o similar."
                if len(areas_cumplen) == 1
                else ", por lo que se deben mantener las condiciones actuales o similares."
            )
            doc.add_paragraph(texto_base + texto_areas + connector)

        # Escenario 2: Algunas cumplen, otras no
        elif areas_cumplen and areas_no_cumplen:
            forma_cumple = "cumple" if len(areas_cumplen) == 1 else "cumplen"
            forma_no_cumple = "no cumple" if len(areas_no_cumplen) == 1 else "no cumplen"
            texto_cumplen = generar_texto_areas(areas_cumplen, forma_cumple)
            texto_no_cumplen = generar_texto_areas(areas_no_cumplen, forma_no_cumple)

            texto_conclusion = texto_base + f"es posible concluir que {texto_cumplen}."
            connector_medidas = (
                f"Respecto de {texto_no_cumplen}, se debe adoptar la medida prescrita para su solución."
                if len(areas_no_cumplen) == 1
                else f"Respecto de {texto_no_cumplen}, se deben adoptar las medidas prescritas para su solución."
            )
            doc.add_paragraph(texto_conclusion)
            doc.add_paragraph(connector_medidas)

        # Escenario 3: Ninguna cumple
        elif areas_no_cumplen and not areas_cumplen:
            forma_no_cumple = "no cumple" if len(areas_no_cumplen) == 1 else "no cumplen"
            texto_areas = generar_texto_areas(areas_no_cumplen, forma_no_cumple)
            connector = (
                ", por lo que se debe adoptar la medida prescrita para su solución."
                if len(areas_no_cumplen) == 1
                else ", por lo que se deben adoptar las medidas prescritas para su solución."
            )
            doc.add_paragraph(texto_base + texto_areas + connector)

        else:
            doc.add_paragraph("No se encontraron áreas evaluadas para generar un resumen.")

    # -------------------------------
    # 2) ANTECEDENTES DE LA ACTIVIDAD
    # -------------------------------
    if not df_visitas.empty:
        row_visita = df_visitas.iloc[0]
        doc.add_paragraph()
        paragraph = doc.add_heading("Antecedentes de la actividad", level=2)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Se extrae información del centro para complementar el texto
        razon_social = row_centro.get("razon_social", "") if not df_centros.empty else ""
        nombre_ct = row_centro.get("nombre_ct", "") if not df_centros.empty else ""
        direccion_ct = row_centro.get("direccion_ct", "")
        comuna_ct = row_centro.get("comuna_ct", "")
        direccion_completa = f"{direccion_ct}, {comuna_ct}"
        fecha_visita = str(row_visita.get("fecha_visita", ""))
        if fecha_visita:
            fecha_obj = datetime.strptime(fecha_visita, "%Y-%m-%d")  # Convertir a objeto datetime
            fecha_visita = fecha_obj.strftime("%d-%m-%Y")

        hora_visita = str(row_visita.get("hora_visita", ""))
        if hora_visita:
            hora_obj = datetime.strptime(hora_visita, "%H:%M:%S")
            hora_visita = hora_obj.strftime("%H:%M")
        personal_visita = row_visita.get("nombre_personal_visita", "")
        cargo_visita = row_visita.get("cargo_personal_visita", "")
        consultor_ist = row_visita.get("consultor_ist", "")
        temperatura = row_visita.get("temperatura_dia","")
        doc.add_paragraph(
            f"A solicitud de {razon_social} se realizó una evaluación de confort térmico en el centro de trabajo {nombre_ct}, ubicado en {direccion_completa}.")
        doc.add_paragraph()
        doc.add_paragraph(
            f"La visita se efectuó el {fecha_visita} a las {hora_visita} a cargo de {consultor_ist}, profesional consultor/a de IST. En representación de la empresa adherente la visita contó con la participación de {personal_visita} ({cargo_visita}). Cabe destacar que la temperatura máxima registrada durante la jornada fue de {temperatura}°C.")
        doc.add_paragraph()
        # --- Agregar áreas evaluadas ---
        if not df_mediciones.empty:
            areas = df_mediciones["nombre_area"].unique()
            if len(areas) == 1:
                doc.add_paragraph(f"El área evaluada fue {join_with_and(areas)}.")
            else:
                doc.add_paragraph(f"Las áreas evaluadas fueron {join_with_and(areas)}.")
        # --- FIN DE AGREGAR ---
    else:
        doc.add_paragraph("No se encontró información de visita para este CUV.")

    # -------------------------------
    # 3) METODOLOGÍA DE LAS MEDICIONES Y EVALUACIONES
    # -------------------------------
    doc.add_paragraph()
    paragraph = doc.add_heading("Metodología de las mediciones y evaluaciones", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # --- Agregar detalles de equipos utilizando datos de df_equipos ---
    if not df_visitas.empty and not df_equipos.empty:
        row_visita = df_visitas.iloc[0]
        # Obtener equipo de temperatura
        equipo_temp_cod = row_visita.get('equipo_temp', '')
        equipo_temp_detalles = df_equipos[df_equipos['id_equipo'] == equipo_temp_cod]
        if not equipo_temp_detalles.empty:
            eq_temp = equipo_temp_detalles.iloc[0]
            equipo_temp_text = f"{eq_temp.get('nombre_equipo', '')} {eq_temp.get('marca_equipo', '')} {eq_temp.get('modelo_equipo', '')}"
        else:
            equipo_temp_text = equipo_temp_cod
        # Obtener equipo de velocidad del aire
        equipo_vel_cod = row_visita.get('equipo_vel_air', '')
        equipo_vel_detalles = df_equipos[df_equipos['id_equipo'] == equipo_vel_cod]
        if not equipo_vel_detalles.empty:
            eq_vel = equipo_vel_detalles.iloc[0]
            equipo_vel_text = f"{eq_vel.get('nombre_equipo', '')} {eq_vel.get('marca_equipo', '')} {eq_vel.get('modelo_equipo', '')}"
        else:
            equipo_vel_text = equipo_vel_cod
        doc.add_paragraph(
            f"Las condiciones del centro de trabajo se midieron utilizando los equipos {equipo_temp_text} y {equipo_vel_text}.")
        doc.add_paragraph()
    # --- FIN DE AGREGAR ---
    doc.add_paragraph(
        "La medición se realizó utilizando la metodología de FANGER para evaluación de confort térmico en espacios interiores de acuerdo a la Nota técnica N°47 del Instituto de Salud Pública.")
    doc.add_paragraph()
    doc.add_paragraph(
        "Se utiliza el estándar de vestimenta y tasa de actividad metabólica detalladas en la siguiente lista:")
    # Verificar que existan mediciones
    if not df_mediciones.empty:
        # Eliminar duplicados según las columnas deseadas
        df_lista = df_mediciones.drop_duplicates(subset=["puesto_trabajo", "clo", "met"])

        # Iterar sobre cada fila y crear un párrafo con la información formateada.
        for _, row_med in df_lista.iterrows():
            puesto = row_med.get("puesto_trabajo", "")
            clo_val = row_med.get("clo", "")
            met_val = row_med.get("met", "")

            # Formatear numéricamente y, por ejemplo, reemplazar el punto por coma en el valor de clo:
            try:
                clo_formateado = f"{float(clo_val):.2f}".replace(".", ",")
            except Exception:
                clo_formateado = str(clo_val)
            try:
                met_formateado = f"{float(met_val):.2f}"
            except Exception:
                met_formateado = str(met_val)

            # Generar el texto para el cargo
            linea = (f"- Para el cargo de {puesto} se utilizó un {clo_formateado} clo para vestimenta "
                     f"y se estimó en {met_formateado} met su actividad metabólica.")

            # Agregar la línea como un párrafo independiente
            doc.add_paragraph(linea)
    else:
        doc.add_paragraph("No se encontraron datos de vestimenta y actividad metabólica.")

    # -------------------------------
    # Resultados de mediciones y evaluación
    # -------------------------------
    doc.add_page_break()

    # Crear listas para las áreas que cumplen y no cumplen
    areas_cumplen = []
    areas_no_cumplen = []

    paragraph = doc.add_heading("Resultados de mediciones y evaluación", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not df_mediciones.empty:
        # Definir el orden y nombres de columnas a mostrar
        columnas_resumen = [
            "Área", "Sector", "Temp. bulbo seco(°C)", "Temp. globo(°C)",
            "Humedad relativa (%)", "Velocidad del aire(m/s)", "PPD", "PMV",
            "Estándar de confortabilidad térmica PMV  [-1,+1]"
        ]
        tabla_resumen = doc.add_table(rows=1, cols=len(columnas_resumen))
        tabla_resumen.style = 'Table Grid'
        # Formatear la primera fila para que tenga el estilo deseado (negrita y letras blancas)

        hdr_cells = tabla_resumen.rows[0].cells
        for idx, col_name in enumerate(columnas_resumen):
            hdr_cells[idx].text = col_name

        # Agrupar el DataFrame por área
        grouped = df_mediciones.groupby("nombre_area")
        for area, group in grouped:
            # Si hay más de una medición para el área, agregar línea de promedios
            if len(group) > 1:
                try:
                    avg_t_bul = group["t_bul_seco"].astype(float).mean()
                except Exception:
                    avg_t_bul = 0
                try:
                    avg_t_globo = group["t_globo"].astype(float).mean()
                except Exception:
                    avg_t_globo = 0
                try:
                    avg_hum = group["hum_rel"].astype(float).mean()
                except Exception:
                    avg_hum = 0
                try:
                    avg_vel = group["vel_air"].astype(float).mean()
                except Exception:
                    avg_vel = 0

                # Calcular promedios de met y clo; asignar valores por defecto si son 0
                try:
                    avg_met = group["met"].astype(float).mean()
                    if avg_met == 0:
                        avg_met = 1.1
                except Exception:
                    avg_met = 1.1
                try:
                    avg_clo = group["clo"].astype(float).mean()
                    if avg_clo == 0:
                        avg_clo = 0.5
                except Exception:
                    avg_clo = 0.5

                try:
                    results = pmv_ppd_iso(
                        tdb=avg_t_bul,
                        tr=avg_t_globo,
                        vr=avg_vel,
                        rh=avg_hum,
                        met=avg_met,
                        clo=avg_clo,
                        model="7730-2005",
                        limit_inputs=False,
                        round_output=True
                    )
                except Exception as e:
                    logging.error("Error al calcular pmv_ppd_iso para el área %s: %s", area, e)
                    results = None

                if results is not None:
                    if isinstance(results, dict):
                        avg_ppd = float(results.get("ppd", 0))
                        avg_pmv = float(results.get("pmv", 0))
                    elif hasattr(results, "ppd") and hasattr(results, "pmv"):
                        avg_ppd = float(results.ppd)
                        avg_pmv = float(results.pmv)
                    else:
                        avg_ppd = 0
                        avg_pmv = 0
                else:
                    avg_ppd = 0
                    avg_pmv = 0

                analisis = interpret_pmv(avg_pmv)

                # Agregar la línea de promedios
                row_cells = tabla_resumen.add_row().cells
                row_cells[0].text = str(area)  # Se muestra el área en la línea de promedios
                row_cells[1].text = "Promedio"  # Sector
                row_cells[2].text = f"{avg_t_bul:.2f}"
                row_cells[3].text = f"{avg_t_globo:.2f}"
                row_cells[4].text = f"{avg_hum:.1f}"
                row_cells[5].text = f"{avg_vel:.2f}"
                row_cells[6].text = f"{avg_ppd:.2f}"
                row_cells[7].text = f"{avg_pmv:.2f}"
                row_cells[8].text = f"{analisis}"  # Placeholder

                if analisis.upper() == "CUMPLE":
                    areas_cumplen.append(area)
                elif analisis.upper() == "NO CUMPLE":
                    areas_no_cumplen.append(area)

                # Si la celda de "Área" contiene información, poner toda la fila en negrita.
                if row_cells[0].text.strip():
                    set_row_bold(tabla_resumen.rows[-1])

            # Agregar las filas individuales del grupo:
            # Si el grupo tiene más de una fila (se agregó la línea de promedios), dejar la celda de área en blanco para todas las filas individuales.
            for _, row_med in group.iterrows():
                row_cells = tabla_resumen.add_row().cells
                if len(group) > 1:
                    row_cells[0].text = ""  # Dejar en blanco en filas individuales si hay resumen
                else:
                    row_cells[0].text = str(row_med.get("nombre_area", ""))
                if len(group) == 1:
                    row_cells[1].text = ""  # Dejar en blanco en filas individuales si hay resumen
                else:
                    row_cells[1].text = " - " + str(row_med.get("sector_especifico", ""))
                try:
                    t_bul = float(row_med.get("t_bul_seco", 0))
                    row_cells[2].text = f"{t_bul:.2f}"
                except:
                    row_cells[2].text = str(row_med.get("t_bul_seco", ""))
                try:
                    t_globo = float(row_med.get("t_globo", 0))
                    row_cells[3].text = f"{t_globo:.2f}"
                except:
                    row_cells[3].text = str(row_med.get("t_globo", ""))
                try:
                    hum = float(row_med.get("hum_rel", 0))
                    row_cells[4].text = f"{hum:.1f}"
                except:
                    row_cells[4].text = str(row_med.get("hum_rel", ""))
                try:
                    vel = float(row_med.get("vel_air", 0))
                    row_cells[5].text = f"{vel:.2f}"
                except:
                    row_cells[5].text = str(row_med.get("vel_air", ""))
                try:
                    ppd = float(row_med.get("ppd", 0))
                    row_cells[6].text = f"{ppd:.2f}"
                except:
                    row_cells[6].text = str(row_med.get("ppd", ""))
                try:
                    pmv = float(row_med.get("pmv", 0))
                    row_cells[7].text = f"{pmv:.2f}"
                except:
                    row_cells[7].text = str(row_med.get("pmv", ""))

                if len(group) > 1:
                    row_cells[8].text = ""  # Dejar en blanco en filas individuales si hay resumen
                else:
                    row_cells[8].text = str(row_med.get("resultado_medicion", ""))

                if row_cells[0].text.strip():
                    (set_row_bold
                     (tabla_resumen.rows[-1]))

                if len(group) == 1:
                    if row_med.get("resultado_medicion").upper() == "CUMPLE":
                        areas_cumplen.append(area)
                    elif row_med.get("resultado_medicion").upper() == "NO CUMPLE":
                        areas_no_cumplen.append(area)

        merge_column_cells(tabla_resumen, 0)
        merge_column_cells(tabla_resumen, 8)
        format_row(tabla_resumen.rows[0])
        set_column_width(tabla_resumen, 0, Cm(2.5))
        set_column_width(tabla_resumen, 1, Cm(2))
        set_column_width(tabla_resumen, 2, Cm(1.5))
        set_column_width(tabla_resumen, 3, Cm(1.5))
        set_column_width(tabla_resumen, 4, Cm(1.5))
        set_column_width(tabla_resumen, 5, Cm(1.5))
        set_column_width(tabla_resumen, 6, Cm(1.5))
        set_column_width(tabla_resumen, 7, Cm(1.5))
        set_column_width(tabla_resumen, 8, Cm(1.5))
    else:
        doc.add_paragraph("No se encontraron mediciones para detallar.")

    # Encabezado principal del contenido: Conclusiones
    doc.add_paragraph()
    paragraph = doc.add_heading("Conclusiones", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not df_centros.empty:
        row_centro = df_centros.iloc[0]
        nombre_ct = row_centro.get("nombre_ct", "")
        # Construir las cadenas de texto según las áreas que cumplen y las que no cumplen

        if areas_cumplen:
            if len(areas_cumplen) == 1:
                # Forma singular
                cumplen_text = f"el área de {areas_cumplen[0]}"
            else:
                # Forma plural, usando la función para unir con "y"
                cumplen_text = f"las áreas {join_with_and(areas_cumplen)}"
        else:
            cumplen_text = None  # O dejarlo en cadena vacía, según convenga

        if areas_no_cumplen:
            if len(areas_no_cumplen) == 1:
                no_cumplen_text = f"el área {areas_no_cumplen[0]}"
            else:
                no_cumplen_text = f"las áreas {join_with_and(areas_no_cumplen)}"
        else:
            no_cumplen_text = None

        # Generar la redacción final de las conclusiones
        if cumplen_text and no_cumplen_text:
            # Caso 2: Existen áreas que cumplen y áreas que no cumplen.
            doc.add_paragraph(
                f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, es posible concluir que {cumplen_text} "
                f"{'cumple' if len(areas_cumplen) == 1 else 'cumplen'} con el estándar de confort térmico, por lo que se recomienda mantener las condiciones actuales o similares."
            )
            doc.add_paragraph(
                f"Respecto a {no_cumplen_text} que NO {'cumple' if len(areas_no_cumplen) == 1 else 'cumplen'} con el estándar, se deben adoptar las medidas prescritas a continuación para corregir las condiciones."
            )
        elif cumplen_text and not no_cumplen_text:
            # Caso 1: Sólo existen áreas que cumplen.
            doc.add_paragraph(
                f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, se concluye que {cumplen_text} "
                f"{'cumple' if len(areas_cumplen) == 1 else 'cumplen'} con el estándar de confort térmico, por lo que se recomienda mantener las condiciones actuales o similares."
            )
        elif no_cumplen_text and not cumplen_text:
            # Caso 3: Sólo existen áreas que no cumplen.
            doc.add_paragraph(
                f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, se concluye que {no_cumplen_text} "
                f"{'NO cumple' if len(areas_no_cumplen) == 1 else 'NO cumplen'} con el estándar de confort térmico, por lo que se deben adoptar las medidas prescritas a continuación para corregir las condiciones."
            )
        else:
            # En caso de que no haya información suficiente
            doc.add_paragraph(
                "No se encontró información suficiente en las mediciones para emitir una conclusión sobre el confort térmico."
            )


    # -------------------------------
    # 4) MEDIDAS CORRECTIVAS
    # -------------------------------
    agregar_medidas_correctivas(doc, df_mediciones, areas_no_cumplen)
    doc.add_paragraph()

    # -------------------------------
    # 5) VIGENCIA DEL INFORME
    # -------------------------------
    # Encabezado principal del contenido: Vigencia del informe
    doc.add_paragraph()
    paragraph = doc.add_heading("Vigencia del informe", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not df_visitas.empty:
        doc.add_paragraph("En términos generales, el presente informe tiene validez de 3 años, a excepción que existan cambios en la situación, del tipo ingenieril o administrativo, que presupongan modificación a las condiciones encontradas al momento de la medición, lo cual implicará realizar una nueva evaluación en un plazo menor al señalado.")
        doc.add_paragraph()
        if len(areas_no_cumplen) > 0:
            doc.add_paragraph("Cuando se concreten los cambios indicados, la empresa deberá informar al IST el detalle de los mismos, de forma tal de programar las gestiones a realizar, las que considerarán previamente un informe de Verificación y Control y posteriormente, de acuerdo a sus resultados, la nueva evaluación de higiene ocupacional correspondiente.")
            doc.add_paragraph()
        doc.add_paragraph("Estos resultados de evaluación representan las condiciones existentes del ambiente y lugar de trabajo al momento de realizar las mediciones.")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    if not df_visitas.empty:
        row_visita = df_visitas.iloc[0]
        consultor_ist = row_visita.get("consultor_ist", "")

    # Agregar párrafo para el consultor, centrado y en negrita
    p_consultor = doc.add_paragraph()
    p_consultor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_consultor = p_consultor.add_run(consultor_ist)
    run_consultor.bold = True

    # Agregar párrafo para la profesión, centrado
    p_profesion = doc.add_paragraph("[__COMPLETAR_PROFESIÓN__]")
    p_profesion.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar párrafo para el zonal, centrado
    p_zonal = doc.add_paragraph("[__COMPLETAR_ZONAL__]")
    p_zonal.alignment = WD_ALIGN_PARAGRAPH.CENTER




    # -------------------------------
    # 6) ANEXOS
    # -------------------------------
    doc.add_page_break()
    doc.add_heading("Anexo 1. Características de los equipos de medición utilizado", level=2)
    # En lugar de mostrar el listado completo, se mostrará solo la información relacionada.
    if not df_visitas.empty and not df_equipos.empty:
        row_visita = df_visitas.iloc[0]
        # Obtener los códigos de equipos que están en uso en la visita
        equipo_temp_cod = row_visita.get('equipo_temp', '')
        equipo_vel_cod = row_visita.get('equipo_vel_air', '')
        codigos_en_uso = [equipo_temp_cod, equipo_vel_cod]

        # Filtrar df_equipos para que solo incluya las filas donde 'id_equipo' está en codigos_en_uso
        df_equipos_filtrado = df_equipos[df_equipos['id_equipo'].isin(codigos_en_uso)]

        # Definir el mapeo de campos a mostrar
        field_mapping = {
            "nombre_equipo": "Tipo de equipo",
            "cod_equipo": "Código",
            "n_serie_equipo": "Número de serie",
            "marca_equipo": "Marca",
            "modelo_equipo": "Modelo",
            "fecha_calibracion": "Ultima calibración",
            "prox_calibracion": "Próxima calibración",
            "empresa_certificadora": "Empresa certificadora",
            "num_certificado": "Número de certificado",
            "url_certificado": "Respaldo certificado"
        }

        if not df_equipos_filtrado.empty:
            for idx, row_eq in df_equipos_filtrado.iterrows():
                # Crear una tabla de dos columnas, una para el campo y otra para el valor
                tabla_equipo = doc.add_table(rows=len(field_mapping), cols=2)
                tabla_equipo.style = 'Table Grid'
                for row_num, (key, display_name) in enumerate(field_mapping.items()):
                    tabla_equipo.rows[row_num].cells[0].text = display_name
                    if key == "url_certificado":
                        url = str(row_eq.get(key, ""))
                        if url.strip():
                            qr_img = generate_qr_code(url)
                            cell = tabla_equipo.rows[row_num].cells[1]
                            cell.text = ""
                            run = cell.paragraphs[0].add_run()
                            run.add_break()
                            run.add_picture(qr_img, width=Inches(1))
                            run.add_break()
                        else:
                            tabla_equipo.rows[row_num].cells[1].text = ""
                    else:
                        tabla_equipo.rows[row_num].cells[1].text = str(row_eq.get(key, ""))
                        set_column_width(tabla_equipo, 0, Cm(4))
                        set_column_width(tabla_equipo, 1, Cm(14))

                doc.add_paragraph("")  # Separador entre tablas
        else:
            doc.add_paragraph("No se encontró información de equipos de medición relacionados con la visita.")

    doc.add_paragraph()
    table_calib = doc.add_table(rows=0, cols=2)
    table_calib.style = 'Table Grid'
    if not df_visitas.empty:
        add_row(table_calib, "Verificación de parámetros de equipos en terreno", "")
        add_row(table_calib, "Patrón TBS", row_visita.get('patron_tbs', ''))
        add_row(table_calib, "Verificación TBS inicial", row_visita.get('ver_tbs_ini', ''))
        add_row(table_calib, "Verificación TBS final", row_visita.get('ver_tbs_fin', ''))
        add_row(table_calib, "Patrón TBH", row_visita.get('patron_tbh', ''))
        add_row(table_calib, "Verificación TBH inicial", row_visita.get('ver_tbh_ini', ''))
        add_row(table_calib, "Verificación TBH final", row_visita.get('ver_tbh_fin', ''))
        add_row(table_calib, "Patrón TG", row_visita.get('patron_tg', ''))
        add_row(table_calib, "Verificación TG inicial", row_visita.get('ver_tg_ini', ''))
        add_row(table_calib, "Verificación TG final", row_visita.get('ver_tg_fin', ''))
    else:
        add_row(table_calib, "D. Detalles de equipos y calibración", "No se encontró información de visita.")
    set_column_width(table_calib, 0, Cm(4))
    set_column_width(table_calib, 1, Cm(14))

    ##########################
    #Observaciones de áreas
    ##########################

    doc.add_page_break()
    doc.add_heading("Anexo 2. Características de las areas evaluadas", level=2)

    # Crear la tabla con 3 columnas: Área, Características constructivas y Condiciones de ventilación
    tabla_caract = doc.add_table(rows=1, cols=3)
    tabla_caract.style = 'Table Grid'

    # Agregar la fila de encabezado
    hdr_cells = tabla_caract.rows[0].cells
    hdr_cells[0].text = "Área"
    hdr_cells[1].text = "Características constructivas"
    hdr_cells[2].text = "Condiciones de ventilación"

    # Agrupar el DataFrame por "nombre_area"
    grouped = df_mediciones.groupby("nombre_area")
    for area, group in grouped:
        # Usar el primer registro del grupo para extraer los datos de instalación
        registro = group.iloc[0]

        # Agregar una fila para el área con sus respectivos datos
        row_cells = tabla_caract.add_row().cells
        row_cells[0].text = area
        row_cells[1].text = str(registro["caract_constructivas"])
        row_cells[2].text = str(registro["ingreso_salida_aire"])

    set_column_width(tabla_caract, 0, Cm(2.5))
    set_column_width(tabla_caract, 1, Cm(7))
    set_column_width(tabla_caract, 2, Cm(7))
    format_row(tabla_caract.rows[0])



    '''
    # Agrupar el DataFrame por "nombre_area"
    grouped = df_mediciones.groupby("nombre_area")
    for area, group in grouped:
        # Agregar un título para el área
        doc.add_paragraph(f"Área: {area}", style='Heading2')

        # Usar el primer registro del grupo para extraer los datos de instalación
        registro = group.iloc[0]

        # Definir los pares (etiqueta, valor) a mostrar
        campos = [
            ("cond_techumbre", registro["cond_techumbre"]),
            ("obs_techumbre", registro["obs_techumbre"]),
            ("cond_paredes", registro["cond_paredes"]),
            ("obs_paredes", registro["obs_paredes"]),
            ("cond_vantanal", registro["cond_vantanal"]),
            ("obs_ventanal", registro["obs_ventanal"]),
            ("cond_aire_acond", registro["cond_aire_acond"]),
            ("obs_aire_acond", registro["obs_aire_acond"]),
            ("cond_ventiladores", registro["cond_ventiladores"]),
            ("obs_ventiladores", registro["obs_ventiladores"]),
            ("cond_inyeccion_extraccion", registro["cond_inyeccion_extraccion"]),
            ("obs_inyeccion_extraccion", registro["obs_inyeccion_extraccion"]),
            ("cond_ventanas", registro["cond_ventanas"]),
            ("obs_ventanas", registro["obs_ventanas"]),
            ("cond_puertas", registro["cond_puertas"]),
            ("obs_puertas", registro["obs_puertas"]),
            ("cond_otras", registro["cond_otras"]),
            ("obs_otras", registro["obs_otras"])
        ]

        # Crear la tabla con 2 columnas; la cantidad de filas es 1 (opcional, para encabezado) + len(campos)
        tabla_instalacion = doc.add_table(rows=1, cols=2)
        tabla_instalacion.style = 'Table Grid'

        # Opcional: agregar una fila de encabezado para identificar las columnas
        hdr_cells = tabla_instalacion.rows[0].cells
        hdr_cells[0].text = "Etiqueta"
        hdr_cells[1].text = "Contenido"

        # Agregar una fila por cada campo con su etiqueta y valor
        for etiqueta, valor in campos:
            row_cells = tabla_instalacion.add_row().cells
            row_cells[0].text = etiqueta
            row_cells[1].text = str(valor)

    '''

    # -------------------------------
    # Finaliza el documento y lo retorna como BytesIO
    # -------------------------------
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# -----------------------------------------------
# FUNCIÓN PRINCIPAL (OPCIONAL) PARA GENERAR EL INFORME
# -----------------------------------------------
def generar_informe(cuv: str):
    """
    Función principal para generar el informe a partir del CUV.
    Se obtienen los dataframes necesarios (usando las funciones de data_access)
    y se genera el documento Word.
    """
    from data_access import get_centro, get_visita, get_mediciones, get_equipos
    df_centros = get_centro(cuv)
    df_visitas = get_visita(cuv)
    if df_visitas.empty:
        logging.error("No se encontró ninguna visita para el CUV proporcionado.")
        return
    # Se selecciona la visita más reciente
    visita_id = df_visitas.iloc[0].get("id_visita")
    df_mediciones = get_mediciones(visita_id)
    df_equipos = get_equipos()
    buffer = generar_informe_en_word(df_centros, df_visitas, df_mediciones, df_equipos)
    output_filename = f"informe_confort_termico_{cuv}.docx"
    with open(output_filename, "wb") as f:
        f.write(buffer.getbuffer())
    logging.info(f"Informe generado y guardado como {output_filename}")


if __name__ == "__main__":
    # Ejemplo de llamada. Modifica el CUV según sea necesario.
    cuv_input = "114123"  # Cambia el valor según el CUV a consultar
    generar_informe(cuv_input)
