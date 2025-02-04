import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime
import requests
import re

def descargar_imagen_gdrive(url_foto: str) -> BytesIO or None:
    """
    Dada una URL de Google Drive (tipo 'https://drive.google.com/open?id=ABC123'),
    intenta descargar el archivo y devolverlo como BytesIO.
    Retorna None si no se puede descargar o no tenga permisos adecuados.
    """
    match = re.search(r"id=(.*)", url_foto)
    if not match:
        return None

    file_id = match.group(1)
    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"

    try:
        response = requests.get(download_url)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            return None
    except:
        return None

def generar_informe_en_word(df_filtrado: pd.DataFrame,
                            df_info_cuv: pd.DataFrame) -> BytesIO:
    """
    Genera un informe en Word (DOCX) combinando:
      - df_info_cuv: Info del CUV (RUT, RAZÓN SOCIAL, etc.)
      - df_filtrado: Datos Generales / Mediciones del CSV principal
    """

    doc = Document()

    # Portada
    doc.add_heading("EVALUACIÓN CONFORT TÉRMICO", 0)
    doc.add_paragraph(
        f"Fecha de generación del informe: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
    )

    # --- 1) Información de CUV (df_info_cuv) ---
    doc.add_heading("Identificación de empresa y centro de trabajo", level=1)
    if not df_info_cuv.empty:
        row_cuv_info = df_info_cuv.iloc[0]
        doc.add_heading("A. Información empresa", level=2)
        doc.add_paragraph(f"Razón Social: {row_cuv_info.get('RAZÓN SOCIAL','')}")
        doc.add_paragraph(f"RUT: {row_cuv_info.get('RUT','')}")
        doc.add_heading("B. Información centro de trabajo", level=2)
        doc.add_paragraph(f"CUV: {row_cuv_info.get('CUV','')}")
        doc.add_paragraph(f"Nombre de Local: {row_cuv_info.get('Nombre de Local','')}")
        doc.add_paragraph(f"Dirección: {row_cuv_info.get('Dirección','')}")
        doc.add_paragraph(f"Comuna: {row_cuv_info.get('Comuna','')}")
        doc.add_paragraph(f"Región: {row_cuv_info.get('Región','')}")
    else:
        doc.add_paragraph("No se encontró información adicional (RUT, Razón Social, etc.) para este CUV.")

    # --- 2) Datos Generales y Mediciones ---
    datos_generales = df_filtrado[df_filtrado["Que seccion quieres completar"] == "Datos generales"]
    mediciones = df_filtrado[df_filtrado["Que seccion quieres completar"] == "Medición de un area"]

    # 2.1) Datos Generales
    if not datos_generales.empty:
        for idx, row in datos_generales.iterrows():
            doc.add_heading("C. Antecedentes generales de la evaluación", level=2)
            doc.add_paragraph(f"Fecha visita: {row.get('Fecha visita', '')}")
            doc.add_paragraph(f"Hora medición: {row.get('Hora medicion', '')}")
            doc.add_paragraph(f"Temperatura máxima del día: {row.get('Temperatura máxima del día', '')} °C")
            doc.add_paragraph(f"Nombre del personal SMU: {row.get('Nombre del personal SMU', '')}")
            doc.add_paragraph(f"Cargo: {row.get('Cargo', '')}")

            # Ejemplo de correo
            doc.add_paragraph(f"Profesional IST: {row.get('Dirección de correo electrónico','')}")
            doc.add_paragraph("... A partir del registro de correo se buscan datos del profesional IST...")
    else:
        doc.add_paragraph("No se han encontrado filas de 'Datos Generales' para este CUV.")

    # 2.2) Mediciones de un área
    if not mediciones.empty:
        for idx, row in mediciones.iterrows():
            doc.add_page_break()
            doc.add_heading(f"Medición área: {row.get('Area o sector','')}", level=2)
            doc.add_paragraph(f"Especificación sector: {row.get('Especificación sector','')}")
            doc.add_paragraph(
                f"Temperatura bulbo seco (°C): {row.get('Temperatura bulbo seco (°C) ejemplo 25.3','')}"
            )
            doc.add_paragraph(
                f"Temperatura globo (°C): {row.get('Temperatura globo (°C)  ejemplo 24.8','')}"
            )
            doc.add_paragraph(
                f"Humedad relativa (%): {row.get('Humedad relativa (%)  ejemplo 18','')}"
            )
            doc.add_paragraph(
                f"Velocidad del aire (m/s): {row.get('Velocidad del aire (m/s)  ejemplo 0.3','')}"
            )
            # Agrega más columnas si deseas

            # Fotografía(s)
            evidencia = row.get("Evidencia fotografica", "")
            if pd.notnull(evidencia) and isinstance(evidencia, str) and evidencia.strip():
                doc.add_paragraph("Fotografías asociadas:")
                for foto_url in [f.strip() for f in evidencia.split(",") if f.strip()]:
                    imagen = descargar_imagen_gdrive(foto_url)
                    if imagen is not None:
                        doc.add_picture(imagen, width=Inches(4))
                    else:
                        doc.add_paragraph(f"No se pudo descargar la imagen: {foto_url}")
    else:
        doc.add_paragraph("No se han encontrado filas de 'Medición de un area' para este CUV.")

    # --- Pie de informe ---
    doc.add_paragraph("Informe versión prueba")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
