#!/usr/bin/env python3
import logging
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from pythermalcomfort.models import pmv_ppd_iso
import qrcode
from PIL import ImageOps  # Asegúrate de tener Pillow instalado
from io import BytesIO
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from pdf2image import convert_from_path
import requests
import os
from natsort import natsorted
from collections import OrderedDict


# Configuración básica del logging
# logging.basicConfig(level=logging.INFO)
logging.basicConfig(level=logging.DEBUG)


# -----------------------------------------------
# Nuevas funciones
# -----------------------------------------------


def calcular_analisis_area(group):
    """Determina si un área cumple o no basado en mediciones."""
    if len(group) > 1:
        # Cálculo de promedios
        avg_t_bul = group["t_bul_seco"].astype(float).mean()
        avg_t_globo = group["t_globo"].astype(float).mean()
        avg_hum = group["hum_rel"].astype(float).mean()
        avg_vel = group["vel_air"].astype(float).mean()
        avg_met = group["met"].astype(float).mean() or 1.1
        avg_clo = group["clo"].astype(float).mean() or 0.5

        # Cálculo PMV/PPD
        try:
            results = pmv_ppd_iso(
                tdb=avg_t_bul,
                tr=avg_t_globo,
                vr=avg_vel,
                rh=avg_hum,
                met=avg_met,
                clo=avg_clo,
                model="7730-2005",
                limit_inputs=False
            )
            pmv = results.pmv if hasattr(results, 'pmv') else results.get('pmv', 0)
            analisis = interpret_pmv(pmv)
        except Exception:
            analisis = "NO CUMPLE"
    else:
        analisis = group.iloc[0].get("resultado_medicion", "NO CUMPLE").upper()

    return analisis


def procesar_areas(df_mediciones):
    """Procesa todas las áreas para determinar cumplimiento."""
    areas_cumplen = []
    areas_no_cumplen = []

    if not df_mediciones.empty:
        grouped = df_mediciones.groupby("nombre_area")
        for area, group in grouped:
            analisis = calcular_analisis_area(group)
            if analisis == "CUMPLE":
                areas_cumplen.append(area)
            else:
                areas_no_cumplen.append(area)
    return areas_cumplen, areas_no_cumplen


# -----------------------------------------------
# FUNCIONES AUXILIARES PARA EL DOCUMENTO WORD
# -----------------------------------------------


def format_columns(df, columns, mode="title"):
    """
    Formatea las columnas especificadas en un DataFrame eliminando espacios al inicio y al final,
    y transformando el texto según el modo seleccionado.

    Parámetros:
        df (DataFrame): El DataFrame a modificar.
        columns (list): Lista de nombres de columnas a formatear.
        mode (str): Modo de transformación, puede ser:
            - "title": Convierte el texto a Título (cada palabra con su primera letra en mayúscula).
            - "capitalize": Solo la primera letra de la primera palabra en mayúscula.
            - "upper": Todo el texto en mayúsculas.
            (El valor por defecto es "title".)

    Retorna:
        DataFrame: El DataFrame con las columnas formateadas.
    """
    for col in columns:
        if col in df.columns:
            # Convertir a cadena y eliminar espacios al inicio y final
            df[col] = df[col].astype(str).str.strip()
            # Aplicar la transformación según el modo
            if mode == "title":
                df[col] = df[col].str.lower().str.title()
            elif mode == "capitalize":
                df[col] = df[col].str.lower().str.capitalize()
            elif mode == "upper":
                df[col] = df[col].str.upper()
            else:
                raise ValueError(f"Modo '{mode}' no reconocido. Use 'title', 'capitalize' o 'upper'.")
    return df


def set_column_width(table, col_index, width):
    """
    Establece el ancho de la columna 'col_index' de la tabla 'table' al valor 'width' (un objeto de docx.shared, por ejemplo Cm(4)).
    """
    # width._value retorna el ancho en unidades dxa (1 dxa = 1/20 de punto)
    for cell in table.columns[col_index].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width.inches * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)


def merge_column_cells(tabla, col_index, start_row=1):
    """
    Fusiona celdas consecutivas en la columna 'col_index' de la tabla 'tabla'
    a partir de la fila 'start_row' (por defecto se salta el encabezado).
    Se fusionan las celdas cuando se encuentran celdas vacías, es decir, se asume
    que solo la primera celda del grupo contiene el valor.
    """
    num_filas = len(tabla.rows)
    current_value = None
    start_index = None

    for i in range(start_row, num_filas):
        cell = tabla.cell(i, col_index)
        text = cell.text.strip()
        # Si la celda tiene contenido, se considera el inicio de un grupo
        if text:
            # Si ya había un grupo iniciado y abarca más de una fila, se fusionan las celdas anteriores
            if current_value is not None and start_index is not None and i - start_index > 1:
                primera_celda = tabla.cell(start_index, col_index)
                for j in range(start_index + 1, i):
                    primera_celda = primera_celda.merge(tabla.cell(j, col_index))
            current_value = text
            start_index = i
    # Fusionar el último grupo pendiente (si abarca más de una fila)
    if current_value is not None and start_index is not None and num_filas - start_index > 1:
        primera_celda = tabla.cell(start_index, col_index)
        for j in range(start_index + 1, num_filas):
            primera_celda = primera_celda.merge(tabla.cell(j, col_index))


def generate_qr_code(url, border=10, box_size=2):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=box_size,
        border=2  # Borde interno del QR
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
    # Agregar un borde (padding) blanco alrededor de la imagen
    padded_img = ImageOps.expand(img, border=border, fill='white')
    bio = BytesIO()
    padded_img.save(bio, format="PNG")
    bio.seek(0)
    return bio


def join_with_and(items):
    """
    Une los elementos de la lista usando la conjunción "y" para el último elemento.
    Ejemplos:
      - ['Oficina']             => "Oficina"
      - ['Oficina', 'Sala']     => "Oficina y Sala"
      - ['Oficina', 'Sala', 'Bodega'] => "Oficina, Sala y Bodega"
    """
    items = list(items)  # Convertir a lista para evitar ambigüedades
    if not items:
        return ""
    elif len(items) == 1:
        return items[0]
    elif len(items) == 2:
        return f"{items[0]} y {items[1]}"
    else:
        return ", ".join(items[:-1]) + " y " + items[-1]


def get_or_add_lang(rPr):
    """
    Obtiene o agrega el elemento de idioma en la configuración de la fuente.
    """
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    return lang


def interpret_pmv(pmv_value):
    """Ejemplo de interpretación simple:
       - CUMPLE si -1 <= pmv <= 1
       - NO CUMPLE en caso contrario.
       Ajusta según tu criterio."""
    if pmv_value <= -1 or pmv_value >= 1:
        return "NO CUMPLE"
    else:
        return "CUMPLE"


'''
def interpret_pmv(pmv_value):
    if pmv_value >= 1:
        return "NO CUMPLE"
    elif pmv_value > -1:
        return "CUMPLE"
    else:
        return "NO CUMPLE"
'''

def descargar_imagen_gdrive(url_foto: str) -> BytesIO or None:
    """
    Dada una URL de Google Drive, intenta descargar el archivo y devolverlo como BytesIO.
    Retorna None si no se puede descargar o no tiene permisos adecuados.
    """
    match = re.search(r"id=(.*)", url_foto)
    if not match:
        return None
    file_id = match.group(1)
    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
    try:
        response = requests.get(download_url)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            return None
    except Exception as e:
        return None


# -----------------------------------------------------
# NUEVAS FUNCIONES PARA CONFIGURAR ESTILOS DE MANERA MODULAR
# -----------------------------------------------------

def set_style_language(style, lang_code="es-CL"):
    """
    Asigna el código de idioma a un estilo dado.
    Se configura para el texto normal, Asia Oriental y bidireccional.
    """
    rPr = style.element.get_or_add_rPr()
    lang = get_or_add_lang(rPr)
    lang.set(qn('w:val'), lang_code)
    lang.set(qn('w:eastAsia'), lang_code)
    lang.set(qn('w:bidi'), lang_code)


def apply_style_properties(style, properties):
    """
    Aplica las propiedades definidas en el diccionario 'properties' al estilo 'style'.
    Las propiedades pueden incluir: font_name, font_size, font_color, bold, italic,
    space_before, space_after, alignment y lang_code.
    """
    font = style.font
    font.name = properties.get("font_name", font.name)
    font.size = properties.get("font_size", font.size)
    font.color.rgb = properties.get("font_color", font.color.rgb)
    font.bold = properties.get("bold", font.bold)
    font.italic = properties.get("italic", font.italic)

    para_format = style.paragraph_format
    para_format.space_before = properties.get("space_before", para_format.space_before)
    para_format.space_after = properties.get("space_after", para_format.space_after)
    para_format.alignment = properties.get("alignment", para_format.alignment)

    # Aplicar idioma
    set_style_language(style, properties.get("lang_code", "es-CL"))


# Diccionario de configuración de estilos
from docx.enum.text import WD_ALIGN_PARAGRAPH

style_configurations = {
    "Normal": {
        "font_name": "Calibri",
        "font_size": Pt(10),
        "font_color": RGBColor(0x00, 0x00, 0x00),
        "bold": False,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "lang_code": "es-CL"
    },
    "Heading 1": {
        "font_name": "Calibri",
        "font_size": Pt(14),
        "font_color": RGBColor(0, 0, 0),
        "bold": True,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(12),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "lang_code": "es-CL"
    },
    "Heading 2": {
        "font_name": "Calibri",
        "font_size": Pt(12),
        "font_color": RGBColor(79, 11, 123),
        "bold": True,
        "italic": False,
        "space_before": Pt(12),
        "space_after": Pt(12),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "lang_code": "es-CL"
    },
    "Heading 3": {
        "font_name": "Calibri",
        "font_size": Pt(11),
        "font_color": RGBColor(0x30, 0x30, 0x30),
        "bold": True,
        "italic": False,
        "space_before": Pt(8),
        "space_after": Pt(8),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "lang_code": "es-CL"
    },
    "TablaTexto": {  # Estilo personalizado para el contenido de las tablas
        "font_name": "Calibri",
        "font_size": Pt(10),
        "font_color": RGBColor(0, 0, 0),
        "bold": False,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "lang_code": "es-CL"
    },
    "Centrado": {
        "font_name": "Calibri",
        "font_size": Pt(10),
        "font_color": RGBColor(0, 0, 0),
        "bold": False,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "lang_code": "es-CL"
    },
    "Centrado Bold": {
        "font_name": "Calibri",
        "font_size": Pt(10),
        "font_color": RGBColor(0, 0, 0),
        "bold": True,
        "italic": False,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "lang_code": "es-CL"
    }
}


def look_informe(doc: Document):
    """
    Configura el documento: establece orientación, márgenes, estilos y el idioma.
    """
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Aplicar configuración a los estilos definidos en nuestro diccionario
    for style_name, properties in style_configurations.items():
        try:
            style = doc.styles[style_name]
            apply_style_properties(style, properties)
        except KeyError:
            # Si el estilo no existe, se continúa
            continue



def set_vertical_alignment(doc: Document, section_index=0, alignment='top'):
    """
    Ajusta la alineación vertical de la sección especificada.
    """
    section = doc.sections[section_index]
    sectPr = section._sectPr
    vAlign = sectPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        sectPr.append(vAlign)
    alignment = alignment.lower()
    if alignment not in ['top', 'center', 'both']:
        alignment = 'top'
    vAlign.set(qn('w:val'), alignment)


# Función auxiliar para poner en negrita todas las celdas de una fila
def set_row_bold(row):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def apply_style(cell, shading_color="4f0b7b"):
    """
    Aplica un formato de fondo (con el color indicado), negrita y tamaño de fuente a una celda.
    """
    # Crear el elemento de sombreado con el color especificado
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shading_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    # Para cada párrafo y cada "run" dentro de la celda, aplicar color de fuente blanco, negrita y tamaño 12
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)


def format_row(row, shading_color="4f0b7b"):
    """
    Aplica el formato de estilo (fondo, negrita, color de fuente y tamaño) a todas las celdas de una fila.
    """
    for cell in row.cells:
        apply_style(cell, shading_color)


def add_row(table, label, value="", first=False):
    """
    Agrega una fila a la tabla.

    Si 'value' es una cadena vacía, se interpreta como una fila cabecera;
    se realiza un merge de la primera y segunda celda, se asigna el texto 'label'
    y se aplica el estilo (usando apply_style).

    Si 'value' no es vacío, se asigna 'label' a la primera celda y 'value' a la segunda.
    Además, si se especifica first=True, se aplica el mismo estilo a todas las celdas de la fila.
    """
    row = table.add_row()
    cells = row.cells
    if value == "":
        # Fila cabecera: se fusionan las dos primeras celdas y se formatea el resultado.
        merged_cell = cells[0].merge(cells[1])
        merged_cell.text = label
        apply_style(merged_cell)
    else:
        # Fila normal: se asigna el texto a la primera y segunda celda.
        cells[0].text = label
        cells[1].text = "" if (value is None or (isinstance(value, str) and value.strip() == "")) else str(value)
        if first:
            format_row(row)


def agregar_contenido(cell, items):
    """
    Agrega cada elemento de 'items' a la celda como párrafo y, posteriormente,
    elimina el primer párrafo si está vacío y limpia la celda para eliminar
    espacios y saltos de línea extra, de modo que solo se conserven los párrafos
    con contenido.
    """
    # 1) Agregar cada elemento como párrafo
    for item in items:
        cell.add_paragraph(item)

    if cell.paragraphs:
        parrafo_a_eliminar = cell.paragraphs[0]
        # Accede al elemento XML y elimínalo del padre
        p_element = parrafo_a_eliminar._element
        p_element.getparent().remove(p_element)


def agregar_medidas_por_accion(table, medidas):
    """
    Agrega filas a la tabla para cada medida.
    Por cada acción de cada medida, se crea una fila nueva,
    manteniendo el contenido de 'areas' y 'plazo' en cada fila.
    """
    for medida in medidas:
        for accion in medida['acciones']:
            row_cells = table.add_row().cells
            # Agrega el contenido de 'areas' en la primera celda.
            agregar_contenido(row_cells[0], medida['areas'])
            # Agrega la acción actual en la segunda celda (se pasa en forma de lista).
            agregar_contenido(row_cells[1], [accion])
            # Asigna el plazo en la tercera celda.
            row_cells[2].text = medida['plazo']
            # Formatea verticalmente centrado cada celda.
            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table





####
# Añade al final de doc_utils.py

def generar_recomendaciones(pmv, tdb_initial, tr_initial, vr, rh, met, clo):
    recomendaciones = []

    # Calcular diferencias térmicas importantes
    dif_temp = tr_initial - tdb_initial

    # 1. Estrategias principales según condición térmica
    estrategias_base = []

    if pmv > 1.0:  # Ambiente caluroso
        estrategias_base = [
            {
                'condicion': True,  # Siempre aplica para calor
                'tipo': 'enfriamiento',
                'mensaje': "Refrigeración activa requerida",
                'acciones': [
                    "- Implementar [__] equipos enfriadores de aire (según las dimensiones de las áreas), con el fin de enfriar el aire.",
                    "- Consultar con el proveedor el óptimo uso del equipo por ejemplo: periodicidad de suministrar agua helada, hielo o implemento refrigerante autorizado para el equipamiento adquirido, con el fin de estar constantemente enfriando durante toda la jornada laboral el área, especialmente en periodo de mayor temperaturas o época estival.",
                    "- Llevar una Bitácora o Registro de la actividad en lo referido al uso de enfriador(es)."
                ],
                'plazo': '6 meses desde la recepción del presente informe técnico'
            },
            {
                'condicion': dif_temp > 2.0,
                'tipo': 'aislamiento',
                'mensaje': "Reducción de carga térmica",
                'acciones': [
                    "- (REVISAR SI CORRESPONDE) Instalar materiales aislantes en techos/paredes",
                    "- (REVISAR SI CORRESPONDE) Implementar protecciones solares reflectivas",
                    "- (REVISAR SI CORRESPONDE) Aislar fuentes de calor radiante"
                ],
                'plazo': '3 meses desde la recepción del presente informe técnico'
            },
            {
                'condicion': vr < 0.2,
                'tipo': 'ventilacion',
                'mensaje': f"Aumentar ventilación",
                'acciones': [
                    "- Implementar [__] ventiladores industriales, con el fin de generar corrientes de aire, las cuales ayudarán a mejorar condiciones de confort térmico en dicha área.",
                    "- (REVISAR SI CORRESPONDE) Implementar sistemas de extracción forzada",
                    "- (REVISAR SI CORRESPONDE) Optimizar ventilación cruzada"
                ],
                'plazo': '3 meses desde la recepción del presente informe técnico'
            }
        ]
    '''
    #NOTA:
    #Se desactivan las recomendaciones por frio, dado que el foco de las evaluaciones son por calor en esta temporada.
    #NOTA
    
    
    elif pmv < -1.0:  # Ambiente frío
        estrategias_base = [
            {
                'condicion': True,
                'tipo': 'calefaccion',
                'mensaje': "Protección contra el frío",
                'acciones': [
                    "- (REVISAR SI CORRESPONDE) Implementar sistemas de calefacción radiante",
                    "- (REVISAR SI CORRESPONDE) Mejorar aislamiento térmico en envolvente",
                    "- (REVISAR SI CORRESPONDE) Optimizar sellado de infiltraciones"
                ],
                'plazo': '3 meses desde la recepción del presente informe técnico'
            }
        ]
    '''

    # Filtrar y agregar estrategias válidas
    for estrategia in estrategias_base:
        if estrategia['condicion']:
            recomendaciones.append({
                'tipo': estrategia['tipo'],
                'categoria': estrategia['mensaje'],
                'nivel': 'Prioridad 1',
                'mensaje': estrategia['mensaje'],
                'acciones': estrategia.get('acciones', []),
                'plazo': estrategia.get('plazo', 'Inmediato')
            })

    '''
    # 2. Mantenimiento (siempre aplica)
    mantenimiento = {
        'preventivo': [
            "Programar mantención de equipos con proveedores",
            "Calendario de limpieza de filtros/paneles",
            "Verificación mensual de sistemas"
        ],
        'correctivo': [
            "Reparación de sistemas de ventilación",
            "Ajuste de equipos de climatización",
            "Registro de intervenciones técnicas"
        ],
        'control': [
            f"Regulación térmica (23-26°C) con registro",
            "Monitoreo continuo de parámetros ambientales"
        ]
    }
    recomendaciones.append({
        'tipo': 'mantenimiento',
        'categoria': 'Gestión Técnica',
        'nivel': 'Prioridad 2',
        'mensaje': "Programa de mantenimiento integral",
        'acciones': mantenimiento,
        'plazo': 'Continuo'
    })

    # 3. Medidas administrativas (siempre aplican)
    recomendaciones.append({
        'tipo': 'administrativa',
        'categoria': 'Comunicación',
        'nivel': 'Prioridad 3',
        'acciones': [
            "Informar formalmente a todo el personal sobre riesgos térmicos (DS N°44)",
            "Establecer registros firmados de capacitación"
        ],
        'plazo': '30 días desde la recepción del presente informe técnico'
    })
    '''

    return recomendaciones

def crear_tabla_recomendaciones(doc, tipo_medida, medidas):
    """Crea tabla de recomendaciones por tipo de medida"""
    # Configurar encabezados según el tipo
    headers = ["Área", "Prescripción de medidas", "Plazo"]
    col_widths = [Cm(2), Cm(14), Cm(4)]

    # Crear tabla
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Configurar anchos de columna
    for idx, width in enumerate(col_widths):
        set_column_width(table, idx, width)

    # Formatear cabecera
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    format_row(table.rows[0], shading_color="4F0B7B")  # Morado corporativo

    table = agregar_medidas_por_accion(table, medidas)

    '''
    # Llenar con medidas
    for medida in medidas:
        row_cells = table.add_row().cells
        agregar_contenido(row_cells[0],medida['areas'])
        agregar_contenido(row_cells[1], medida['acciones'])
        row_cells[2].text = medida['plazo']
        # Formato vertical centrado
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table
    '''

def agregar_medidas_correctivas(doc, df_mediciones, areas_no_cumplen):
    # 1. Medidas Ingenieriles (solo para áreas no conformes)
    medidas_ingenieriles = []
    if areas_no_cumplen:
        for area in areas_no_cumplen:
            grupo = df_mediciones[df_mediciones['nombre_area'] == area]
            avg_params = {
                'tdb': grupo['t_bul_seco'].mean(),
                'tr': grupo['t_globo'].mean(),
                'vr': grupo['vel_air'].mean(),
                'rh': grupo['hum_rel'].mean(),
                'met': grupo['met'].mean(),
                'clo': grupo['clo'].mean(),
            }

            try:
                results = pmv_ppd_iso(
                    tdb=avg_params['tdb'],
                    tr=avg_params['tr'],
                    vr=avg_params['vr'],
                    rh=avg_params['rh'],
                    met=avg_params['met'],
                    clo=avg_params['clo'],
                    model="7730-2005",
                    limit_inputs=False,
                    round_output=True
                )
            except Exception as e:
                logging.error("Error al calcular pmv_ppd_iso para el área %s: %s", area, e)
                results = None

            if results is not None:
                if isinstance(results, dict):
                    avg_ppd = float(results.get("ppd", 0))
                    avg_pmv = float(results.get("pmv", 0))
                elif hasattr(results, "ppd") and hasattr(results, "pmv"):
                    avg_ppd = float(results.ppd)
                    avg_pmv = float(results.pmv)
                else:
                    avg_ppd = 0
                    avg_pmv = 0
            else:
                avg_ppd = 0
                avg_pmv = 0

            avg_params.update({
                'pmv': avg_pmv,
                'ppd':avg_ppd
            })

            recs = generar_recomendaciones(
                pmv=avg_params['pmv'],
                tdb_initial=avg_params['tdb'],
                tr_initial=avg_params['tr'],
                vr=avg_params['vr'],
                rh=avg_params['rh'],
                met=avg_params['met'],
                clo=avg_params['clo']
            )

            for rec in recs:
                if rec['tipo'] in ['ventilacion', 'enfriamiento', 'aislamiento', 'calefaccion']:
                    medidas_ingenieriles.append({
                        #'tipo_medida': rec['categoria'],
                        'areas': [area],
                        'acciones': rec['acciones'],
                        'plazo': rec['plazo']
                    })

    # 2. Medidas Administrativas (siempre se incluyen)
    medidas_administrativas = [
        {
            #'tipo_medida': 'Comunicación',
            'areas': ['Todas'],
            'acciones': [
                "- Informar a cada persona trabajadora acerca de los riesgos que entrañan sus labores, de las medidas preventivas, de los métodos y/o procedimientos de trabajo correctos, acorde a lo identificado por la empresa. Además de lo señalado previamente, la entidad empleadora deberá informar de manera oportuna y adecuada el resultado del presente informe técnico.",
                "- Realizar capacitaciones (teóricas/prácticas) periódicas en prevención de riesgos laborales, con la finalidad de garantizar el aprendizaje efectivo y eficaz, dejando registro de dichas capacitaciones y evaluaciones. En el marco de los artículos 15° y 16° del Párrafo IV del D.S 44 “Aprueba nuevo reglamento sobre gestión preventiva de los riesgos laborales para un entorno de trabajo seguro y saludable."
            ],
            'plazo': '30 días desde la recepción del presente informe técnico'
        },
        {
            #'tipo_medida': 'Mantenimiento',
            'areas': areas_no_cumplen if areas_no_cumplen else ['Todas'],
            'acciones': [
                "- Realizar mantención preventiva en los equipos de climatización, con el fin de identificar desgastes y prevenir fallas. Se debe seguir un cronograma establecido y registrar cada intervención.",
                "- Ejecutar reparaciones en equipos de climatización al detectar fallas en su funcionamiento, restableciendo su operatividad de manera oportuna y registrando las acciones realizadas",
                "- Implementar un monitoreo continuo de los parámetros de confort térmico entre 23 a 26°C en verano y 18 a 21 en invierno, manteniendo un registro sistemático de las mediciones y ajustes efectuados"
            ],
            'plazo': 'Continuo'
        }
    ]

    # Agregar secciones al documento
    doc.add_heading("4.1 Medidas de Carácter Ingenieril", level=3)
    if medidas_ingenieriles:
        crear_tabla_recomendaciones(doc, "Ingenieril", medidas_ingenieriles)
    else:
        doc.add_paragraph("No se requieren medidas ingenieriles para este caso.")

    doc.add_paragraph()
    doc.add_heading("4.2 Medidas de Carácter Administrativo", level=3)
    crear_tabla_recomendaciones(doc, "Administrativa", medidas_administrativas)

# Función para generar texto de áreas cumplen/no cumplen
def generar_texto_areas(areas, tipo):
    if not areas:
        return ""
    areas_formateadas = join_with_and(areas)
    if tipo == "cumplen":
        return f"las áreas {areas_formateadas} cumplen con el estándar de confort térmico"
    else:
        return f"las áreas {areas_formateadas} no cumplen con el estándar"

######

# -----------------------------------------------
# FUNCIÓN PARA GENERAR EL DOCUMENTO WORD
# -----------------------------------------------
def generar_informe_en_word(df_centros, df_visitas, df_mediciones, df_equipos) -> BytesIO:
    """
    Genera el informe en Word utilizando:
      - df_centros: información del centro de trabajo (tabla higiene_Centros_Trabajo)
      - df_visitas: información de visitas (tabla higiene_Visitas); se selecciona la visita más reciente.
      - df_mediciones: mediciones asociadas a la visita (tabla higiene_Mediciones)
      - df_equipos: información de equipos de medición (tabla higiene_Equipos_Medicion)
    """

    format_columns(df_visitas, ['nombre_personal_visita', 'consultor_ist'], mode="title")
    format_columns(df_visitas, 'cargo_personal_visita', mode="capitalize")
    format_columns(df_mediciones, ['nombre_area', 'sector_especifico','puesto_trabajo'], mode="capitalize")



    doc = Document()
    look_informe(doc)
    set_vertical_alignment(doc, section_index=0, alignment='top')

    # Cabecera con logo
    section = doc.sections[0]
    section.header_distance = Inches(0.4)
    header = section.header
    if header.paragraphs:
        paragraph = header.paragraphs[0]
    else:
        paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.add_picture('IST.jpg', width=Cm(2))

    # Título del informe: se alinea a la derecha
    titulo = doc.add_heading("INFORME EVALUACIÓN CONFORT TÉRMICO", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Párrafo con el código: se crea un párrafo nuevo, se añade el texto en negrita y se alinea a la derecha
    parrafo_codigo = doc.add_paragraph()
    run_codigo = parrafo_codigo.add_run("CODIGO: [COMPLETAR]")
    run_codigo.bold = True
    parrafo_codigo.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # A partir de aquí, el contenido se alineará a la izquierda (valor por defecto)
    paragraph = doc.add_heading("1. Antecedentes", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("Por solicitud del área de prevención de la empresa, se realiza evaluación de Confort Térmico, para determinar condición en que se encuentran los trabajadores que se desempeñan en áreas o sectores de trabajo, con el fin de que la empresa pueda adoptar oportuna y eficazmente medidas que permitan mejorar las condiciones evaluadas, según corresponda.")
    doc.add_paragraph()

    # -------------------------------
    # 1) IDENTIFICACIÓN ACTIVIDAD
    # -------------------------------

    if not df_centros.empty and not df_visitas.empty:
        row_centro = df_centros.iloc[0]
        row_visita = df_visitas.iloc[0]

        # Tabla 1: Información de la empresa
        table_empresa = doc.add_table(rows=0, cols=2)
        table_empresa.style = 'Table Grid'
        add_row(table_empresa, "1.1 Información empresa")
        add_row(table_empresa, "Razón Social", row_centro.get('razon_social', ''))
        add_row(table_empresa, "RUT", row_centro.get('rut', ''))
        add_row(table_empresa, "CIIU", row_centro.get('CIIU', '[COMPLETAR]'))

        # Espacio entre tablas (opcional)
        doc.add_paragraph()

        # Tabla 2: Información centro de trabajo
        table_centro = doc.add_table(rows=0, cols=2)
        table_centro.style = 'Table Grid'
        add_row(table_centro, "1.2 Información centro de trabajo")
        add_row(table_centro, "CUV / Código IST", row_centro.get('cuv', ''))
        add_row(table_centro, "Nombre de Local", row_centro.get('nombre_ct', ''))
        add_row(table_centro, "Dirección", row_centro.get('direccion_ct', ''))
        add_row(table_centro, "Comuna", row_centro.get('comuna_ct', ''))
        add_row(table_centro, "Región", row_centro.get('region_ct', ''))

        # Espacio entre tablas (opcional)
        doc.add_paragraph()

        # Tabla 3: Información de la visita
        table_visita = doc.add_table(rows=0, cols=2)
        table_visita.style = 'Table Grid'
        add_row(table_visita, "1.3 Información de la visita")
        #add_row(table_visita, "Motivo de la actividad", row_visita.get('motivo_evaluacion', ''))
        add_row(table_visita, "Motivo de la actividad", "Programa de trabajo")
        add_row(table_visita, "Fecha actividad de terreno", row_visita.get('fecha_visita', ''))
        add_row(table_visita, "Hora actividad de terreno", row_visita.get('hora_visita', ''))
        add_row(table_visita, "Temperatura ambiental exterior", f"{row_visita.get('temperatura_dia', '')}°C")
        add_row(table_visita, "Fecha emisión informe", "[COMPLETAR]")
        add_row(table_visita, "Profesional consultor/a de IST", row_visita.get('consultor_ist', ''))
        add_row(table_visita, "Acompañante empresa", row_visita.get('nombre_personal_visita', ''))
        add_row(table_visita, "Cargo de la persona que acompaña visita", row_visita.get('cargo_personal_visita', ''))
        add_row(table_visita, "Revisor del informe", "[COMPLETAR]")
        add_row(table_visita, "Jefatura responsable IST", "[COMPLETAR]")
        add_row(table_visita, "Destinatario informe", "[COMPLETAR]")

    else:
        # En caso de no existir información, se crea una tabla con el mensaje de error
        table_error = doc.add_table(rows=0, cols=2)
        add_row(table_error, "Información de empresa/centro", "No se encontró información para este CUV.")

    # Encabezado principal del contenido: Metodología
    paragraph = doc.add_heading("2. Metodología", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Primer párrafo
    doc.add_paragraph(
        "El concepto de “confort térmico” describe el estado mental de una persona en términos de percibir un ambiente demasiado caluroso o demasiado frío. "
        "Por otro lado, también se puede definir como una manifestación subjetiva de conformidad o satisfacción entre el trabajador con el ambiente térmico existente."
    )

    # Segundo párrafo con palabras en negrita
    doc.add_paragraph()
    p = doc.add_paragraph(
        "El presente informe utiliza la metodología de FANGER para evaluación de confort térmico en espacios interiores de acuerdo a la Nota técnica N°47 del Instituto de Salud Pública. "
        "De esta forma, los diferentes puestos de trabajo son evaluados y calificados en cada caso como "
    )
    p.add_run('"Cumple"').bold = True
    p.add_run(" o ")
    p.add_run('"No cumple"').bold = True

    # Tercer párrafo con palabras en negrita
    doc.add_paragraph()
    p2 = doc.add_paragraph(
        "Los alcances de las calificaciones específicas para cada área o sector evaluado, corresponde al cumplimiento del Voto Medio Estimado "
    )
    p2.add_run("PMV").bold = True
    p2.add_run(
        " (Predicted Mean Vote), equivalente a una condición media deseable que indican la sensación térmica media de un entorno y ")
    p2.add_run("PPD").bold = True
    p2.add_run(
        " (Predicted Percentage Dissatisfied) correspondiente al porcentaje de personas que sentirán algún grado de disconfort en un ambiente de trabajo evaluado.")

    # -------------------------------
    # Resultados de mediciones y evaluación
    # -------------------------------
    # Crear listas para las áreas que cumplen y no cumplen
    areas_cumplen = []
    areas_no_cumplen = []

    paragraph = doc.add_heading("3. Resultados de las mediciones y evaluación", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def generar_tabla_resumen(doc, df_mediciones):
        # Verificamos si el DataFrame tiene datos
        if df_mediciones.empty:
            doc.add_paragraph("No se encontraron mediciones para detallar.")
            return

        # Definimos las columnas en el orden requerido
        columnas_resumen = [
            "Área",
            "Estándar confortabilidad",
            "Puesto de trabajo",
            "Temp. bulbo seco(°C)",
            "Temp. globo(°C)",
            "Humedad relativa (%)",
            "Velocidad del aire(m/s)",
            "PPD",
            "PMV"
        ]

        # Creamos la tabla con tantas columnas como la lista anterior
        tabla_resumen = doc.add_table(rows=1, cols=len(columnas_resumen))
        tabla_resumen.style = 'Table Grid'

        # Encabezados
        hdr_cells = tabla_resumen.rows[0].cells
        for idx, col_name in enumerate(columnas_resumen):
            hdr_cells[idx].text = col_name

        # Agrupamos por área
        grouped = df_mediciones.groupby("nombre_area")

        # Recorremos cada grupo (cada área)
        for area, group in grouped:
            # Si hay múltiples mediciones en el área, calculamos promedio;
            # si solo hay una, usamos directamente esa.
            if len(group) > 1:
                # Calculamos promedios
                avg_t_bul = group["t_bul_seco"].astype(float).mean()
                avg_t_globo = group["t_globo"].astype(float).mean()
                avg_hum = group["hum_rel"].astype(float).mean()
                avg_vel = group["vel_air"].astype(float).mean()
                avg_met = group["met"].astype(float).mean()
                avg_clo = group["clo"].astype(float).mean()

                # Calcular pmv/ppd a partir de la función pmv_ppd_iso
                # (ajusta según tu propia lógica)
                try:
                    avg_ppd = pmv_ppd_iso(
                        tdb=avg_t_bul,
                        tr=avg_t_globo,
                        vr=avg_vel,
                        rh=avg_hum,
                        met=avg_met,  # Ajusta si necesitas
                        clo=avg_clo,  # Ajusta si necesitas
                        model="7730-2005",
                        limit_inputs=False,
                        round_output=True
                    ).ppd
                    avg_pmv = pmv_ppd_iso(
                        tdb=avg_t_bul,
                        tr=avg_t_globo,
                        vr=avg_vel,
                        rh=avg_hum,
                        met=avg_met,  # Ajusta si necesitas
                        clo=avg_clo,  # Ajusta si necesitas
                        model="7730-2005",
                        limit_inputs=False,
                        round_output=True
                    ).pmv
                except Exception as e:
                    logging.error("Error al calcular pmv_ppd_iso para el área %s: %s", area, e)
                #    avg_ppd = 0
                #    avg_pmv = 0

                # Interpretación del PMV para mostrar "CUMPLE" o "NO CUMPLE"
                analisis = interpret_pmv(avg_pmv)

                # Puesto de trabajo
                valores_unicos = list(OrderedDict.fromkeys(group["puesto_trabajo"].dropna()))
                puesto_trabajo = "\n".join(str(x) for x in valores_unicos)

                # Creamos la fila final
                row_cells = tabla_resumen.add_row().cells
                paragraph_0 = row_cells[0].paragraphs[0]
                run_0 = paragraph_0.add_run(str(area))
                run_0.bold = True
                paragraph_1 = row_cells[1].paragraphs[0]
                run_1 = paragraph_1.add_run(analisis.upper())
                run_1.bold = True
                row_cells[2].text = puesto_trabajo
                row_cells[3].text = f"{avg_t_bul:.2f}"
                row_cells[4].text = f"{avg_t_globo:.2f}"
                row_cells[5].text = f"{avg_hum:.1f}"
                row_cells[6].text = f"{avg_vel:.2f}"
                row_cells[7].text = f"{avg_ppd:.2f}"
                row_cells[8].text = f"{avg_pmv:.2f}"

            else:
                # Solo hay una medición en el área, la usamos directamente
                row = group.iloc[0]
                t_bul = float(row.get("t_bul_seco", 0))
                t_globo = float(row.get("t_globo", 0))
                hum = float(row.get("hum_rel", 0))
                vel = float(row.get("vel_air", 0))
                ppd = float(row.get("ppd", 0))
                pmv = float(row.get("pmv", 0))
                analisis = row.get("resultado_medicion", "")  # "CUMPLE" o "NO CUMPLE"
                puesto_trabajo = str(row.get("puesto_trabajo", ""))

                # Creamos la fila con los datos únicos
                row_cells = tabla_resumen.add_row().cells
                paragraph_0 = row_cells[0].paragraphs[0]
                run_0 = paragraph_0.add_run(str(area))
                run_0.bold = True
                paragraph_1 = row_cells[1].paragraphs[0]
                run_1 = paragraph_1.add_run(analisis.upper())
                run_1.bold = True
                row_cells[2].text = puesto_trabajo
                row_cells[3].text = f"{t_bul:.2f}"
                row_cells[4].text = f"{t_globo:.2f}"
                row_cells[5].text = f"{hum:.1f}"
                row_cells[6].text = f"{vel:.2f}"
                row_cells[7].text = f"{ppd:.2f}"
                row_cells[8].text = f"{pmv:.2f}"

        # Opcional: Ajustar anchos de columna si lo deseas
        set_column_width(tabla_resumen, 0, Cm(3))
        set_column_width(tabla_resumen, 1, Cm(3))
        set_column_width(tabla_resumen, 2, Cm(3))
        set_column_width(tabla_resumen, 3, Cm(1.5))
        set_column_width(tabla_resumen, 4, Cm(1.5))
        set_column_width(tabla_resumen, 5, Cm(1.5))
        set_column_width(tabla_resumen, 6, Cm(1.5))
        set_column_width(tabla_resumen, 7, Cm(1.5))
        set_column_width(tabla_resumen, 8, Cm(1.5))

        # Formato de la fila de encabezado (opcional)
        format_row(tabla_resumen.rows[0])

    generar_tabla_resumen(doc, df_mediciones)

    # Procesar áreas antes del resumen
    areas_cumplen, areas_no_cumplen = procesar_areas(df_mediciones)

    # Encabezado principal del contenido: Conclusiones
    doc.add_paragraph()

    if not df_centros.empty:
        row_centro = df_centros.iloc[0]
        nombre_ct = row_centro.get("nombre_ct", "")
        # Construir las cadenas de texto según las áreas que cumplen y las que no cumplen

        if areas_cumplen:
            if len(areas_cumplen) == 1:
                # Forma singular
                cumplen_text = f"el área de {areas_cumplen[0]}"
            else:
                # Forma plural, usando la función para unir con "y"
                cumplen_text = f"las áreas {join_with_and(areas_cumplen)}"
        else:
            cumplen_text = None  # O dejarlo en cadena vacía, según convenga

        if areas_no_cumplen:
            if len(areas_no_cumplen) == 1:
                no_cumplen_text = f"el área {areas_no_cumplen[0]}"
            else:
                no_cumplen_text = f"las áreas {join_with_and(areas_no_cumplen)}"
        else:
            no_cumplen_text = None

        # Generar la redacción final de las conclusiones

    # Generar la redacción final de las conclusiones
    if cumplen_text and no_cumplen_text:
        # Caso 2: Existen áreas que cumplen y áreas que no cumplen.
        doc.add_paragraph(
            f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, se concluye que {cumplen_text} "
            f"{'cumple' if len(areas_cumplen) == 1 else 'cumplen'} con el estándar de confort térmico, establecido mediante la metodología de Fanger. "
            f"Esto significa que, al registrarse un PMV entre -1 y +1, el PPD, que expresa el porcentaje de personas que experimentan disconfort con la temperatura, "
            f"resulta inferior al 25%. Por ello, se recomienda mantener las condiciones actuales o similares."
        )
        doc.add_paragraph()
        doc.add_paragraph(
            f"Respecto de {no_cumplen_text} {'NO cumple' if len(areas_no_cumplen) == 1 else 'NO cumplen'} con el mismo estándar, ya que el PMV se encuentra fuera del rango de -1 a +1, "
            f"lo que implica que el PPD resulta superior al 25%. Por lo tanto, se deben adoptar las medidas correctivas para alcanzar las condiciones de confort deseadas."
        )
    elif cumplen_text and not no_cumplen_text:
        # Caso 1: Sólo existen áreas que cumplen.
        doc.add_paragraph(
            f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, se concluye que {cumplen_text} "
            f"{'cumple' if len(areas_cumplen) == 1 else 'cumplen'} con el estándar de confort térmico, establecido mediante la metodología de Fanger. "
            f"Esto significa que, al registrarse un PMV entre -1 y +1, el PPD, que expresa el porcentaje de personas que experimentan disconfort con la temperatura, "
            f"resulta inferior al 25%. Por ello, se recomienda mantener las condiciones actuales o similares."
        )
    elif no_cumplen_text and not cumplen_text:
        # Caso 3: Sólo existen áreas que no cumplen.
        doc.add_paragraph(
            f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, se concluye que {no_cumplen_text} "
            f"{'NO cumple' if len(areas_no_cumplen) == 1 else 'NO cumplen'} con el estándar de confort térmico, establecido mediante la metodología de Fanger. "
            f"Esto significa que, al registrarse un PMV fuera del rango de -1 a +1, el PPD, que expresa el porcentaje de personas que experimentan disconfort con la temperatura, "
            f"resulta superior al 25%. Por lo tanto, se deben adoptar las medidas correctivas para alcanzar las condiciones de confort deseadas."
        )
    else:
        # En caso de que no haya información suficiente
        doc.add_paragraph(
            "No se encontró información suficiente en las mediciones para emitir una conclusión sobre el confort térmico."
        )

    '''
    if not df_centros.empty:
        row_centro = df_centros.iloc[0]
        nombre_ct = row_centro.get("nombre_ct", "")
        # Construir las cadenas de texto según las áreas que cumplen y las que no cumplen

        if areas_cumplen:
            if len(areas_cumplen) == 1:
                # Forma singular
                cumplen_text = f"el área de {areas_cumplen[0]}"
            else:
                # Forma plural, usando la función para unir con "y"
                cumplen_text = f"las áreas {join_with_and(areas_cumplen)}"
        else:
            cumplen_text = None  # O dejarlo en cadena vacía, según convenga

        if areas_no_cumplen:
            if len(areas_no_cumplen) == 1:
                no_cumplen_text = f"el área {areas_no_cumplen[0]}"
            else:
                no_cumplen_text = f"las áreas {join_with_and(areas_no_cumplen)}"
        else:
            no_cumplen_text = None

        # Generar la redacción final de las conclusiones
        if cumplen_text and no_cumplen_text:
            # Caso 2: Existen áreas que cumplen y áreas que no cumplen.
            doc.add_paragraph(
                f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, es posible concluir que {cumplen_text} "
                f"{'cumple' if len(areas_cumplen) == 1 else 'cumplen'} con el estándar de confort térmico, por lo que se recomienda mantener las condiciones actuales o similares."
            )
            doc.add_paragraph(
                f"Respecto a {no_cumplen_text} que NO {'cumple' if len(areas_no_cumplen) == 1 else 'cumplen'} con el estándar, se deben adoptar las medidas prescritas a continuación para corregir las condiciones."
            )
        elif cumplen_text and not no_cumplen_text:
            # Caso 1: Sólo existen áreas que cumplen.
            doc.add_paragraph(
                f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, se concluye que {cumplen_text} "
                f"{'cumple' if len(areas_cumplen) == 1 else 'cumplen'} con el estándar de confort térmico, por lo que se recomienda mantener las condiciones actuales o similares."
            )
        elif no_cumplen_text and not cumplen_text:
            # Caso 3: Sólo existen áreas que no cumplen.
            doc.add_paragraph(
                f"Efectuadas mediciones de confort térmico en el local {nombre_ct}, se concluye que {no_cumplen_text} "
                f"{'NO cumple' if len(areas_no_cumplen) == 1 else 'NO cumplen'} con el estándar de confort térmico, por lo que se deben adoptar las medidas prescritas a continuación para corregir las condiciones."
            )
        else:
            # En caso de que no haya información suficiente
            doc.add_paragraph(
                "No se encontró información suficiente en las mediciones para emitir una conclusión sobre el confort térmico."
            )

    '''

    # -------------------------------
    # 4) MEDIDAS CORRECTIVAS
    # -------------------------------

    paragraph = doc.add_heading("4. Prescripción de medidas", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Asumiendo que row_centro ya está definido y contiene la información de la empresa:
    razon_social = row_centro.get('razon_social', 'RENDIC HERMANOS S.A.')

    # Luego, en el cuerpo del documento:
    doc.add_paragraph(
        "Conforme al artículo 68 de la Ley N° 16.744, la implementación de las medidas prescritas por este organismo "
        "administrador es de carácter obligatoria, por lo que su incumplimiento podrá ser sancionado con el recargo de "
        "la cotización adicional diferenciada, sin perjuicio de las demás sanciones que correspondan."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f"No obstante, {razon_social} podrá implementar otras medidas técnicas y/o administrativas equivalentes a las "
        "señaladas en el presente informe y que contribuyan a disminuir la exposición de sus trabajadores, debiendo "
        "informar a IST, quien evaluará su efectividad una vez implementadas. Adicionalmente, en el caso de que las áreas "
        "de trabajo sean operadas por contratistas, el mandante debe informar obligatoriamente a todos sus contratistas los "
        "riesgos a los que están expuestos."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Acorde a las condiciones existentes al momento de las mediciones, al resultado de las mismas y a las conclusiones "
        "obtenidas, se establecen las siguientes medidas de control:"
    )

    agregar_medidas_correctivas(doc, df_mediciones, areas_no_cumplen)

    doc.add_paragraph()

    # -------------------------------
    # 5) VIGENCIA DEL INFORME
    # -------------------------------
    # Encabezado principal del contenido: Vigencia del informe
    paragraph = doc.add_heading("5. Vigencia del informe", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not df_visitas.empty:
        doc.add_paragraph("En términos generales, el presente informe tiene validez de 3 años, a excepción que existan cambios en la situación, del tipo ingenieril o administrativo, que presupongan modificación a las condiciones encontradas al momento de la medición, lo cual implicará realizar una nueva evaluación en un plazo menor al señalado.")
        doc.add_paragraph()
        if len(areas_no_cumplen) > 0:
            doc.add_paragraph("Cuando se concreten los cambios indicados, la empresa deberá informar al IST el detalle de los mismos, de forma tal de programar las gestiones a realizar, las que considerarán previamente un informe de Verificación y Control y posteriormente, de acuerdo a sus resultados, la nueva evaluación de higiene ocupacional correspondiente.")
            doc.add_paragraph()
        doc.add_paragraph("Estos resultados de evaluación representan las condiciones existentes del ambiente y lugar de trabajo al momento de realizar las mediciones.")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()


    if not df_visitas.empty:
        row_visita = df_visitas.iloc[0]
        consultor_ist = row_visita.get("consultor_ist", "")

    # Agregar párrafo para el consultor, centrado y en negrita
    p_consultor = doc.add_paragraph()
    p_consultor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_consultor = p_consultor.add_run(consultor_ist)
    run_consultor.bold = True

    # Agregar párrafo para la profesión, centrado
    p_profesion = doc.add_paragraph("[__COMPLETAR_PROFESIÓN__]")
    p_profesion.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar párrafo para el zonal, centrado
    p_zonal = doc.add_paragraph("[__COMPLETAR_ZONAL__]")
    p_zonal.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # -------------------------------
    # 6) ANEXOS
    # -------------------------------
    doc.add_page_break()
    doc.add_heading("Anexo 1. Consideraciones técnicas de la evaluación", level=2)

    # Agrega un párrafo con el título para la tabla
    doc.add_heading("a)     Verificación en terreno de parámetros de los equipos", level=3)

    # Crea la tabla con 4 columnas y aplica un estilo
    table_calib = doc.add_table(rows=0, cols=4)
    table_calib.style = 'Table Grid'

    # Agrega la fila de encabezado
    hdr_cells = table_calib.add_row().cells
    hdr_cells[0].text = "Temperatura"
    hdr_cells[1].text = "Patrón"
    hdr_cells[2].text = "Verificación inicial"
    hdr_cells[3].text = "Verificación final"

    # Si hay datos en df_visitas, agrega las filas para cada equipo
    if not df_visitas.empty:
        # Para TBS
        row_cells = table_calib.add_row().cells
        row_cells[0].text = "TBS"
        row_cells[1].text = str(row_visita.get('patron_tbs', ''))
        row_cells[2].text = str(row_visita.get('ver_tbs_ini', ''))
        row_cells[3].text = str(row_visita.get('ver_tbs_fin', ''))

        # Para TBH
        row_cells = table_calib.add_row().cells
        row_cells[0].text = "TBH"
        row_cells[1].text = str(row_visita.get('patron_tbh', ''))
        row_cells[2].text = str(row_visita.get('ver_tbh_ini', ''))
        row_cells[3].text = str(row_visita.get('ver_tbh_fin', ''))

        # Para TG
        row_cells = table_calib.add_row().cells
        row_cells[0].text = "TG"
        row_cells[1].text = str(row_visita.get('patron_tg', ''))
        row_cells[2].text = str(row_visita.get('ver_tg_ini', ''))
        row_cells[3].text = str(row_visita.get('ver_tg_fin', ''))
    else:
        # Si no hay información, se agrega una fila de error
        row_cells = table_calib.add_row().cells
        row_cells[0].text = "D. Detalles de equipos y calibración"
        row_cells[1].text = "No se encontró información de visita."
        # Se pueden unir las celdas restantes para que el mensaje quede centrado
        merged = row_cells[0].merge(row_cells[1])
        merged = merged.merge(row_cells[2]).merge(row_cells[3])

    # Configurar el ancho de cada columna (opcional)
    set_column_width(table_calib, 0, Cm(4.25))
    set_column_width(table_calib, 1, Cm(4.25))
    set_column_width(table_calib, 2, Cm(4.25))
    set_column_width(table_calib, 3, Cm(4.25))
    format_row(table_calib.rows[0])


    doc.add_paragraph()
    doc.add_heading("b)     Caracterización de vestimenta utilizada y tasa metabólica", level=3)

    doc.add_paragraph("Para efectos del presente informe, se considera el valor de 0,5 Clo, que es equivalente a ropa normal de trabajo y una tasa metabólica (Mets) de 1,1 a 1,2 (valor equivalente a 109 kcal/hrs.) en función de los componentes de la actividad para el área Línea de Cajas o Sala de ventas, y una tasa metabólica de 1,89 Mets (valor equivalente a 187 kcal/hrs.), en función de la profesión, para el área Bodega o Recepción.")

    doc.add_paragraph()
    doc.add_heading("c)     Características generales de las areas evaluadas", level=3)

    # Crear la tabla con 3 columnas: Área, Características constructivas y Condiciones de ventilación
    tabla_caract = doc.add_table(rows=1, cols=3)
    tabla_caract.style = 'Table Grid'

    # Agregar la fila de encabezado
    hdr_cells = tabla_caract.rows[0].cells
    hdr_cells[0].text = "Área"
    hdr_cells[1].text = "Características constructivas"
    hdr_cells[2].text = "Condiciones de ventilación"

    # Agrupar el DataFrame por "nombre_area"
    grouped = df_mediciones.groupby("nombre_area")
    for area, group in grouped:
        # Usar el primer registro del grupo para extraer los datos de instalación
        registro = group.iloc[0]

        # Agregar una fila para el área con sus respectivos datos
        row_cells = tabla_caract.add_row().cells
        row_cells[0].text = area
        row_cells[1].text = str(registro["caract_constructivas"])
        row_cells[2].text = str(registro["ingreso_salida_aire"])

    set_column_width(tabla_caract, 0, Cm(3))
    set_column_width(tabla_caract, 1, Cm(7))
    set_column_width(tabla_caract, 2, Cm(7))
    format_row(tabla_caract.rows[0])




    # Salto de página y título del anexo
    doc.add_page_break()
    doc.add_heading("Anexo 2. Instrumentos de medición utilizados", level=2)

    if not df_visitas.empty and not df_equipos.empty:
        row_visita = df_visitas.iloc[0]
        # Obtener los códigos de equipos que están en uso en la visita
        equipo_temp_cod = row_visita.get('equipo_temp', '')
        equipo_vel_cod = row_visita.get('equipo_vel_air', '')
        codigos_en_uso = [equipo_temp_cod, equipo_vel_cod]

        # Filtrar df_equipos para que solo incluya las filas donde 'id_equipo' está en codigos_en_uso
        df_equipos_filtrado = df_equipos[df_equipos['id_equipo'].isin(codigos_en_uso)]

        # Definir el mapeo de campos a mostrar
        field_mapping = {
            "nombre_equipo": "Tipo de equipo",
            "cod_equipo": "Código",
            "n_serie_equipo": "Número de serie",
            "marca_equipo": "Marca",
            "modelo_equipo": "Modelo",
            "fecha_calibracion": "Última calibración",
            "prox_calibracion": "Próxima calibración",
            "empresa_certificadora": "Empresa certificadora",
            "num_certificado": "Número de certificado",
            "url_certificado": "Respaldo certificado"
        }

        if not df_equipos_filtrado.empty:
            for idx, row_eq in df_equipos_filtrado.iterrows():
                # Crear una tabla de dos columnas para los datos del equipo
                tabla_equipo = doc.add_table(rows=len(field_mapping), cols=2)
                tabla_equipo.style = 'Table Grid'
                for row_num, (key, display_name) in enumerate(field_mapping.items()):
                    tabla_equipo.rows[row_num].cells[0].text = display_name
                    if key == "url_certificado":
                        url = str(row_eq.get(key, ""))
                        if url.strip():
                            # Genera un código QR (función asumida)
                            qr_img = generate_qr_code(url)
                            cell = tabla_equipo.rows[row_num].cells[1]
                            cell.text = ""
                            run = cell.paragraphs[0].add_run()
                            run.add_break()
                            run.add_picture(qr_img, width=Inches(1))
                            run.add_break()
                        else:
                            tabla_equipo.rows[row_num].cells[1].text = ""
                    else:
                        tabla_equipo.rows[row_num].cells[1].text = str(row_eq.get(key, ""))
                # Ajustar anchos de columnas para toda la tabla (se recomienda hacerlo fuera del bucle interno)
                set_column_width(tabla_equipo, 0, Cm(3.5))
                set_column_width(tabla_equipo, 1, Cm(13.5))

                doc.add_paragraph("")  # Separador entre tablas

            for idx, row_eq in enumerate(df_equipos_filtrado.itertuples(), 1):
                id_equipo = str(row_eq.id_equipo)  # Asegúrate que este campo coincide con tus directorios

                # Ruta al directorio de imágenes para este equipo
                img_dir = os.path.join("imagenes_pdf", id_equipo)

                try:
                    if os.path.exists(img_dir) and os.path.isdir(img_dir):
                        # Obtener todas las imágenes ordenadas numéricamente
                        imagenes = natsorted([
                            os.path.join(img_dir, f)
                            for f in os.listdir(img_dir)
                            if f.lower().endswith(('.png', '.jpg', '.jpeg'))
                        ])

                        # Insertar todas las imágenes en el documento
                        for img_path in imagenes:
                            # Añadir imagen ocupando el ancho completo de la página
                            doc.add_picture(img_path, width=Cm(17))
                    else:
                        doc.add_paragraph(f"No se encontraron imágenes para el equipo {id_equipo}")
                except Exception as e:
                    doc.add_paragraph(f"Error al cargar imágenes para equipo {id_equipo}: {str(e)}")

        else:
            doc.add_paragraph("No se encontró información de equipos de medición relacionados con la visita.")
    else:
        doc.add_paragraph("No se encontró información de la visita o de los equipos.")

    # (Continúa el resto del script si es necesario)

    # -------------------------------
    # Finaliza el documento y lo retorna como BytesIO
    # -------------------------------
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer