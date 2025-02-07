############### Parte 1: Importaciones, Configuración Inicial y Conexión a la Base de Datos ################

# Importar las bibliotecas necesarias
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import unicodedata
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import logging
import os
import base64
from datetime import timedelta, datetime, date
from st_aggrid import AgGrid, GridOptionsBuilder
from dotenv import load_dotenv
import mysql.connector
from mysql.connector import Error
from streamlit_cookies_controller import CookieController
from utils.helpers import autenticar_usuario, login, check_session, logout
import time

cookie_controller = CookieController(key="app_cookies")

# Inicializar sesión del usuario
if "data_user" not in st.session_state:
    st.session_state["data_user"] = None

# Cargar variables de entorno desde un archivo .env
load_dotenv()

# Configuración de logging para monitorear la aplicación.
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Configuración de la base de datos para SQL Server utilizando variables de entorno
server = os.getenv('DB_HOST')  # Usa DB_HOST en lugar de DB_SERVER
database = os.getenv('DB_NAME')
username = os.getenv('DB_USERNAME')
password = os.getenv('DB_PASS')
port = os.getenv('DB_PORT')

def make_sidebar():
    with st.sidebar:
        if st.session_state.get("data_user", None):
            st.write(f"👋 Hola, {st.session_state['data_user']['nombre']}!")

            if st.button("Cerrar sesión"):
                cookie_controller.set("user_data", "", max_age=0)
                # cookie_controller.remove("user_data")
                st.session_state["data_user"] = None
                st.success("Has cerrado sesión correctamente.")
                time.sleep(2)
                st.rerun()

def checkear_session():
    user_data = cookie_controller.get("user_data")
    if user_data:
        st.session_state["data_user"] = user_data
    return user_data

# Lógica de login persistente
if not checkear_session():
    # Mostrar formulario de login
    st.title("Login")
    username = st.text_input("Usuario:", key="username_input")
    password = st.text_input("Contraseña:", type="password", key="password_input")

    if st.button("Iniciar sesión"):

            user_data = autenticar_usuario(username, password)
            if user_data:
                # Guardar datos en cookies
                cookie_controller.set("user_data", user_data)
                st.session_state["data_user"] = user_data
                st.success("Login exitoso")
                time.sleep(1)
                # st.switch_page("pages/home.py")
                st.rerun()
            else:
                st.error("Usuario o contraseña incorrectos.")

else:
    # Configurar la página principal de Streamlit
    st.set_page_config(
        page_title="Generador de Informes de Riesgos Psicosociales",
        layout="wide"
    )
    # Título principal de la aplicación
    st.title("Generador de Informes de Riesgos Psicosociales")
    st.write("""
    Esta aplicación permite generar informes técnicos basados en datos de riesgos psicosociales.
    Por favor, cargue los archivos necesarios y siga las instrucciones.
    """)
    make_sidebar()

    # Función para establecer la conexión con la base de datos
    def get_db_connection():
        """
        Establece una conexión con la base de datos SQL Server.
        Retorna el objeto de conexión si es exitosa, de lo contrario detiene la aplicación.
        """
        try:
            connection = mysql.connector.connect(
                host=server,
                database=database,
                user=username,
                password=password,
                port=port
            )
            if connection.is_connected():
                logging.info("Conexión a la base de datos establecida exitosamente.")
                return connection
        except Error as e:
            st.error(f"Error al conectar a la base de datos: {e}")
            logging.error(f"Error al conectar a la base de datos: {e}")
            st.stop()


    # Función para consultar una tabla específica filtrando por CUV
    def consultar_tabla(tabla, cuv=None, columnas=None):
        """
        Realiza una consulta SQL a una tabla específica filtrando por CUV si se proporciona.

        Parámetros:
        - tabla (str): Nombre de la tabla en la base de datos.
        - cuv (str, opcional): Valor del CUV a filtrar.
        - columnas (list[str], opcional): Lista de columnas a seleccionar. Si no se especifica, se seleccionan todas.

        Retorna:
        - pd.DataFrame: DataFrame con los resultados de la consulta.
        """
        tablas_permitidas = {
            "informe_CEAL__Summary",
            "informe_CEAL__basecompleta",
            "informe_CEAL__df_porcentajes_niveles",
            "informe_CEAL__df_res_dimTE3",
            "informe_CEAL__df_resumen",
            "informe_CEAL__resultado",
            "informe_CEAL__top_glosas",
            "informe_CEAL__fileresultados",
            "informe_CEAL__ciiu",
            "informe_CEAL__rec"
        }

        if tabla not in tablas_permitidas:
            st.error(f"Tabla '{tabla}' no permitida.")
            return pd.DataFrame()

        columnas_sql = ", ".join(columnas) if columnas else "*"
        query = f"SELECT {columnas_sql} FROM {tabla}"
        params = [cuv] if cuv else None
        if cuv:
            query += " WHERE CUV = %s"  # Usa %s si estás en MySQL, o ? si es SQLite

        connection = get_db_connection()

        if connection:
            try:
                # Usar read_sql_query con parámetros
                if params:
                    df = pd.read_sql_query(query, connection, params=params)
                else:
                    df = pd.read_sql_query(query, connection)

                logging.info(
                    f"Consulta ejecutada en la tabla '{tabla}' para CUV: {cuv}" if cuv else f"Consulta ejecutada en la tabla '{tabla}'"
                )
                return df
            except Exception as e:
                st.error(f"Error al consultar la tabla '{tabla}': {e}")
                logging.error(f"Error al consultar la tabla '{tabla}': {e}")
                return pd.DataFrame()
            finally:
                connection.close()
                logging.info("Conexión a la base de datos cerrada.")
        else:
            return pd.DataFrame()


    # Función para extraer y validar 'codigo_ciiu'
    def extraer_codigo_ciiu(df, columna='CIIU_CT'):
        """
        Extrae y valida el 'codigo_ciiu' de la columna especificada en el DataFrame.

        Parámetros:
        - df (pd.DataFrame): DataFrame que contiene la columna 'CIIU_CT'.
        - columna (str): Nombre de la columna de donde extraer el 'codigo_ciiu'.

        Retorna:
        - int o None: 'codigo_ciiu' si es válido, de lo contrario None.
        """
        if columna not in df.columns:
            st.error(f"No se encontró la columna '{columna}' en la tabla 'Base Completa'.")
            return None

        # Extraer el primer valor no nulo de 'CIIU_CT'
        primer_valor = df[columna].dropna().iloc[0] if not df[columna].dropna().empty else None

        if not primer_valor:
            st.error("No se encontró un valor válido en 'CIIU_CT'.")
            return None

        # Procesar el valor para extraer 'codigo_ciiu'
        if isinstance(primer_valor, str):
            partes = primer_valor.split('_')
            if len(partes) < 2:
                st.error(f"El formato de '{columna}' es inválido: '{primer_valor}'. Se esperaba al menos un '_' separador.")
                return None
            codigo_ciiu_str = partes[-1]
        else:
            codigo_ciiu_str = str(primer_valor)

        # Validar que 'codigo_ciiu' sea numérico
        if not codigo_ciiu_str.isdigit():
            st.error(f"El valor de 'codigo_ciiu' extraído no es numérico: '{codigo_ciiu_str}'.")
            return None

        # Convertir a entero
        codigo_ciiu = int(codigo_ciiu_str)

        # Validar la longitud del código (ejemplo: asumiendo que debe tener entre 1 y 2 dígitos)
        if len(codigo_ciiu_str) > 5:
            codigo_ciiu = int(codigo_ciiu_str[:2])
        elif len(codigo_ciiu_str) > 1:
            codigo_ciiu = int(codigo_ciiu_str[:1])

        return codigo_ciiu


    # Función para procesar columnas de fecha
    def procesar_columna_fecha(df, columna, formato='%d-%m-%Y'):
        """
        Convierte una columna de fechas en el DataFrame al formato deseado.

        Parameters:
        - df (pd.DataFrame): El DataFrame que contiene la columna.
        - columna (str): El nombre de la columna a procesar.
        - formato (str): El formato al que se desea convertir la fecha (por defecto '%d-%m-%Y').

        Returns:
        - pd.DataFrame: El DataFrame con la columna procesada.
        """
        if columna in df.columns:
            # Convertir la columna a tipo datetime automáticamente
            df[columna] = pd.to_datetime(df[columna], errors='coerce')

            # Aplicar el formato deseado si la conversión a datetime fue exitosa
            df[columna] = df[columna].apply(lambda x: x.strftime(formato) if pd.notna(x) else 'N/A')
            return df
        else:
            raise ValueError(f"La columna '{columna}' no se encontró en el DataFrame.")




        ############### Parte 2: Definición de Tablas y Búsqueda por CUV ################

    # Lista de tablas a consultar en la base de datos
    tablas_a_consultar = [
        "informe_CEAL__Summary",
        "informe_CEAL__basecompleta",
        "informe_CEAL__df_porcentajes_niveles",
        "informe_CEAL__df_res_dimTE3",
        "informe_CEAL__df_resumen",
        "informe_CEAL__resultado",
        "informe_CEAL__top_glosas",
        "informe_CEAL__fileresultados"
    ]

    # Mapeo de nombres de tablas a nombres amigables para su visualización en la interfaz
    nombres_amigables = {
        "informe_CEAL__Summary": "Summary",
        "informe_CEAL__basecompleta": "BaseCompleta",
        "informe_CEAL__df_porcentajes_niveles": "Porcentajes Niveles",
        "informe_CEAL__df_res_dimTE3": "Res Dim TE3",
        "informe_CEAL__df_resumen": "Resumen",
        "informe_CEAL__resultado": "Resultado",
        "informe_CEAL__top_glosas": "Top Glosas",
        "informe_CEAL__fileresultados": "Filas Resultados",
        "informe_CEAL__ciiu": "CIIU",
        "informe_CEAL__rec": "Recomendaciones"
    }

    # Título y cuadro de texto para ingresar el CUV
    st.header("Aplicación de Búsqueda por CUV")
    cuv_valor = st.text_input("Ingresa el CUV que deseas buscar:", "")

    # Inicializar variables en st.session_state
    for var in ['combined_df_base_complet3', 'df_res_com', 'summary_df', 'df_porcentajes_niveles', 'df_res_dimTE3','confirmadas_df',
                'df_resumen', 'df_resultados_porcentaje', 'top_glosas', 'df_ciiu', 'df_recomendaciones', 'df_resultados', 'interpretaciones_df']:
        st.session_state.setdefault(var, pd.DataFrame())


    if 'interpretaciones_df' not in st.session_state or st.session_state['interpretaciones_df'].empty:
        st.warning("No se encontraron interpretaciones en la sesión.")
    else:
        interpretaciones_df = st.session_state['interpretaciones_df']

    if 'confirmadas_df' not in st.session_state or st.session_state['confirmadas_df'].empty:
        st.warning("No se encontraron medidas confirmadas en la sesión.")
    else:
        confirmadas_df = st.session_state['confirmadas_df']


    # Inicializar variables de estado al inicio
    if 'interpretaciones_temporales' not in st.session_state:
        st.session_state['interpretaciones_temporales'] = {}


    # Justo después de inicializar las variables en session_state
    if 'selected_cuv' not in st.session_state:
        st.session_state.selected_cuv = ""  # Inicializamos con un string vacío o None


    # Botón para ejecutar la búsqueda
    if st.button("Buscar"):
        if not cuv_valor.strip():
            st.warning("Por favor, ingresa un valor de CUV antes de continuar.")
        else:
            st.header(f"Resultados para CUV: {cuv_valor}")

            resultados = {}

            # Consultar recomendaciones y ciiu (sin filtrar por ahora, a menos que sea necesario)
            resultados["df_recomendaciones"] = consultar_tabla("informe_CEAL__rec")
            resultados["df_ciiu"] = consultar_tabla("informe_CEAL__ciiu")

            for tabla in tablas_a_consultar:
                nombre_amigable = nombres_amigables.get(tabla, tabla)

                if tabla == "informe_CEAL__fileresultados":
                    columnas_fileresultados = ['CUV', 'Folio']
                    df_resultados = consultar_tabla(tabla, cuv_valor, columnas=columnas_fileresultados)
                    df_res_com = consultar_tabla(tabla, cuv_valor)  # Todas las columnas
                    resultados["df_resultados"] = df_resultados
                    resultados["df_res_com"] = df_res_com
                else:
                    df = consultar_tabla(tabla, cuv_valor)
                    resultados[nombre_amigable] = df

            # Ahora todos los DataFrames en 'resultados' ya están filtrados por el CUV ingresado.
            # Pasamos estos DF a st.session_state directamente.
            st.session_state.combined_df_base_complet3 = resultados.get("BaseCompleta", pd.DataFrame())
            st.session_state.summary_df = resultados.get("Summary", pd.DataFrame())
            st.session_state.df_porcentajes_niveles = resultados.get("Porcentajes Niveles", pd.DataFrame())
            st.session_state.df_res_dimTE3 = resultados.get("Res Dim TE3", pd.DataFrame())
            st.session_state.df_resumen = resultados.get("Resumen", pd.DataFrame())
            st.session_state.df_resultados_porcentaje = resultados.get("Resultado", pd.DataFrame())
            st.session_state.top_glosas = resultados.get("Top Glosas", pd.DataFrame())
            st.session_state.df_resultados = resultados.get("df_resultados", pd.DataFrame())
            st.session_state.df_res_com = resultados.get("df_res_com", pd.DataFrame())
            st.session_state.df_ciiu = resultados.get("df_ciiu", pd.DataFrame())
            st.session_state.df_recomendaciones = resultados.get("df_recomendaciones", pd.DataFrame())

            # OPCIONAL: Guardar el CUV seleccionado en session_state, si se necesita
            st.session_state.selected_cuv = cuv_valor

            # A partir de aquí, todos los datos ya están filtrados por el CUV y disponibles en st.session_state.
            # No volver a filtrar por CUV en el resto del código.

            df_ciiu = st.session_state.df_ciiu
            codigo_ciiu = extraer_codigo_ciiu(st.session_state.df_res_com)

            if codigo_ciiu is not None:
                st.write(f"**Código CIIU Extraído:** {codigo_ciiu}")
            else:
                st.error(f"No se pudo determinar el valor de CIIU para el CUV {cuv_valor}.")

            # Continuar con el procesamiento solo si df_res_com no está vacío
            if not st.session_state.df_res_com.empty:
                # Asegurar que ciertas columnas sean de tipo string
                if 'CUV' in st.session_state.df_res_com.columns:
                    st.session_state.df_res_com['CUV'] = st.session_state.df_res_com['CUV'].astype(str)

                # Mostrar las fechas antes de la conversión
                st.subheader("Fechas Antes de la Conversión")
                st.write("Fecha_Inicio:", st.session_state.df_res_com['Fecha_Inicio'].head())
                st.write("Fecha_Fin:", st.session_state.df_res_com['Fecha_Fin'].head())

                # Mostrar algunos ejemplos de las fechas
                st.write("Ejemplos de 'Fecha_Inicio':")
                st.write(st.session_state.df_res_com['Fecha_Inicio'].head())

                st.write("Ejemplos de 'Fecha_Fin':")
                st.write(st.session_state.df_res_com['Fecha_Fin'].head())


                # Procesar columnas de fecha
                columnas_fecha = ['Fecha_Inicio', 'Fecha_Fin']
                try:
                    for columna in columnas_fecha:
                        st.session_state.df_res_com = procesar_columna_fecha(st.session_state.df_res_com, columna)
                except ValueError as e:
                    st.error(e)
            else:
                st.info("No se encontraron registros en 'Filas Resultados' para el CUV proporcionado.")

            # Mostrar las fechas después de la conversión
            st.subheader("Fechas Después de la Conversión")
            st.write("Fecha_Inicio:", st.session_state.df_res_com['Fecha_Inicio'].head())
            st.write("Fecha_Fin:", st.session_state.df_res_com['Fecha_Fin'].head())


            # Visualización básica de los resultados
            st.subheader("Filas Resultados (CUV y Folio)")
            if not st.session_state.df_resultados.empty:
                st.dataframe(st.session_state.df_resultados)
            else:
                st.info("No se encontraron registros en 'Filas Resultados' para el CUV proporcionado.")

            st.subheader("Filas Resultados (Todas las Columnas)")
            if not st.session_state.df_res_com.empty:
                st.dataframe(st.session_state.df_res_com)
            else:
                st.info("No se encontraron registros en 'Filas Resultados (Todas las Columnas)' para el CUV proporcionado.")

            # Visualizar las otras tablas consultadas
            otras_tablas = {
                "BaseCompleta": st.session_state.combined_df_base_complet3,
                "Summary": st.session_state.summary_df,
                "Porcentajes Niveles": st.session_state.df_porcentajes_niveles,
                "Res Dim TE3": st.session_state.df_res_dimTE3,
                "Resumen": st.session_state.df_resumen,
                "Resultado": st.session_state.df_resultados_porcentaje,
                "Top Glosas": st.session_state.top_glosas,
                "CIIU": st.session_state.df_ciiu,
                "Recomendaciones": st.session_state.df_recomendaciones
            }

            for nombre, df in otras_tablas.items():
                st.subheader(nombre)
                if not df.empty:
                    st.dataframe(df)
                else:
                    st.info(f"No se encontraron registros en '{nombre}' para el CUV proporcionado.")


    combined_df_base_complet3 = st.session_state.combined_df_base_complet3
    summary_df = st.session_state.summary_df
    df_porcentajes_niveles = st.session_state.df_porcentajes_niveles
    df_res_dimTE3 = st.session_state.df_res_dimTE3
    df_resumen = st.session_state.df_resumen
    df_resultados_porcentaje = st.session_state.df_resultados_porcentaje
    top_glosas = st.session_state.top_glosas
    df_ciiu = st.session_state.df_ciiu
    df_recomendaciones = st.session_state.df_recomendaciones
    df_res_com = st.session_state.df_res_com
    df_resultados = st.session_state.df_resultados

    #for col in df.select_dtypes(include='object').columns:
    #    if df[col].str.isnumeric().all():
    #        df[col] = pd.to_numeric(df[col], errors='coerce')
    #    else:
    #        df[col] = df[col].astype(str).fillna('')


    ############### Parte 3: Procesamiento posterior a la búsqueda, Funciones de Formateo y Auxiliares ###############

    columnas_fecha = ['Fecha_Inicio', 'Fecha_Fin']

    # Asegúrate que df_res_com esté disponible en session_state antes de este bloque
    if "df_res_com" in st.session_state and not st.session_state.df_res_com.empty:
        df_res_com = st.session_state.df_res_com.copy()
        for columna in columnas_fecha:
            df_res_com = procesar_columna_fecha(df_res_com, columna)
        # Si quieres actualizar el session_state con el df procesado
        st.session_state.df_res_com = df_res_com
    else:
        st.warning("No hay datos en df_res_com para procesar.")



    # Definir funciones auxiliares
    def normalizar_texto(texto):
        if isinstance(texto, str):
            texto = texto.strip().lower()
            texto = ''.join(
                c for c in unicodedata.normalize('NFD', texto)
                if unicodedata.category(c) != 'Mn'
            )
            return texto
        else:
            return ''


    def obtener_dimm_por_dimension(nombre_dimension):
        nombre_dimension_normalizado = normalizar_texto(nombre_dimension)
        # df_dimensiones debería estar definido o cargado desde alguna parte
        # Aquí debes asegurarte de tener el DataFrame 'dimensiones' disponible
        # Por ejemplo, podrías cargarlo desde un archivo o definirlo manualmente
        # Supongamos que está en 'df_res_dimTE3'
        # Ajusta esto según tu estructura de datos
        df_dimensiones = df_res_dimTE3[['Dimensión', 'dimm']].drop_duplicates()
        df_dimensiones['dimension_normalizada'] = df_dimensiones['Dimensión'].apply(normalizar_texto)
        resultado = df_dimensiones[df_dimensiones['dimension_normalizada'] == nombre_dimension_normalizado]
        if not resultado.empty:
            return resultado.iloc[0]['dimm']
        else:
            st.warning(f"No se encontró el código 'dimm' para la dimensión '{nombre_dimension}'.")
            return None


    def agregar_tabla_ges_por_dimension(doc, df, cuv, df_recomendaciones, df_resultados_porcentaje,
                                        df_porcentajes_niveles, top_glosas, datos):
        """
        Agrega una tabla de dimensiones y GES para un CUV específico en el documento de Word.

        Parámetros:
        doc (Document): El objeto del documento de Word.
        df (pd.DataFrame): DataFrame con los datos de dimensiones y GES filtrados.
        cuv (str): El CUV específico para el que se generará la tabla.
        df_recomendaciones (pd.DataFrame): DataFrame con las recomendaciones por dimensión.
        """

        # Filtrar el DataFrame para el CUV específico
        df_revision = df[df['CUV'] == cuv]
        unique_te3 = df_revision['TE3'].dropna().unique()

        if len(unique_te3) < 2:
            noges = 1
            st.warning(f"Solo hay un GES para el CUV {cuv}. No se generarán recomendaciones por GES.")
        else:
            noges = 0

        # Filtrar el DataFrame para el CUV específico y puntajes 1 y 2
        df_cuv = df[(df['CUV'] == cuv) & (df['Puntaje'].isin([1, 2]))]

        if df_cuv.empty:
            st.warning(f"No hay datos con puntaje 1 o 2 para el CUV {cuv}.")
            return

        # Agrupar por 'Dimensión' y combinar los valores únicos de 'TE3' en una lista separada por "; "
        resultado = df_cuv.groupby('Dimensión')['TE3'].unique().reset_index()
        resultado['GES'] = resultado['TE3'].apply(lambda x: '; '.join(map(str, x)))

        # Limpiar el campo 'GES' reemplazando ciertos caracteres
        resultado['GES'] = resultado['GES'].str.replace('|', '_', regex=False) \
            .str.replace(':', '_', regex=False) \
            .str.replace('?', '_', regex=False)

        # Asegúrate de que los valores en 'CIIU' son strings y extraer la parte necesaria
        datos['CIIU'] = datos['CIIU'].apply(lambda x: x.split('_')[-1] if isinstance(x, str) else x)

        # Filtrar el DataFrame para el CUV específico y obtener el valor único de CIIU
        ciiu_valor = datos.loc[datos['CUV'] == cuv, 'CIIU'].copy()

        if len(ciiu_valor) > 0:
            ciiu_unico = ciiu_valor.iloc[0]
            if isinstance(ciiu_unico, str) and ciiu_unico.isdigit():
                ciiu = int(ciiu_unico[:2]) if len(ciiu_unico) > 5 else int(ciiu_unico[:1])

            else:
                print("El valor de CIIU no es numérico.")
                ciiu = None
        else:
            print("CUV no encontrado en la tabla de datos.")
            ciiu = None

        if ciiu is None:
            st.error(f"No se pudo determinar el valor de CIIU para el CUV {cuv}.")
            return

        # Crear la tabla en el documento
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        column_widths = [Inches(0.5), Inches(0.5), Inches(0.5), Inches(7), Inches(0.5), Inches(0.5)]

        # Configurar el ancho de cada columna
        for col_idx, width in enumerate(column_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = width

        # Agregar encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Dimensión en riesgo'
        hdr_cells[1].text = 'Preguntas clave'
        hdr_cells[2].text = 'Explicación'
        hdr_cells[3].text = 'Medidas propuestas'
        hdr_cells[4].text = 'Fecha monitoreo'
        hdr_cells[5].text = 'Responsable seguimiento'

        # Asegurarse de que 'Descripción' sea string y reemplazar NaN
        df_resultados_porcentaje['Descripción'] = df_resultados_porcentaje['Descripción'].fillna('').astype(str)

        # Rellenar la tabla con los datos de 'Dimensión' y 'GES'
        for _, row in resultado.iterrows():
            dim = row['Dimensión']
            ges = row['GES']

            # Obtener las recomendaciones para esta dimensión
            recomendaciones = df_recomendaciones[
                (df_recomendaciones['Dimensión'] == dim) &
                (df_recomendaciones['ciiu'] == str(ciiu))
                ]['Recomendación'].tolist()
            medidas_propuestas = '\n'.join([f"- {rec}" for rec in recomendaciones]) if recomendaciones else 'N/A'

            # Obtener la descripción relacionada desde df_resultados_porcentaje
            descripcion = df_resultados_porcentaje[
                (df_resultados_porcentaje['Dimensión'] == dim) &
                (df_resultados_porcentaje['CUV'] == cuv)
                ]['Descripción'].values

            # Filtrar solo cadenas no vacías
            descripcion = [desc for desc in descripcion if isinstance(desc, str) and desc.strip() != '']

            descripcion2 = [
                f"{desc} en {ges}"
                for desc in df_porcentajes_niveles[
                    (df_porcentajes_niveles['Dimensión'] == dim) &
                    (df_porcentajes_niveles['CUV'] == cuv) &
                    (df_porcentajes_niveles['TE3'] == ges) &
                    (df_porcentajes_niveles['Descripción'].str.strip() != '')
                    ]['Descripción'].values
            ]

            descripcion2_text = '\n'.join(descripcion2).replace("[", "").replace("]", "").replace("'",
                                                                                                  "") if descripcion2 else ""

            # Construir descripcion_text
            descripcion_text = ""
            if len(descripcion) > 0 and isinstance(descripcion[0], str) and len(descripcion[0]) > 0:
                descripcion_text = descripcion[0] + " para todo el centro de trabajo\n"
            elif len(descripcion) > 1 and isinstance(descripcion[1], str) and len(descripcion[1]) > 0:
                descripcion_text = descripcion[1] + " para todo el centro de trabajo\n"
            elif len(descripcion) > 2 and isinstance(descripcion[2], str) and len(descripcion[2]) > 0:
                descripcion_text = descripcion[2] + " para todo el centro de trabajo\n"
            else:
                descripcion_text = ""

            # Verificar si hay múltiples GES
            if noges == 1:
                descripcion2_text = ""
                print(f"Solo hay un GES para el CUV {cuv}. No se generarán recomendaciones por GES.")

            # Obtener las preguntas clave desde top_glosas
            filtro_glosas = top_glosas[(top_glosas['Dimensión'] == dim) & (top_glosas['CUV'] == cuv)]
            preguntas = filtro_glosas['Pregunta'].tolist()
            preguntas_text = '\n'.join(preguntas) if preguntas else 'N/A'

            # Rellenar las celdas de la tabla
            row_cells = table.add_row().cells
            row_cells[0].text = f"{descripcion_text}{descripcion2_text}".strip()
            row_cells[1].text = preguntas_text.strip()
            row_cells[2].text = ''  # Espacio para 'Explicación'
            row_cells[3].text = medidas_propuestas.strip()
            row_cells[4].text = ''  # Fecha de monitoreo
            row_cells[5].text = ''  # Responsable seguimiento

        # Ajustar el tamaño de fuente de las celdas (opcional)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Ajusta el tamaño de la fuente


    def convertir_figura_a_imagen(fig, formato='png', dpi=150):
        """
        Convierte una figura de Matplotlib a una cadena base64.

        Args:
            fig (matplotlib.figure.Figure): La figura de Matplotlib a convertir.
            formato (str): El formato de la imagen (por defecto 'png').
            dpi (int): Resolución de la imagen (puntos por pulgada).

        Returns:
            str: La imagen codificada en base64.
        """
        buf = BytesIO()
        fig.savefig(buf, format=formato, bbox_inches='tight', dpi=dpi)
        buf.seek(0)
        img_base64 = base64.b64encode(buf.read()).decode()
        return img_base64


    def agregar_tabla_ges_por_dimension_streamlit(df, cuv, df_recomendaciones, df_porcentajes_niveles, top_glosas, df_res_com):
        """
        Prepara una estructura de datos con medidas propuestas por dimensión y retorna una lista de diccionarios.
        """

        # Verificar tipos de entrada
        dataframes = {
            "df": df,
            "df_recomendaciones": df_recomendaciones,
            "df_porcentajes_niveles": df_porcentajes_niveles,
            "top_glosas": top_glosas,
            "df_res_com": df_res_com
        }

        for name, df_input in dataframes.items():
            if not isinstance(df_input, pd.DataFrame):
                st.error(f"El parámetro '{name}' debe ser un DataFrame.")
                return []

        # Convertir 'Puntaje' a numérico
        try:
            df['Puntaje'] = pd.to_numeric(df['Puntaje'], errors='coerce')
            df_porcentajes_niveles['Puntaje'] = pd.to_numeric(df_porcentajes_niveles['Puntaje'], errors='coerce')
        except Exception as e:
            st.error(f"Error al convertir 'Puntaje' a numérico: {e}")
            return []

        # Filtrar el DataFrame para el CUV y puntajes 1 y 2
        df_cuv = df[(df['CUV'] == cuv) & (df['Puntaje'].isin([1, 2]))]
        if df_cuv.empty:
            st.warning(f"No hay datos con puntaje 1 o 2 para el CUV {cuv}.")
            return []

        # Agrupar por 'Dimensión' y obtener los valores únicos de 'TE3'
        resultado = df_cuv.groupby('Dimensión')['TE3'].unique().reset_index()
        resultado['GES'] = resultado['TE3'].apply(lambda x: '; '.join(map(str, x)))

        # Limpiar el campo 'GES' reemplazando ciertos caracteres
        resultado['GES'] = resultado['GES'].str.replace(r'[|:?]', '_', regex=True)

        # Extraer el valor único de CIIU para el CUV
        try:
            df_res_com['CIIU'] = df_res_com['CIIU'].apply(lambda x: x.split('_')[-1] if isinstance(x, str) else x)
            ciiu_valor = df_res_com.loc[df_res_com['CUV'] == cuv, 'CIIU']

            if not ciiu_valor.empty:
                ciiu_unico = ciiu_valor.iloc[0]
                if isinstance(ciiu_unico, str) and ciiu_unico.isdigit():
                    ciiu = int(ciiu_unico[:2]) if len(ciiu_unico) > 5 else int(ciiu_unico[:1])
                else:
                    st.error("El valor de CIIU no es numérico.")
                    return []
            else:
                st.error(f"No se encontró el valor de CIIU para el CUV {cuv}.")
                return []
        except Exception as e:
            st.error(f"Error al procesar el valor de CIIU: {e}")
            return []

        # Crear descripciones en df_porcentajes_niveles
        try:
            df_porcentajes_niveles['Descripción'] = df_porcentajes_niveles.apply(
                lambda row: f"{row['Porcentaje']}% Riesgo {row['Nivel']}, {row['Respuestas']} personas"
                if row['Puntaje'] in [1, 2] else "",
                axis=1
            )
            df_porcentajes_niveles['Descripción'] = df_porcentajes_niveles['Descripción'].fillna('').astype(str)
        except Exception as e:
            st.error(f"Error al generar las descripciones en df_porcentajes_niveles: {e}")
            return []

        # Inicializar lista para almacenar resultados
        dimensiones = []

        for _, row in resultado.iterrows():
            dim = row['Dimensión']
            ges = row['GES']

            # Obtener recomendaciones para la dimensión y el CIIU
            recomendaciones = df_recomendaciones[
                (df_recomendaciones['Dimensión'] == dim) &
                (df_recomendaciones['ciiu'] == str(ciiu))
            ]['Recomendación'].tolist()
            medidas_propuestas = recomendaciones if recomendaciones else ['N/A']

            # Generar descripciones de riesgo
            descripcion = [
                f"{desc} en {ges}"
                for desc in df_porcentajes_niveles[
                    (df_porcentajes_niveles['Dimensión'] == dim) &
                    (df_porcentajes_niveles['CUV'] == cuv) &
                    (df_porcentajes_niveles['TE3'] == ges) &
                    (df_porcentajes_niveles['Descripción'].str.strip() != '')
                ]['Descripción'].values
            ]
            descripcion_text = '\n'.join(descripcion).strip() if descripcion else 'N/A'

            # Obtener preguntas clave desde top_glosas
            filtro_glosas = top_glosas[(top_glosas['Dimensión'] == dim) & (top_glosas['CUV'] == cuv)]
            preguntas = filtro_glosas['Pregunta'].tolist()
            preguntas_text = '\n'.join([f"- {pregunta}" for pregunta in preguntas]) if preguntas else 'N/A'

            # Agregar a la lista de resultados
            dimensiones.append({
                'GES': ges,
                'Dimensión en riesgo': dim,
                'Descripción riesgo': descripcion_text,
                'Preguntas clave': preguntas_text,
                'Interpretación del grupo de discusión': "",
                'Medidas propuestas': medidas_propuestas
            })

        return dimensiones


    # Función auxiliar para manejar NaN
    def safe_get(d, key):
        value = d.get(key, 'N/A')
        if pd.isna(value):
            return 'N/A'
        return value



    def mostrar_datos(datos):
        """
        Muestra los datos de una empresa formateados en Markdown.

        Parameters:
        - datos (dict): Diccionario con los datos de la empresa.
        - estado (dict): Diccionario con el estado de riesgo.
        """
        # Manejo seguro para 'CIIU_CT'
        ciiu = safe_get(datos, 'CIIU_CT')
        if ciiu != 'N/A' and isinstance(ciiu, str):
            ciiu = ciiu.split('_')[-1]
        else:
            ciiu = 'N/A'

        # Manejo de fechas
        fecha_inicio = safe_get(datos, 'Fecha_Inicio')
        fecha_fin = safe_get(datos, 'Fecha_Fin')

        # Convertir fechas a formato string si no son 'N/A'
        if fecha_inicio != 'N/A':
            fecha_inicio = pd.to_datetime(fecha_inicio).strftime('%d-%m-%Y') if not pd.isna(fecha_inicio) else 'N/A'
        if fecha_fin != 'N/A':
            fecha_fin = pd.to_datetime(fecha_fin).strftime('%d-%m-%Y') if not pd.isna(fecha_fin) else 'N/A'

        contenido = f"""
        **Razón Social:** {safe_get(datos, 'Nombre_Empresa')}  
        **RUT:** {safe_get(datos, 'RUT_Empresa')}  
        **Nombre del centro de trabajo:** {safe_get(datos, 'Nombre_Centro_de_Trabajo')}  
        **CUV:** {safe_get(datos, 'CUV')}  
        **CIIU:** {ciiu}  
        **Fecha de activación del cuestionario:** {fecha_inicio}  
        **Fecha de cierre del cuestionario:** {fecha_fin}  
        **Universo de trabajadores de evaluación:** {safe_get(datos, 'Nº_Trabajadores_CT')}  
        **Nivel de riesgo:** {safe_get(datos, 'Nivel_de_riesgo')}
        """
        st.markdown(contenido)


    # Filtrar los datos para el CUV seleccionado
    datos = df_res_com

    # Convertir DataFrames a diccionarios
    datos_dict = datos.iloc[0].to_dict() if not datos.empty else {}

    # Mostrar los datos
    st.subheader("Información de la Empresa")

    st.write("Contenido de 'datos':", datos_dict)

    # Llamar a la función con diccionarios si 'datos_dict' no está vacío
    if datos_dict:
        mostrar_datos(datos_dict)
    else:
        st.info("No hay datos disponibles para mostrar.")


    # Manejar casos donde los datos están vacíos
    if datos.empty:
        if 'selected_cuv' in st.session_state and st.session_state.selected_cuv:
            st.error(f"No se encontraron datos para el CUV {st.session_state.selected_cuv}.")
        else:
            st.error("No se encontraron datos. Aún no se ha seleccionado ningún CUV.")
        st.stop()  # Detenemos la ejecución para no intentar datos.iloc[0] más adelante
    else:
        # Asegurarse de que 'selected_cuv' está definido
        if 'selected_cuv' not in st.session_state:
            st.session_state['selected_cuv'] = datos_dict.get('CUV', 'N/A')

        datos_first = datos.iloc[0]




    def generar_grafico_principal(df):
        if df.empty:
            st.warning("No se encontraron datos para el CUV seleccionado.")
            return None

        try:
            # Convertir 'Porcentaje' a numérico y manejar errores
            if 'Porcentaje' in df.columns:
                df['Porcentaje'] = pd.to_numeric(df['Porcentaje'], errors='coerce')

            # Asegurarse de que 'Nivel' sea string
            if 'Nivel' in df.columns:
                df['Nivel'] = df['Nivel'].astype(str)

            # Pivotear el DataFrame
            df_pivot = df.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[::-1]
            df_pivot.index = df_pivot.index.astype(str)  # Asegurar que los índices sean strings
            df_pivot.index = df_pivot.index.map(lambda x: f"dim_{x}")
        except Exception as e:
            st.error(f"Error al pivotear los datos: {e}")
            return None

        # Crear la figura
        fig, ax = plt.subplots(figsize=(12, 8))

        # Configuración de los niveles y colores
        niveles = ["Bajo", "Medio", "Alto"]
        colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
        posiciones = np.arange(len(df_pivot.index))
        ancho_barra = 0.2

        for i, nivel in enumerate(niveles):
            if nivel in df_pivot.columns:
                valores = df_pivot[nivel]
                ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                        label=f"Riesgo {nivel.lower()} (%)", color=colores.get(nivel, 'grey'))
            else:
                st.warning(f"Nivel '{nivel}' no encontrado en las columnas de pivot.")

        # Línea de referencia en 50%
        ax.axvline(50, color="blue", linestyle="--", linewidth=1)
        ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {st.session_state.selected_cuv}", pad=50)
        ax.set_xlabel("Porcentaje")
        ax.set_ylabel("Dimensiones")
        ax.set_xlim(0, 100)
        ax.set_yticks(posiciones + ancho_barra)
        ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

        # Leyenda y ajustes de diseño
        fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.6, 0.96), ncol=3)
        plt.subplots_adjust(left=0.3, top=0.85)
        plt.tight_layout()

        return fig


    # Sección 3: Generación y visualización de gráficos
    st.header("Gráfico general del centro de trabajo")

    st.write("Tipos de columnas en df_resultados_porcentaje:", st.session_state.df_resultados_porcentaje.dtypes)
    st.write("Contenido de df_resultados_porcentaje:", st.session_state.df_resultados_porcentaje.head(20))



    if st.session_state.df_resultados_porcentaje.empty:
        st.warning("No hay datos en df_resultados_porcentaje para generar el gráfico principal.")
    else:
        fig_principal = generar_grafico_principal(st.session_state.df_resultados_porcentaje)



    if fig_principal:
        # Convertir la figura a una imagen base64
        img_base64 = convertir_figura_a_imagen(fig_principal, formato='png', dpi=150)
        # Construir la cadena de la imagen en formato data URI
        img_uri = f"data:image/png;base64,{img_base64}"
        # Mostrar la imagen con un ancho fijo (por ejemplo, 800 píxeles)
        st.markdown(
            f"<div style='display: flex; justify-content: center;'><img src='data:image/png;base64,{img_base64}' width='800'></div>",
            unsafe_allow_html=True
        )
    else:
        st.warning("fig.principal No se pudo generar el gráfico principal.")


    # ---- NUEVA SECCIÓN PARA MOSTRAR DIMENSIONES EN RIESGO ----
    #st.header("4. Dimensiones en Riesgo")

    # Obtener dimensiones en riesgo

    if 'Puntaje' in st.session_state.df_resultados_porcentaje.columns:
        st.session_state.df_resultados_porcentaje['Puntaje'] = pd.to_numeric(st.session_state.df_resultados_porcentaje['Puntaje'], errors='coerce')

    dimensiones_riesgo_alto = st.session_state.df_resultados_porcentaje[
        st.session_state.df_resultados_porcentaje['Puntaje'] == 2
        ]['Dimensión'].tolist()

    dimensiones_riesgo_medio = st.session_state.df_resultados_porcentaje[
        st.session_state.df_resultados_porcentaje['Puntaje'] == 1
        ]['Dimensión'].tolist()

    dimensiones_riesgo_bajo = st.session_state.df_resultados_porcentaje[
        st.session_state.df_resultados_porcentaje['Puntaje'] == -2
        ]['Dimensión'].tolist()

    # Mostrar dimensiones en riesgo en la web
    st.subheader("Dimensiones en riesgo")
    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("**Alto:**")
        if dimensiones_riesgo_alto:
            st.write(", ".join(dimensiones_riesgo_alto))
        else:
            st.write("Ninguna")

    with col2:
        st.markdown("**Medio:**")
        if dimensiones_riesgo_medio:
            st.write(", ".join(dimensiones_riesgo_medio))
        else:
            st.write("Ninguna")

    with col3:
        st.markdown("**Bajo:**")
        if dimensiones_riesgo_bajo:
            st.write(", ".join(dimensiones_riesgo_bajo))
        else:
            st.write("Ninguna")


    def generar_graficos_por_te3(df):
        if df.empty:
            st.warning("No se encontraron datos para el CUV seleccionado.")
            return []

        te3_values = df['TE3'].unique()

        if len(te3_values) <= 1:
            st.info("Solo hay un GES para este CUV. No se generarán gráficos adicionales.")
            return []

        figs_te3 = []

        for te3 in te3_values:
            # Filtrar datos para el valor específico de TE3
            df_te3 = df[df['TE3'] == te3]
            if df_te3.empty:
                continue

            try:
                # Convertir columnas a los tipos adecuados
                if 'Porcentaje' in df_te3.columns:
                    df_te3.loc[:, 'Porcentaje'] = pd.to_numeric(df_te3['Porcentaje'], errors='coerce')

                if 'Puntaje' in df_te3.columns:
                    df_te3.loc[:, 'Puntaje'] = pd.to_numeric(df_te3['Puntaje'], errors='coerce')

                if 'Nivel' in df_te3.columns:
                    df_te3.loc[:, 'Nivel'] = df_te3['Nivel'].astype(str)

                # Pivotear el DataFrame
                df_pivot = df_te3.pivot(index="Dimensión", columns="Nivel", values="Porcentaje")
                df_pivot = df_pivot.infer_objects(copy=False)  # Asegurar que no hay objetos ambiguos
                df_pivot = df_pivot.fillna(0).iloc[::-1]  # Llenar valores NaN
            except Exception as e:
                st.error(f"Error al pivotear los datos para TE3 {te3}: {e}")
                continue

            # Crear la figura
            fig, ax = plt.subplots(figsize=(12, 8))

            # Configuración de los niveles y colores
            niveles = ["Bajo", "Medio", "Alto"]
            colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
            posiciones = np.arange(len(df_pivot.index))
            ancho_barra = 0.2

            for i, nivel in enumerate(niveles):
                if nivel in df_pivot.columns:
                    valores = df_pivot[nivel]
                    ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                            label=f"Riesgo {nivel.lower()} (%)", color=colores.get(nivel, 'grey'))
                else:
                    st.warning(f"Nivel '{nivel}' no encontrado en las columnas de pivot para TE3 {te3}.")

            # Línea de referencia en 50%
            ax.axvline(50, color="blue", linestyle="--", linewidth=1)
            ax.set_title(
                f"Porcentaje de trabajadores por nivel de riesgo - CUV {st.session_state.selected_cuv}, TE3 {te3}",
                pad=50)
            ax.set_xlabel("Porcentaje")
            ax.set_ylabel("Dimensiones")
            ax.set_xlim(0, 100)
            ax.set_yticks(posiciones + ancho_barra)
            ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

            # Leyenda y ajustes de diseño
            fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.5, 0.96), ncol=3)
            plt.tight_layout()

            figs_te3.append((fig, te3))

        return figs_te3


    # Sección 4: Generación de gráficos por TE3
    st.header("Generación de gráficos por GES")

    # Generar y mostrar los gráficos por TE3
    figs_te3 = generar_graficos_por_te3(st.session_state.df_porcentajes_niveles)

    if figs_te3:
        for fig_te3, te3 in figs_te3:
            st.subheader(f"Gráfico para GES: {te3}")
            # Convertir la figura a una imagen base64
            img_base64_te3 = convertir_figura_a_imagen(fig_te3, formato='png')
            # Construir la cadena de la imagen en formato data URI
            img_uri_te3 = f"data:image/png;base64,{img_base64_te3}"
            # Mostrar la imagen con un ancho fijo (por ejemplo, 800 píxeles)
            #st.image(img_uri_te3, width=800, caption=f"Gráfico para GES: {te3}")

            st.markdown(
                f"<div style='display: flex; justify-content: center;'><img src='data:image/png;base64,{img_base64_te3}' width='800'></div>",
                unsafe_allow_html=True
            )

            if 'Puntaje' in df_porcentajes_niveles.columns:
                df_porcentajes_niveles['Puntaje'] = pd.to_numeric(
                    df_porcentajes_niveles['Puntaje'], errors='coerce')

            dimensiones_riesgo_alto = df_porcentajes_niveles[
                (df_porcentajes_niveles['TE3'] == te3) &
                (df_porcentajes_niveles['Puntaje'] == 2)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_medio = df_porcentajes_niveles[
                (df_porcentajes_niveles['TE3'] == te3) &
                (df_porcentajes_niveles['Puntaje'] == 1)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_bajo = df_porcentajes_niveles[
                (df_porcentajes_niveles['TE3'] == te3) &
                (df_porcentajes_niveles['Puntaje'] == -2)
                ]['Dimensión'].tolist()

            # Crear tres columnas
            col1, col2, col3 = st.columns(3)

            with col1:
                st.markdown("**Alto:**")
                if dimensiones_riesgo_alto:
                    st.write(", ".join(dimensiones_riesgo_alto))
                else:
                    st.write("Ninguna")

            with col2:
                st.markdown("**Medio:**")
                if dimensiones_riesgo_medio:
                    st.write(", ".join(dimensiones_riesgo_medio))
                else:
                    st.write("Ninguna")

            with col3:
                st.markdown("**Bajo:**")
                if dimensiones_riesgo_bajo:
                    st.write(", ".join(dimensiones_riesgo_bajo))
                else:
                    st.write("Ninguna")


    else:
        st.info("No se generaron gráficos adicionales por GES.")


    # Sección 5: Prescripciones de medidas
    st.header("Prescripciones de medidas")

    dimensiones_te3 = agregar_tabla_ges_por_dimension_streamlit(df_res_dimTE3, st.session_state.selected_cuv,
                                                                df_recomendaciones,
                                                                df_porcentajes_niveles, top_glosas, df_res_com)


    #st.session_state.dimensiones_te3 = dimensiones_te3

    # Función para actualizar la numeración
    def actualizar_numeracion(df):
        df = df.reset_index(drop=True)
        df['N°'] = df.index + 1
        return df


    # Agrupar las dimensiones por GES correctamente desglosando valores combinados
    ges_groups = {}
    for dimension in dimensiones_te3:
        ges_values = dimension['GES'].split(";")  # Dividir los valores de GES
        for ges in ges_values:
            ges = ges.strip()  # Eliminar espacios adicionales
            if ges not in ges_groups:
                ges_groups[ges] = []
            # Crear una copia de la dimensión para evitar duplicados en diferentes GES
            dimension_copy = dimension.copy()
            dimension_copy['GES'] = ges  # Asignar el GES desglosado
            ges_groups[ges].append(dimension_copy)


    # Definir la función para formatear opciones del selectbox
    def format_option(option):
        """
        Formatea las opciones del selectbox con un nombre y una fecha.
        option: tuple -> (nombre, fecha)
        """
        nombre, fecha = option

        # Asegurarse de que fecha no sea una Serie
        if isinstance(fecha, pd.Series):
            fecha = fecha.iloc[0]  # Tomar el primer valor si es una Serie

        # Validar que la fecha no sea nula
        if pd.notnull(fecha):
            return f"{nombre} - {fecha.strftime('%d-%m-%Y')}"
        else:
            return nombre

    # Crear un diccionario temporal para almacenar interpretaciones en st.session_state si no existe
    if 'interpretaciones_temporales' not in st.session_state:
        st.session_state['interpretaciones_temporales'] = {}

    # Procesar cada GES
    for ges, dimensiones in ges_groups.items():
        st.header(f"GES: {ges}")

        # Procesar cada Dimensión dentro del GES
        for idx, dimension in enumerate(dimensiones, 1):
            st.subheader(f"Dimensión: {dimension['Dimensión en riesgo']}")
            st.write(f"**Descripción del riesgo:** {dimension['Descripción riesgo']}")
            st.write("**Preguntas clave:**")
            st.write(dimension["Preguntas clave"])

            # Gestionar la interpretación del grupo de discusión
            interpretacion_key = f"interpretacion_{ges}_{idx}"

            # Establecer un valor inicial en el diccionario temporal si no existe
            if interpretacion_key not in st.session_state['interpretaciones_temporales']:
                st.session_state['interpretaciones_temporales'][interpretacion_key] = st.session_state.get(
                    interpretacion_key, "")

            # Mostrar el cuadro de texto para la interpretación
            interpretacion = st.text_area(
                label="Interpretación del grupo de discusión",
                value=st.session_state['interpretaciones_temporales'][interpretacion_key],
                height=150,
                key=interpretacion_key
            )

            # Actualizar el valor en el diccionario temporal sin modificar el valor en el estado directamente asociado al widget
            st.session_state['interpretaciones_temporales'][interpretacion_key] = interpretacion

            # Gestionar medidas propuestas
            st.write("#### Medidas propuestas")
            session_key = f"measures_{ges}_{idx}"
            if session_key not in st.session_state:
                st.session_state[session_key] = pd.DataFrame([
                    {'N°': i + 1, 'GES': ges, 'Dimensión': dimension['Dimensión en riesgo'],
                     'Medida': medida, 'Fecha monitoreo': '', 'Responsable': '',
                     'Activo': True, 'Seleccionada': False}
                    for i, medida in enumerate(dimension['Medidas propuestas'])
                ])
            else:
                st.session_state[session_key] = actualizar_numeracion(st.session_state[session_key])

            df = st.session_state[session_key]
            medidas_list = [""] + df.loc[df['Activo'], 'Medida'].tolist()  # Añadir opción vacía al inicio
            selected_measure = st.selectbox(
                "Seleccione una medida para editar o deje vacío para crear una nueva",
                medidas_list,
                key=f"select_{ges}_{idx}"
            )

            if selected_measure:  # Si selecciona una medida existente
                medida_idx = df[df['Medida'] == selected_measure].index[0]
                st.write("#### Editar medida seleccionada")
            else:  # Si no selecciona nada, permite crear una nueva medida
                st.write("#### Crear una nueva medida")
                medida_idx = None

            # Calcular las fechas de corto, mediano y largo plazo
            fecha_fin = pd.to_datetime("2024-11-26")  # Reemplaza con tu lógica para obtener 'Fecha Fin'
            fecha_corto_plazo = fecha_fin + timedelta(days=240)
            fecha_mediano_plazo = fecha_fin + timedelta(days=330)
            fecha_largo_plazo = fecha_fin + timedelta(days=420)

            # Crear una lista con las fechas calculadas
            fechas_opciones = [
                ("Corto Plazo (180 días)", fecha_corto_plazo),
                ("Mediano Plazo (270 días)", fecha_mediano_plazo),
                ("Largo Plazo (360 días)", fecha_largo_plazo),
                ("Otra fecha", None)  # Opción para seleccionar una fecha personalizada
            ]

            # Crear formulario para editar o crear medida
            with st.form(key=f"form_{ges}_{idx}"):
                medida = st.text_area(
                    "Descripción de la medida",
                    value=df.at[medida_idx, 'Medida'] if medida_idx is not None else "",
                    key=f"edit_medida_{ges}_{idx}",
                    height=90
                )

                # Mostrar la lista desplegable para seleccionar la fecha de monitoreo
                opcion_seleccionada = st.selectbox(
                    "Selecciona la Fecha de Monitoreo",
                    options=fechas_opciones,
                    format_func=format_option,
                    key=f"select_fecha_{ges}_{idx}"
                )

                # Determinar la fecha seleccionada
                if opcion_seleccionada[1]:
                    # Si se selecciona una de las opciones predefinidas
                    fecha = opcion_seleccionada[1]
                else:
                    # Si se selecciona 'Otra fecha', mostrar el date_input para elegir manualmente
                    fecha_default = pd.to_datetime(df.at[medida_idx, 'Fecha monitoreo']) if (
                            medida_idx is not None and pd.notna(df.at[medida_idx, 'Fecha monitoreo'])
                    ) else datetime.today()
                    fecha = st.date_input(
                        "Selecciona una Fecha de Monitoreo personalizada",
                        value=fecha_default,
                        key=f"edit_fecha_personalizada_{ges}_{idx}"
                    )

                # Asegurarse de que 'fecha' sea un objeto datetime
                if isinstance(fecha, datetime):
                    fecha_formateada = fecha.strftime('%d-%m-%Y')
                elif isinstance(fecha, (pd.Timestamp, date)):
                    fecha_formateada = fecha.strftime('%d-%m-%Y')
                else:
                    st.warning("Fecha inválida. Se asignará la fecha de hoy.")
                    fecha = datetime.today()
                    fecha_formateada = fecha.strftime('%d-%m-%Y')

                responsable = st.text_input(
                    "Responsable",
                    value=df.at[medida_idx, 'Responsable'] if medida_idx is not None else "",
                    key=f"edit_responsable_{ges}_{idx}"
                )

                # Botón para enviar el formulario
                submit_button = st.form_submit_button(label="Confirmar selección o crear nueva medida")

            # Procesar la acción del formulario
            if submit_button:
                if medida_idx is not None:  # Editar medida existente
                    st.session_state[session_key].at[medida_idx, 'Medida'] = medida
                    st.session_state[session_key].at[medida_idx, 'Fecha monitoreo'] = fecha.strftime(
                        '%Y-%m-%d') if fecha else ''
                    st.session_state[session_key].at[medida_idx, 'Responsable'] = responsable
                    st.session_state[session_key].at[medida_idx, 'Seleccionada'] = True
                    st.success("Medida actualizada correctamente")
                else:  # Crear nueva medida
                    nueva_medida = {
                        "N°": len(st.session_state[session_key]) + 1,
                        "GES": ges,
                        "Dimensión": dimension['Dimensión en riesgo'],
                        "Medida": medida,
                        "Fecha monitoreo": fecha.strftime('%d-%m-%Y') if fecha else '',
                        "Responsable": responsable,
                        "Activo": True,
                        "Seleccionada": True
                    }
                    st.session_state[session_key] = pd.concat(
                        [st.session_state[session_key], pd.DataFrame([nueva_medida])],
                        ignore_index=True
                    )
                    st.success("Nueva medida creada correctamente")

    # Botón para guardar todas las interpretaciones
    if st.button("Guardar todas las interpretaciones"):
        for key, interpretacion in st.session_state['interpretaciones_temporales'].items():
            if key not in st.session_state:
                st.session_state[
                    key] = interpretacion  # Guardar en `st.session_state` solo si no ha sido instanciado por el widget
        st.success("Todas las interpretaciones se han guardado correctamente")

    # Nueva Sección: Resumen de datos confirmados
    st.header("5. Resumen de datos confirmados")

    confirmed_measures = []
    interpretaciones_data = []

    for ges, dimensiones in ges_groups.items():
        for idx, dimension in enumerate(dimensiones, 1):
            session_key = f"measures_{ges}_{idx}"
            interpretacion_key = f"interpretacion_{ges}_{idx}"

            # Procesar medidas confirmadas
            if session_key in st.session_state:
                temp_df = st.session_state[session_key].copy()
                temp_df = temp_df[temp_df['Seleccionada']]  # Filtrar solo medidas seleccionadas

                # Agregar información del GES y la Dimensión
                temp_df['Dimensión'] = dimension["Dimensión en riesgo"]
                temp_df['GES'] = ges

                # Añadir la interpretación correspondiente a todas las medidas de la dimensión
                if interpretacion_key in st.session_state['interpretaciones_temporales']:
                    temp_df['Interpretación'] = st.session_state['interpretaciones_temporales'][interpretacion_key]
                else:
                    temp_df['Interpretación'] = ""

                confirmed_measures.append(temp_df)

            # Procesar las interpretaciones de cada dimensión
            if interpretacion_key in st.session_state['interpretaciones_temporales']:
                interpretacion = st.session_state['interpretaciones_temporales'][interpretacion_key]
                interpretaciones_data.append({
                    'GES': ges,
                    'Dimensión': dimension["Dimensión en riesgo"],
                    'Interpretación': interpretacion
                })

    # Mostrar Resumen de Interpretaciones
    #if interpretaciones_data:
    #    interpretaciones_df = pd.DataFrame(interpretaciones_data)
    #    if not interpretaciones_df.empty:
    #        st.write("Las interpretaciones ingresadas hasta el momento:")
    #        st.dataframe(interpretaciones_df[['GES', 'Dimensión', 'Interpretación']])
    #else:
    #    st.info("No hay interpretaciones ingresadas hasta el momento.")

    if interpretaciones_data:
        interpretaciones_df = pd.DataFrame(interpretaciones_data)
        if not interpretaciones_df.empty:
            st.write("Las interpretaciones ingresadas hasta el momento:")

            # Agregar una columna de índice personalizado
            interpretaciones_df.insert(0, 'Índice', range(1, len(interpretaciones_df) + 1))  # Índice desde 1

            # Configurar opciones de la tabla
            gb = GridOptionsBuilder.from_dataframe(interpretaciones_df)
            gb.configure_default_column(
                wrapText=True,  # Ajusta el texto largo a varias líneas
                autoHeight=True,  # Ajusta la altura de la celda automáticamente
            )
            gb.configure_column("Índice", header_name="N°", width=30)  # Configurar columna de índice
            gb.configure_column("GES", header_name="GES", width=100)
            gb.configure_column("Dimensión", header_name="Dimensión", width=100)
            gb.configure_column("Interpretación", header_name="Interpretación", width=400)

            grid_options = gb.build()

            # Mostrar la tabla con AgGrid
            st.write("### Las interpretaciones ingresadas hasta el momento:")
            AgGrid(interpretaciones_df, gridOptions=grid_options, height=400, fit_columns_on_grid_load=True)
        else:
            st.info("No hay interpretaciones ingresadas hasta el momento.")
    else:
        st.info("No hay interpretaciones ingresadas hasta el momento.")




    # Mostrar Resumen de Medidas Confirmadas
    if confirmed_measures:
        confirmadas_df = pd.concat(confirmed_measures, ignore_index=True)
        if not confirmadas_df.empty:
            columnas_a_mostrar = [
                'Índice', 'GES', 'Dimensión', 'Medida',
                'Fecha monitoreo', 'Responsable'
            ]
            # Insertar el índice personalizado
            confirmadas_df.insert(0, 'Índice', range(1, len(confirmadas_df) + 1))  # Índice desde 1

            # Filtrar las columnas en el DataFrame
            confirmadas_df = confirmadas_df[columnas_a_mostrar]

            st.write("### Las siguientes medidas han sido confirmadas hasta el momento:")

            # Configurar las opciones de la tabla
            gb = GridOptionsBuilder.from_dataframe(confirmadas_df)
            gb.configure_default_column(
                wrapText=True,  # Ajusta el texto largo a varias líneas
                autoHeight=True,  # Ajusta la altura de la celda automáticamente
            )
            gb.configure_column("Índice", header_name="N°", width=30)  # Configurar columna de índice
            gb.configure_column("GES", header_name="GES", width=100)
            gb.configure_column("Dimensión", header_name="Dimensión", width=100)
            gb.configure_column("Medida", header_name="Medida", width=250)
            gb.configure_column("Fecha monitoreo", header_name="Fecha monitoreo", width=50)
            gb.configure_column("Responsable", header_name="Responsable", width=50)

            grid_options = gb.build()

            # Mostrar la tabla con AgGrid
            AgGrid(confirmadas_df, gridOptions=grid_options, height=400, fit_columns_on_grid_load=True)

        else:
            st.info("No hay medidas confirmadas hasta el momento.")
    else:
        st.info("No hay medidas confirmadas hasta el momento.")

    # Exportar como CSV las interpretaciones ingresadas
    if 'interpretaciones_df' in locals() and not interpretaciones_df.empty:
        csv_interpretaciones = interpretaciones_df.to_csv(index=False)
        st.download_button(
            label="Descargar archivo CSV con Interpretaciones",
            data=csv_interpretaciones,
            file_name="interpretaciones_ingresadas.csv",
            mime="text/csv",
        )
        st.success("Datos de interpretaciones guardados correctamente.")

    # Exportar como CSV
    if 'confirmadas_df' in locals() and not confirmadas_df.empty:
        csv = confirmadas_df.to_csv(index=False)
        #print(confirmadas_df)
        st.download_button(
            label="Descargar archivo CSV con Medidas Confirmadas",
            data=csv,
            file_name="medidas_seleccionadas.csv",
            mime="text/csv",
        )
        st.success("Datos de medidas guardados correctamente.")
    else:
        st.warning("No se han seleccionado medidas para guardar.")


    # Sección 6: Generación del informe en Word
    st.header("6. Generación del informe en Word")

    def establecer_orientacion_apaisada(doc):
        """
        Configura el documento en orientación horizontal (apaisado).
        """
        section = doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    def ges_contains_te3(ges_string, te3):
        ges_list = [ges.strip().lower() for ges in ges_string.split(';')]
        return te3.strip().lower() in ges_list



    def generar_contenido_word(datos, estado_riesgo, fig_principal, figs_te3, interpretaciones_df, summary_df, confirmadas_df, df_resultados_porcentaje):

        """
        Genera el contenido del informe en un objeto Document de python-docx.
        """
        # Crear un nuevo documento
        doc = Document()
        establecer_orientacion_apaisada(doc)

        # Establecer Calibri como fuente predeterminada para el estilo 'Normal'
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(9)  # Tamaño de fuente opcional; ajusta según prefieras

        # Crear un nuevo estilo llamado 'destacado' con Calibri y tamaño de fuente 12
        destacado = doc.styles.add_style('destacado', 1)  # 1 para párrafos
        destacado_font = destacado.font
        destacado_font.name = 'Calibri'
        destacado_font.size = Pt(12)  # Tamaño de la fuente en puntos

        # Configurar el idioma del documento en español
        lang = doc.styles['Normal'].element
        lang.set(qn('w:lang'), 'es-ES')

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        # Agregar imagen del logo (ajusta la ruta de la imagen a tu ubicación)
        doc.add_picture('IST.jpg', width=Inches(2))  # Ajusta el tamaño según sea necesario
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alinear al centro

        # Título principal
        titulo = doc.add_heading('INFORME TÉCNICO', level=1)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Subtítulo
        subtitulo = doc.add_heading('PRESCRIPCIÓN DE MEDIDAS PARA PROTOCOLO DE VIGILANCIA', level=2)
        subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Subtítulo
        subtitulo = doc.add_heading('DE RIESGOS PSICOSOCIALES EN EL TRABAJO', level=2)
        subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
        doc.add_paragraph()


        # Información general
        p = doc.add_paragraph()
        p.add_run('Razón Social: ').bold = True
        p.add_run(f"{safe_get(datos, 'Nombre_Empresa')}\n")
        p.add_run('RUT: ').bold = True
        p.add_run(f"{safe_get(datos, 'RUT_Empresa')}\n")
        p.add_run('Nombre del centro de trabajo: ').bold = True
        p.add_run(f"{safe_get(datos, 'Nombre_Centro_de_Trabajo')}\n")
        p.add_run('CUV: ').bold = True
        p.add_run(f"{safe_get(datos, 'CUV')}\n")
        p.add_run('CIIU: ').bold = True
        p.add_run(f"{safe_get(datos, 'CIIU')}\n")
        p.add_run('Fecha de activación del cuestionario: ').bold = True
        p.add_run(f"{safe_get(datos, 'Fecha_Inicio')}\n")
        p.add_run('Fecha de cierre del cuestionario: ').bold = True
        p.add_run(f"{safe_get(datos, 'Fecha_Fin')}\n")
        p.add_run('Universo de trabajadores de evaluación: ').bold = True
        p.add_run(f"{safe_get(datos, 'Nº_Trabajadores_CT')}\n")
        p.add_run('Nivel de riesgo: ').bold = True
        p.add_run(f"{safe_get(datos, 'Nivel_de_riesgo')}\n")
        p.paragraph_format.left_indent = Cm(1.5)

        # Salto de página
        doc.add_page_break()

        # Título de sección
        doc.add_heading('RESULTADOS GENERALES', level=2)

        # Información de riesgo general
        p = doc.add_paragraph()
        p.add_run('Nivel de riesgo: ').bold = True
        p.add_run(f"{estado_riesgo}\n")
        p.style.font.size = Pt(12)

        # Insertar imagen del gráfico principal
        if fig_principal:
            img_buffer = BytesIO()
            fig_principal.savefig(img_buffer, format='png')
            img_buffer.seek(0)
            doc.add_picture(img_buffer, width=Inches(6))
            img_buffer.close()
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Obtener dimensiones en riesgo
            dimensiones_riesgo_altog = df_resultados_porcentaje[
                (df_resultados_porcentaje['CUV'] == datos['CUV']) & (
                        df_resultados_porcentaje['Puntaje'] == 2)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_mediog = df_resultados_porcentaje[
                (df_resultados_porcentaje['CUV'] == datos['CUV']) & (
                        df_resultados_porcentaje['Puntaje'] == 1)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_bajog = df_resultados_porcentaje[
                (df_resultados_porcentaje['CUV'] == datos['CUV']) & (
                        df_resultados_porcentaje['Puntaje'] == -2)
                ]['Dimensión'].tolist()

            # Dimensiones en riesgo
            p = doc.add_paragraph()
            p.add_run('Dimensiones en riesgo alto: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_altog) if dimensiones_riesgo_altog else 'Ninguna'}\n")
            p.add_run('Dimensiones en riesgo medio: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_mediog) if dimensiones_riesgo_mediog else 'Ninguna'}\n")
            p.add_run('Dimensiones en riesgo bajo: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_bajog) if dimensiones_riesgo_bajog else 'Ninguna'}\n")


        # Agregar gráficos por TE3
        for fig_te3, te3 in figs_te3:
            # Salto de página
            doc.add_page_break()
            doc.add_heading(f"Resultado para el área o GES: {te3}", level=2)
            doc.add_paragraph()
            doc.add_paragraph()

            # Insertar gráfico
            img_buffer = BytesIO()
            fig_te3.savefig(img_buffer, format='png')
            img_buffer.seek(0)
            doc.add_picture(img_buffer, width=Inches(6))
            img_buffer.close()
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Obtener dimensiones en riesgo alto, medio y bajo
            dimensiones_riesgo_alto = df_porcentajes_niveles[
                (df_porcentajes_niveles['CUV'] == datos['CUV']) &
                (df_porcentajes_niveles['TE3'] == te3) &
                (df_porcentajes_niveles['Puntaje'] == 2)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_medio = df_porcentajes_niveles[
                (df_porcentajes_niveles['CUV'] == datos['CUV']) &
                (df_porcentajes_niveles['TE3'] == te3) &
                (df_porcentajes_niveles['Puntaje'] == 1)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_bajo = df_porcentajes_niveles[
                (df_porcentajes_niveles['CUV'] == datos['CUV']) &
                (df_porcentajes_niveles['TE3'] == te3) &
                (df_porcentajes_niveles['Puntaje'] == -2)
                ]['Dimensión'].tolist()

            # Dimensiones en riesgo alto, medio y bajo
            p = doc.add_paragraph()
            p.add_run('Dimensiones en riesgo alto: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_alto) if dimensiones_riesgo_alto else 'Ninguna'}\n")
            p.add_run('Dimensiones en riesgo medio: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_medio) if dimensiones_riesgo_medio else 'Ninguna'}\n")
            p.add_run('Dimensiones en riesgo bajo: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_bajo) if dimensiones_riesgo_bajo else 'Ninguna'}\n")

            # Interpretaciones por TE3
            def ges_contains_te3(ges_string, te3):
                ges_list = [ges.strip().lower() for ges in ges_string.split(';')]
                return te3.strip().lower() in ges_list

            # Filtrar interpretaciones
            interpretaciones_filtradas = interpretaciones_df[
                interpretaciones_df['GES'].apply(lambda x: ges_contains_te3(x, te3))
            ]

            interpretacion_list = [
                f"{row['Dimensión']}: {row['Interpretación']}"
                for _, row in interpretaciones_filtradas.iterrows()
                if pd.notna(row['Interpretación']) and row['Interpretación'].strip()
            ]

            # Agregar interpretaciones solo si existen
            if interpretacion_list:
                p = doc.add_paragraph()
                p.add_run('Interpretación del grupo de discusión:').bold = True
                p.add_run('\n')  # Salto de línea después del título

                # Agregar las interpretaciones una por línea
                for interpretacion in interpretacion_list:
                    p.add_run(interpretacion).add_break()  # Salto de línea después de cada interpretación

                doc.add_paragraph()  # Salto adicional entre secciones

            # Tabla de medidas propuestas
            doc.add_heading(f"Medidas propuestas para el área o GES: {te3}", level=3)
            doc.add_paragraph()

            # Filtrar confirmadas_df por el TE3 actual
            confirmadas_te3 = confirmadas_df[
                confirmadas_df['GES'].apply(lambda x: ges_contains_te3(x, te3))
            ]

            # Verificar si hay medidas propuestas
            if not confirmadas_te3.empty:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Dimensión'
                hdr_cells[1].text = 'Medida'
                hdr_cells[2].text = 'Fecha monitoreo'
                hdr_cells[3].text = 'Responsable'

                # Llenar la tabla con las medidas filtradas
                for _, row in confirmadas_te3.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row.get('Dimensión', 'N/A'))
                    row_cells[1].text = str(row.get('Medida', 'N/A'))
                    row_cells[2].text = str(row.get('Fecha monitoreo', 'N/A'))
                    row_cells[3].text = str(row.get('Responsable', 'N/A'))
            else:
                # Mensaje si no hay medidas propuestas
                doc.add_paragraph("No hay medidas propuestas para este área o GES.", style='Normal')

        # Retornar el objeto Document
        return doc


    def generar_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles, interpretaciones_df, confirmadas_df):
        if 'CUV' not in df_res_com.columns or 'CUV' not in summary_df.columns:
            st.error("La columna 'CUV' no se encuentra en los DataFrames proporcionados.")
            return None

        datos = df_res_com  # Ya filtrado por CUV
        estado = summary_df

        if datos.empty:
            st.error("No se encontraron datos para el CUV seleccionado.")
            return None
        if estado.empty:
            st.error("No se encontraron datos en 'summary_df' para el CUV seleccionado.")
            return None

        row = datos.iloc[0]
        estado_riesgo = estado['Riesgo'].values[0]

        # Llamadas sin CUV, asumiendo datos filtrados:
        fig_principal = generar_grafico_principal(df_resultados_porcentaje)
        if not fig_principal:
            st.warning("No se pudo generar el gráfico principal.")
            return None

        figs_te3 = generar_graficos_por_te3(df_porcentajes_niveles)

        doc = generar_contenido_word(row, estado_riesgo, fig_principal, figs_te3, interpretaciones_df, summary_df, confirmadas_df, df_resultados_porcentaje)

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return docx_buffer

    # Botón para generar y descargar el informe
    if st.button("Generar informe en Word"):

        if ('df_res_com' in st.session_state and 'summary_df' in st.session_state):
            with st.spinner("Generando el informe, por favor espera..."):
                # Generar el documento
                doc_buffer = generar_informe(st.session_state.df_res_com,
                                             st.session_state.summary_df,
                                             st.session_state.df_resultados_porcentaje,
                                             st.session_state.df_porcentajes_niveles,
                                             interpretaciones_df,
                                             confirmadas_df)

                if doc_buffer:
                    # Botón de descarga
                    st.download_button(
                        label="Descargar informe",
                        data=doc_buffer,
                        file_name=f"Informe_{st.session_state.selected_cuv}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.warning(
                "Los datos necesarios para generar el informe no están disponibles. Asegúrate de haber cargado todos los archivos requeridos.")

