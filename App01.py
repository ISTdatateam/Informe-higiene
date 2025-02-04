import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime


def generar_informe_en_word(df_filtrado: pd.DataFrame) -> BytesIO:
    """
    Genera un informe en Word (DOCX) a partir de las filas filtradas (df_filtrado).
    Retorna un objeto BytesIO con el contenido del archivo .docx.
    """

    doc = Document()

    # Título
    doc.add_heading("Informe de Evaluación Térmica - Medición de un área", 0)
    doc.add_paragraph(
        f"Fecha de generación del informe: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
    )
    doc.add_paragraph("Este documento contiene los datos recopilados para el CUV ingresado.")

    if df_filtrado.empty:
        doc.add_paragraph("No se encontraron registros para este CUV.")
    else:
        # Iteramos sobre cada fila
        for idx, row in df_filtrado.iterrows():
            doc.add_heading(f"Registro {idx + 1}", level=1)

            # Extrae datos (ajusta los nombres de columna a tu CSV)
            marca_temporal = row.get("Marca temporal", "")
            cuv = row.get("CUV", "")
            seccion = row.get("Que seccion quieres completar", "")
            fecha_visita = row.get("Fecha visita", "")
            # ... (añade aquí el resto de columnas que quieras)

            doc.add_paragraph(f"Marca temporal: {marca_temporal}")
            doc.add_paragraph(f"CUV: {cuv}")
            doc.add_paragraph(f"Sección: {seccion}")
            doc.add_paragraph(f"Fecha visita: {fecha_visita}")
            # ... (añade aquí el resto de columnas)

    doc.add_page_break()
    doc.add_paragraph("Fin del informe - Datos generados automáticamente.")

    # Exportar a BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def main():
    st.title("Formulario de Evaluación Térmica - SIN fotos")

    # URL de tu hoja en CSV
    csv_url = (
        "https://docs.google.com/spreadsheets/d/e/"
        "2PACX-1vTPdZTyxM6BDLmnlqe246tfBm7H06vXBdQKruh2mPg-rhQSD8olCS30ej4BdtJ1R__3W6K-Va3hm5Ax/"
        "pub?output=csv"
    )
    df = pd.read_csv(csv_url)

    # Para diagnosticar nombres de columna
    st.write("Columnas detectadas:", df.columns.tolist())

    # 1) Creamos un espacio en session_state para almacenar df_filtrado
    if "df_filtrado" not in st.session_state:
        st.session_state["df_filtrado"] = pd.DataFrame()

    # 2) Input y botón para filtrar
    input_cuv = st.text_input("Ingresa el CUV:")
    if st.button("Buscar"):
        # Normalizamos la columna 'CUV' y filtramos
        df["CUV"] = df["CUV"].astype(str).str.strip()
        input_cuv_str = str(input_cuv).strip()
        st.session_state["df_filtrado"] = df[df["CUV"] == input_cuv_str]

    # 3) Si tenemos datos filtrados en session_state, los mostramos
    if not st.session_state["df_filtrado"].empty:
        st.subheader("Datos Filtrados")
        st.dataframe(st.session_state["df_filtrado"])

        # 4) Botón para generar informe Word
        if st.button("Generar Informe en Word"):
            informe_docx = generar_informe_en_word(st.session_state["df_filtrado"])
            st.download_button(
                label="Descargar Informe",
                data=informe_docx,
                file_name="informe.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("Realiza la búsqueda de un CUV para mostrar datos y generar informe.")


if __name__ == "__main__":
    main()
