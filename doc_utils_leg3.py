import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from datetime import datetime
import requests
import re
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION


def get_or_add_lang(rPr):
    # Busca el elemento 'w:lang'
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        # Si no existe, lo creamos y lo añadimos a rPr
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    return lang


def descargar_imagen_gdrive(url_foto: str) -> BytesIO or None:
    """
    Dada una URL de Google Drive (tipo 'https://drive.google.com/open?id=ABC123'),
    intenta descargar el archivo y devolverlo como BytesIO.
    Retorna None si no se puede descargar o no tenga permisos adecuados.
    """
    match = re.search(r"id=(.*)", url_foto)
    if not match:
        return None

    file_id = match.group(1)
    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"

    try:
        response = requests.get(download_url)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            return None
    except Exception as e:
        return None


def look_informe(doc):
    """
    Configura el documento en orientación horizontal (apaisado) y establece
    los márgenes y estilos para todo el documento.
    """
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    # Estilo Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(9)  # Tamaño de fuente opcional; ajusta según prefieras

    # Estilo destacado
    destacado = doc.styles.add_style('destacado', 1)  # 1 para párrafos
    destacado_font = destacado.font
    destacado_font.name = 'Calibri'
    destacado_font.size = Pt(12)  # Tamaño de la fuente en puntos

    # Idioma Español de Chile
    rpr = style.element.get_or_add_rPr()
    lang = get_or_add_lang(rpr)
    lang.set(qn('w:val'), 'es-CL')
    lang.set(qn('w:eastAsia'), 'es-CL')
    lang.set(qn('w:bidi'), 'es-CL')


def set_vertical_alignment(doc, section_index=0, alignment='top'):
    """
    Ajusta la alineación vertical de la sección especificada.

    Parámetros:
    -----------
    doc : Document
        Objeto Document de python-docx.
    section_index : int
        Índice de la sección que se desea modificar. Por defecto 0 (la primera sección).
    alignment : str
        Valor de alineación vertical.
        Puede ser 'top' (arriba), 'center' (centrado) o 'both'.
    """
    # Obtenemos la sección
    section = doc.sections[section_index]
    # Obtenemos o creamos la propiedad sectPr
    sectPr = section._sectPr

    # Buscamos el elemento vAlign
    vAlign = sectPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        sectPr.append(vAlign)

    # Ajustamos el atributo w:val según el parámetro alignment
    alignment = alignment.lower()
    if alignment not in ['top', 'center', 'both']:
        alignment = 'top'  # Valor por defecto si no es válido

    vAlign.set(qn('w:val'), alignment)


def generar_informe_en_word(df_filtrado: pd.DataFrame,
                            df_info_cuv: pd.DataFrame) -> BytesIO:
    """
    Genera un informe en Word (DOCX) combinando:
      - df_info_cuv: Info del CUV (RUT, RAZÓN SOCIAL, etc.)
      - df_filtrado: Datos Generales / Mediciones del CSV principal
    con el siguiente formato:
      1) Fuente Calibri
      2) Tabla de 2 columnas para datos iniciales (A, B, C)
      3) Resumen en forma de tabla (6 columnas) para todas las mediciones
      4) Desarrollo detallado por cada área, también en tablas de 2 columnas
      5) Fotos embebidas
    """

    '''
    Código comentado (por ejemplo, para pruebas de otros formatos):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    # Agregar imagen del logo
    doc.add_picture('IST.jpg', width=Inches(2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Título principal y subtítulos
    titulo = doc.add_heading('INFORME TÉCNICO', level=1)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitulo = doc.add_heading('PRESCRIPCIÓN DE MEDIDAS PARA PROTOCOLO DE VIGILANCIA', level=2)
    subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitulo = doc.add_heading('DE RIESGOS PSICOSOCIALES EN EL TRABAJO', level=2)
    subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    doc.add_paragraph()

    # Información general (ejemplo)
    p = doc.add_paragraph()
    p.add_run('Razón Social: ').bold = True
    p.add_run(f"{safe_get(datos, 'Nombre_Empresa')}\n")
    # ... más información ...
    p.paragraph_format.left_indent = Cm(1.5)
    '''

    # ============================
    # Crea el documento Word
    # ============================
    doc = Document()
    look_informe(doc)
    set_vertical_alignment(doc, section_index=0, alignment='top')
    section = doc.sections[0]
    section.header_distance = Inches(0.5)
    header = section.header
    if header.paragraphs:
        paragraph = header.paragraphs[0]
    else:
        paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.add_picture('IST.jpg', width=Inches(0.8))  # Ajusta el tamaño según necesites

    # Encabezado principal del contenido
    paragraph = doc.add_heading("INFORME EVALUACIÓN CONFORT TÉRMICO", level=1)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()


    # ============================
    #   1) DATOS INICIALES
    # ============================
    # Creamos una tabla de 2 columnas para agrupar la información (secciones A, B y C)
    table_inicial = doc.add_table(rows=0, cols=2)
    table_inicial.style = 'Table Grid'

    def add_row(table, label, value=""):
        """
        Agrega una fila a la tabla.
        Si 'value' es vacío se asume que se trata de un título; se fusionan ambas celdas
        y se aplica fondo morado con letras blancas.
        """
        row_cells = table.add_row().cells

        if value == "":
            # Fusionar ambas celdas para títulos
            merged_cell = row_cells[0].merge(row_cells[1])
            merged_cell.text = label

            # Aplicar fondo morado (hex "800080")
            shading_elm = parse_xml(r'<w:shd {} w:fill="800080"/>'.format(nsdecls('w')))
            merged_cell._tc.get_or_add_tcPr().append(shading_elm)

            # Cambiar el color de fuente a blanco, negrita y aumentar el tamaño
            for paragraph in merged_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.size = Pt(12)
        else:
            row_cells[0].text = label
            value_str = "" if pd.isna(value) else str(value)
            row_cells[1].text = value_str

    # --- Sección A y B: Información empresa y centro de trabajo ---
    if not df_info_cuv.empty:
        row_cuv_info = df_info_cuv.iloc[0]
        # A. Información empresa
        add_row(table_inicial, "Información empresa")
        add_row(table_inicial, "Razón Social", row_cuv_info.get('RAZÓN SOCIAL', ''))
        add_row(table_inicial, "RUT", row_cuv_info.get('RUT', ''))

        # B. Información centro de trabajo
        add_row(table_inicial, "Información centro de trabajo")
        add_row(table_inicial, "CUV", row_cuv_info.get('CUV', ''))
        add_row(table_inicial, "Nombre de Local", row_cuv_info.get('Nombre de Local', ''))
        add_row(table_inicial, "Dirección", row_cuv_info.get('Dirección', '') + ", " +  row_cuv_info.get('Comuna',''))
    else:
        add_row(table_inicial, "Información de empresa/centro", "No se encontró información para este CUV.")

    # --- Separamos los registros del DF principal ---
    datos_generales = df_filtrado[df_filtrado["Que seccion quieres completar"] == "Datos generales"]
    mediciones = df_filtrado[df_filtrado["Que seccion quieres completar"] == "Medición de un area"]


    # Encabezado principal del contenido
    doc.add_paragraph()
    paragraph = doc.add_heading("Resumen", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    if not datos_generales.empty:
        row_gen = datos_generales.iloc[0]
        doc.add_paragraph("Efectuadas mediciones de confort térmico en el local 'Nombre de Local', es posible concluir que las áreas de 'Area o sector' , cumplen con el estándar de confort térmico, por lo cual se debe mantener las condiciones actuales o similares.")
        doc.add_paragraph("Respecto de las áreas de 'Area o sector', no cumple con el estándar de confort térmico, por lo cual hace necesario adoptar las medidas prescritas detalladas en la tabla 2 para su solución.")



    doc.add_paragraph()
    paragraph = doc.add_heading("Antecedentes de la actividad ", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    if not df_info_cuv.empty:
        row_cuv_info = df_info_cuv.iloc[0]

        # Extraer los datos desde la fila utilizando el nombre de las columnas
        razon_social = row_cuv_info.get("RAZÓN SOCIAL", "")
        rut = row_cuv_info.get("RUT", "")
        cuv = row_cuv_info.get("CUV", "")
        nombre_local = row_cuv_info.get("Nombre de Local", "")
        direccion = row_cuv_info.get("Dirección", "")
        comuna = row_cuv_info.get("Comuna", "")
        direccion_completa = f"{direccion}, comuna de {comuna}"

        # Otros datos que necesites incluir en el texto (estos pueden venir de otra parte de tu código)
        fecha_visita = "20/01/2025"
        nombre_consultor = "IST"
        cargo_empresa = "[Nombre persona empresa]"

        # Mostrar el texto generado
        doc.add_paragraph(f"A solicitud de {razon_social} se realiza una evaluación de confort térmico (calor) en la sucursal {nombre_local}, ubicada en {direccion_completa}.")
        doc.add_paragraph(f"La visita fue realizada el día {fecha_visita} por el consultor {nombre_consultor} en compañía de {cargo_empresa}.")
        doc.add_paragraph("Las áreas evaluadas fueron: Línea de cajas, bodega, etccc.")

    doc.add_paragraph()
    paragraph = doc.add_heading("Metodología de las mediciones y evaluaciones", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    if not datos_generales.empty:
        row_gen = datos_generales.iloc[0]
        doc.add_paragraph("Las condiciones del centro de trabajo se midieron utilizado los equipos Monitor de estrés térmico Equipo 1 y Anemómetro Equipo 2, los cuales cuentan con sus certificados vigentes. (anexo xx)")
        doc.add_paragraph("La medición se realizó utilizando la metodología de FANGER para evaluación de confort térmico en espacios interiores de acuerdo a Nota técnica N°47 del Instituto de Salud Pública")
        doc.add_paragraph("Se utiliza el estándar de vestimenta … y la tasa mets para las actividades realizadas (anexo xxxx) o,5 Clo")





    table_calib = doc.add_table(rows=0, cols=2)
    table_calib.style = 'Table Grid'
    # -- Sección C: Antecedentes generales (en la misma tabla de 2 columnas) --
    if not datos_generales.empty:
        row_gen = datos_generales.iloc[0]
        # Título para la sección de detalles de equipos y calibración
        add_row(table_calib, "D. Detalles de equipos y calibración")
        add_row(table_calib, "Equipo temperatura", row_gen.get('Código equipo temperatura', ''))
        add_row(table_calib, "Equipo velocidad viento", row_gen.get('Codigo equipo 2', ''))

        # Se añaden más detalles relacionados
        add_row(table_calib, "Tipo de vestimenta utilizada", row_gen.get('Tipo de vestimenta utilizada', ''))
        add_row(table_calib, "Motivo de evaluación", row_gen.get('Motivo de evaluación', ''))
        add_row(table_calib, "Comentarios finales de evaluación", row_gen.get('Comentarios finales de evaluación', ''))


    # ============================
    #   2) RESUMEN DE MEDICIONES
    # ============================
    if not mediciones.empty:
        cols_resumen = [
            "Area o sector",
            "Especificación sector",
            "Temperatura bulbo seco",
            "Temperatura globo",
            "Humedad relativa",
            "Velocidad del aire"
        ]

        doc.add_heading("Resultados de mediciones y evaluación  ", level=2)
        doc.add_paragraph()

        #tabla_resumen = doc.add_table(rows=1, cols=len(cols_resumen))
        tabla_resumen = doc.add_table(rows=1, cols=10)
        tabla_resumen.style = 'Table Grid'
        hdr_cells = tabla_resumen.rows[0].cells

        hdr_cells[0].text = "Área de medición"
        hdr_cells[1].text = "Sector"
        hdr_cells[2].text = "Hora de medición"
        hdr_cells[3].text = "Temperatura bulbo seco(°C)"
        hdr_cells[4].text = "Temperatura globo(°C)"
        hdr_cells[5].text = "Humedad relativa (%)"
        hdr_cells[6].text = "Velocidad del aire(m/s)"
        hdr_cells[7].text = "PPD"
        hdr_cells[8].text = "PMV"
        hdr_cells[9].text = "Estándar de confortabilidad térmica PMV [-1,+1]"

        for _, row_med in mediciones.iterrows():
            row_cells = tabla_resumen.add_row().cells
            row_cells[0].text = str(row_med.get("Area o sector", ""))
            row_cells[1].text = str(row_med.get("Especificación sector", ""))
            row_cells[2].text = str(row_med.get("Area o sector", ""))
            row_cells[3].text = str(row_med.get("Temperatura bulbo seco", ""))
            row_cells[4].text = str(row_med.get("Temperatura globo", ""))
            row_cells[5].text = str(row_med.get("Humedad relativa", ""))
            row_cells[6].text = str(row_med.get("Velocidad del aire", ""))
            row_cells[7].text = str(row_med.get("Area o sector", ""))
            row_cells[8].text = str(row_med.get("Area o sector", ""))
            row_cells[9].text = str(row_med.get("Area o sector", ""))
        doc.add_paragraph("")  # Espacio extra

    else:
        doc.add_paragraph("No se han encontrado filas de 'Medición de un area' para este CUV.")

    doc.add_paragraph()
    doc.add_paragraph("PPD: Es la predicción cuantitativa del porcentaje de personas que se sentirían insatisfechas que sienten mucho calor o mucho frío en un medioambiente determinado.")
    doc.add_paragraph("PMV: Índice que estima el valor medio de las respuestas sobre la sensación térmica que emitiría un grupo de personas sometidas al mismo ambiente")

    doc.add_paragraph()
    paragraph = doc.add_heading("Conclusiones", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    if not datos_generales.empty:
        row_gen = datos_generales.iloc[0]
        doc.add_paragraph(
            "De la evaluación de confort térmico realizada en las diferentes áreas del local xx de SMU, cuyos resultados se entregan en la tabla 1, se puede concluir lo siguiente:")
        doc.add_paragraph(
            " [Línea de caja y Oficina] de gerencia] , cumplen con el estándar de confort térmico, por lo cual se debe mantener las condiciones actuales o similares.")
        doc.add_paragraph(
            "Respecto [de las área] de [Bodega] no cumple con el estándar de confort térmico, por lo cual hace necesario adoptar las medidas prescritas para su solución.")

    doc.add_paragraph()
    paragraph = doc.add_heading("Prescripciones de medidas ", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    doc.add_paragraph("Medidas de ingeniería")
    tabla_resumen = doc.add_table(rows=1, cols=3)
    tabla_resumen.style = 'Table Grid'
    hdr_cells = tabla_resumen.rows[0].cells

    hdr_cells[0].text = "Área de medición"
    hdr_cells[1].text = "Sector"
    hdr_cells[2].text = "Hora de medición"

    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_paragraph("Medidas de ingeniería")
    tabla_resumen = doc.add_table(rows=1, cols=len(cols_resumen))
    tabla_resumen.style = 'Table Grid'
    hdr_cells = tabla_resumen.rows[0].cells

    hdr_cells[0].text = "Área de medición"
    hdr_cells[1].text = "Sector"
    hdr_cells[2].text = "Hora de medición"


    doc.add_paragraph()
    paragraph = doc.add_heading("Vigencia del informe y notas", level=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    if not datos_generales.empty:
        row_gen = datos_generales.iloc[0]
        doc.add_paragraph(
            "El presente informe tiene una vigencia de 3 años, en la medida que no cambien las condiciones.")
        doc.add_paragraph(
            "Estos resultados de evaluación representan las condiciones existentes del ambiente y lugar de trabajo al momento de realizar las mediciones.")


    doc.add_paragraph()
    doc.add_page_break()



    # ============================
    #   3) DESARROLLO POR ÁREA
    # ============================

    if not mediciones.empty:
        for idx, row_med in mediciones.iterrows():
            heading_anexoareas = doc.add_heading("Anexo de condiciones observadas en áreas de medición", level=2)
            heading_anexoareas.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_heading(
                f"Área: {row_med.get('Area o sector', '')} - {row_med.get('Especificación sector', '')}",
                level=2
            )

            tabla_area = doc.add_table(rows=0, cols=2)
            tabla_area.style = 'Table Grid'

            def add_area_row(label, value):
                r = tabla_area.add_row().cells
                value_str = "" if pd.isna(value) else str(value)
                r[0].text = label
                r[1].text = value_str

            add_area_row("Puesto de trabajo", row_med.get("Puesto de trabajo", ""))
            add_area_row("Trabajador de pie o sentado", row_med.get("Trabajador de pie o sentado", ""))
            add_area_row("Techumbre", row_med.get("Techumbre", ""))
            add_area_row("Observación techumbre", row_med.get("Observacion techumbre - Indique tipo de material", ""))
            add_area_row("Paredes", row_med.get("Paredes", ""))
            add_area_row("Observación paredes", row_med.get("Observacion paredes", ""))
            add_area_row("Ventanales", row_med.get("Ventanales", ""))
            add_area_row("Observación ventanales", row_med.get("Observacion ventanales", ""))
            add_area_row("Aire acondicionado", row_med.get("Aire acondicionado", ""))
            add_area_row("Observaciones aire acondicionado", row_med.get("Observaciones aire acondicionado", ""))
            add_area_row("Ventiladores", row_med.get("Ventiladores", ""))
            add_area_row("Observaciones ventiladores", row_med.get("Observaciones ventiladores", ""))
            add_area_row("Inyección y/o extracción de aire", row_med.get("Inyección y/o extracción de aire", ""))
            add_area_row("Observaciones inyección y/o extracción de aire",
                         row_med.get("Observaciones inyeccion y/o extracción de aire", ""))
            add_area_row("Ventanas (Ventilación natural)", row_med.get("Ventanas (Ventilación natural)", ""))
            add_area_row("Observaciones ventanas (ventilación natural)",
                         row_med.get("Observaciones ventanas (ventilación natural)", ""))
            add_area_row("Puertas (ventilación natural)", row_med.get("Puertas (ventilación natural)", ""))
            add_area_row("Observaciones puertas (ventilación natural)",
                         row_med.get("Observaciones puertas (ventilación natural)", ""))
            add_area_row("Otras condiciones de disconfort térmico",
                         row_med.get("Otras condiciones de disconfort termico", ""))
            add_area_row("Observaciones sobre otras condiciones de disconfort térmico",
                         row_med.get("Observaciones sobre otras condiciones de disconfort térmico", ""))

            # Fotografías asociadas
            doc.add_paragraph()
            doc.add_paragraph()

            evidencia = row_med.get("Evidencia fotografica", "")
            if pd.notnull(evidencia) and isinstance(evidencia, str) and evidencia.strip():
                doc.add_paragraph("Fotografías asociadas:")
                for foto_url in [f.strip() for f in evidencia.split(",") if f.strip()]:
                    imagen = descargar_imagen_gdrive(foto_url)
                    if imagen is not None:
                        doc.add_picture(imagen, width=Inches(4))
                    else:
                        doc.add_paragraph(f"No se pudo descargar la imagen: {foto_url}")
            doc.add_paragraph("")  # Espacio extra
            doc.add_page_break()
    else:
        doc.add_paragraph("No se encontró información de mediciones para detallar.")

    # ============================
    #   4) APARTADO DETALLES TÉCNICOS DE LA EVALUACIÓN
    # ============================

    heading_anexotec = doc.add_heading("Anexo de detalles técnicos de evaluación", level=2)
    heading_anexotec.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if not datos_generales.empty:
        row_gen = datos_generales.iloc[0]
        # Se crea una tabla de 4 columnas para el nuevo formato:
        table_tec = doc.add_table(rows=0, cols=4)
        table_tec.style = 'Table Grid'
        # Fila de cabecera
        hdr_cells = table_tec.add_row().cells
        hdr_cells[0].text = "Verificación"
        hdr_cells[1].text = "Patrón equipo"
        hdr_cells[2].text = "Inicio medición"
        hdr_cells[3].text = "Final medición"
        # Fila para Valor TBS
        row = table_tec.add_row().cells
        row[0].text = "Valor TBS"
        row[1].text = str(row_gen.get('Patrón TBS', ''))
        row[2].text = str(row_gen.get('Verificación TBS inicial', ''))
        row[3].text = str(row_gen.get('Verificación TBS final', ''))
        # Fila para Valor TBH
        row = table_tec.add_row().cells
        row[0].text = "Valor TBH"
        row[1].text = str(row_gen.get('Patrón TBH', ''))
        row[2].text = str(row_gen.get('Verificación TBH inicial', ''))
        row[3].text = str(row_gen.get('Verificación TBH final', ''))
        # Fila para Valor TG
        row = table_tec.add_row().cells
        row[0].text = "Valor TG"
        row[1].text = str(row_gen.get('Patrón TG', ''))
        row[2].text = str(row_gen.get('Verificación TG inicial', ''))
        row[3].text = str(row_gen.get('Verificación TG final', ''))
    else:
        # En caso de no encontrar datos generales, se crea una tabla informativa
        table_tec = doc.add_table(rows=0, cols=4)
        table_tec.style = 'Table Grid'
        hdr_cells = table_tec.add_row().cells
        hdr_cells[0].text = "Verificación"
        hdr_cells[1].text = "Patrón equipo"
        hdr_cells[2].text = "Inicio medición"
        hdr_cells[3].text = "Final medición"
        row = table_tec.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = "No se han encontrado información para este CUV."
        row[3].text = ""

    # Convertir el documento a BytesIO para descargar
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
