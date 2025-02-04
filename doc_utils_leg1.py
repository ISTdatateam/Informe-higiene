import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime
import requests
import re

def descargar_imagen_gdrive(url_foto: str) -> BytesIO or None:
    """
    Dada una URL de Google Drive (tipo 'https://drive.google.com/open?id=ABC123'),
    intenta descargar el archivo y devolverlo como BytesIO.
    Retorna None si no se puede descargar o no tenga permisos adecuados.
    """
    match = re.search(r"id=(.*)", url_foto)
    if not match:
        return None

    file_id = match.group(1)
    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"

    try:
        response = requests.get(download_url)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            return None
    except:
        return None


def generar_informe_en_word(df_filtrado: pd.DataFrame) -> BytesIO:
    """
    Genera un informe en Word (DOCX) a partir de las filas filtradas (df_filtrado).
    Separa la información en dos bloques:
      1) Datos Generales
      2) Medición de un área (puede haber múltiples entradas para el mismo CUV)
    Incluye inserción de fotografías desde enlaces de Google Drive.
    """

    # Separar los registros según 'Que seccion quieres completar'
    datos_generales = df_filtrado[df_filtrado["Que seccion quieres completar"] == "Datos generales"]
    mediciones = df_filtrado[df_filtrado["Que seccion quieres completar"] == "Medición de un area"]

    doc = Document()

    # --- Portada ---
    doc.add_heading("Informe de Evaluación Térmica", 0)
    doc.add_paragraph(
        f"Fecha de generación del informe: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
    )
    doc.add_paragraph(
        "El presente informe reúne la información recopilada en terreno, "
        "diferenciando los datos generales y las mediciones de diferentes áreas "
        "del lugar evaluado. A continuación, se presentan los resultados."
    )

    # --- 1) DATOS GENERALES ---
    doc.add_heading("1. Datos Generales", level=1)
    if not datos_generales.empty:
        for idx, row in datos_generales.iterrows():
            doc.add_heading(f"Registro de Datos Generales #{idx+1}", level=2)

            doc.add_paragraph(f"Marca temporal: {row.get('Marca temporal', '')}")
            doc.add_paragraph(f"CUV: {row.get('CUV', '')}")
            doc.add_paragraph(f"Fecha visita: {row.get('Fecha visita', '')}")
            doc.add_paragraph(f"Hora medición: {row.get('Hora medicion', '')}")
            doc.add_paragraph(f"Temperatura máxima del día: {row.get('Temperatura máxima del día', '')}")
            doc.add_paragraph(f"Nombre del personal SMU: {row.get('Nombre del personal SMU', '')}")
            doc.add_paragraph(f"Cargo: {row.get('Cargo', '')}")
            doc.add_paragraph(f"Código equipo 1: {row.get('Código equipo 1', '') or row.get('Codigo equipo 1','')}")
            doc.add_paragraph(f"Código equipo 2: {row.get('Código equipo 2', '') or row.get('Codigo equipo 2','')}")
            doc.add_paragraph(f"Verificación TBS patrón A: {row.get('Verificación TBS patrón A', '')}")
            doc.add_paragraph(f"Verificación TBH patrón A: {row.get('Verificación TBH patrón A', '')}")
            doc.add_paragraph(f"Verificación TG patrón A: {row.get('Verificación TG patrón A', '')}")
            doc.add_paragraph(f"Verificación TBS patrón B: {row.get('Verificación TBS patrón B', '')}")
            doc.add_paragraph(f"Verificación TBH patrón B: {row.get('Verificación TBH patrón B', '')}")
            doc.add_paragraph(f"Verificación TG patrón B: {row.get('Verificación TG patrón B', '')}")
            doc.add_paragraph(f"Dirección de correo electrónico: {row.get('Dirección de correo electrónico','')}")

            doc.add_paragraph(" ")  # Espacio en blanco
    else:
        doc.add_paragraph("No se han encontrado filas de 'Datos Generales' para este CUV.")

    # --- 2) MEDICIONES DE UN ÁREA ---
    doc.add_heading("2. Mediciones de un Área", level=1)
    if not mediciones.empty:
        for idx, row in mediciones.iterrows():
            doc.add_heading(f"Medición #{idx+1}", level=2)

            doc.add_paragraph(f"Marca temporal: {row.get('Marca temporal', '')}")
            doc.add_paragraph(f"CUV: {row.get('CUV', '')}")
            doc.add_paragraph(f"Área o sector: {row.get('Area o sector','')}")
            doc.add_paragraph(f"Especificación sector: {row.get('Especificación sector','')}")
            doc.add_paragraph(f"Puesto de trabajo: {row.get('Puesto de trabajo','')}")
            doc.add_paragraph(f"Techumbre: {row.get('Techumbre','')}")
            doc.add_paragraph(
                f"Observacion techumbre: {row.get('Observacion techumbre - Indique tipo de material','')}"
            )
            doc.add_paragraph(f"Paredes: {row.get('Paredes','')}")
            doc.add_paragraph(f"Observacion paredes: {row.get('Observacion paredes','')}")
            doc.add_paragraph(
                f"Temperatura bulbo seco (°C): {row.get('Temperatura bulbo seco (°C)','')}"
            )
            doc.add_paragraph(f"Temperatura globo (°C): {row.get('Temperatura globo (°C)','')}")
            doc.add_paragraph(f"Humedad relativa (%): {row.get('Humedad relativa (%)','')}")
            doc.add_paragraph(
                f"Velocidad del aire (m/s): {row.get('Velocidad del aire (m/s)','')}"
            )
            doc.add_paragraph(
                f"Trabajador de pie o sentado: {row.get('Trabajador de pie o sentado','')}"
            )

            # -- Nuevas columnas de ventilación / confort térmico --
            doc.add_paragraph(f"Ventanales: {row.get('Ventanales','')}")
            doc.add_paragraph(f"Observación ventanales: {row.get('Observacion ventanales','')}")
            doc.add_paragraph(f"Aire acondicionado: {row.get('Aire acondicionado','')}")
            doc.add_paragraph(f"Observaciones aire acondicionado: {row.get('Observaciones aire acondicionado','')}")
            doc.add_paragraph(f"Ventiladores: {row.get('Ventiladores','')}")
            doc.add_paragraph(f"Observaciones ventiladores: {row.get('Observaciones ventiladores','')}")
            doc.add_paragraph(f"Inyección y/o extracción de aire: {row.get('Inyección y/o extracción de aire','')}")
            doc.add_paragraph(f"Observaciones inyeccion y/o extracción de aire: {row.get('Observaciones inyeccion y/o extracción de aire','')}")
            doc.add_paragraph(f"Ventanas (Ventilación natural): {row.get('Ventanas (Ventilación natural)','')}")
            doc.add_paragraph(f"Observaciones ventanas (ventilación natural): {row.get('Observaciones ventanas (ventilación natural)','')}")
            doc.add_paragraph(f"Puertas (ventilación natural): {row.get('Puertas (ventilación natural)','')}")
            doc.add_paragraph(f"Observaciones puertas (ventilación natural): {row.get('Observaciones puertas (ventilación natural)','')}")
            doc.add_paragraph(f"Otras condiciones de disconfort termico: {row.get('Otras condiciones de disconfort termico','')}")
            doc.add_paragraph(
                f"Observaciones sobre otras condiciones de disconfort térmico: {row.get('Observaciones sobre otras condiciones de disconfort térmico','')}"
            )

            # --- Fotografías ---
            evidencia = row.get("Evidencia fotografica", "")
            if pd.notnull(evidencia) and isinstance(evidencia, str) and evidencia.strip():
                doc.add_paragraph("Fotografías asociadas:")
                lista_fotos = [link.strip() for link in evidencia.split(",") if link.strip()]
                for foto_url in lista_fotos:
                    imagen = descargar_imagen_gdrive(foto_url)
                    if imagen is not None:
                        doc.add_picture(imagen, width=Inches(4))
                    else:
                        doc.add_paragraph(f"No se pudo descargar la imagen: {foto_url}")

            doc.add_paragraph(" ")  # Separador visual
    else:
        doc.add_paragraph("No se han encontrado filas de 'Medición de un area' para este CUV.")

    # --- Cierre ---
    doc.add_page_break()
    doc.add_paragraph("Fin del informe - Datos generados automáticamente.")
    doc.add_paragraph(
        "Nota: La información y observaciones incluidas en este documento "
        "se basan en los registros recopilados durante la visita al lugar "
        "y pueden estar sujetas a validación adicional."
    )

    # Guardar el documento en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
